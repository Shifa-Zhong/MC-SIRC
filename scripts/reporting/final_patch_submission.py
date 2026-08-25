#!/usr/bin/env python3
"""Apply final terminology synchronization after the revision-5 document build."""

from __future__ import annotations

from pathlib import Path
import re
import shutil
import subprocess
import sys

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.reporting import finalize_revision5_submission as finalizer


def replace_everywhere(document, old, new):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    count = 0
    for paragraph in paragraphs:
        if old not in paragraph.text:
            continue
        for run in paragraph.runs:
            if old in run.text:
                count += run.text.count(old)
                run.text = run.text.replace(old, new)
        if old in paragraph.text:
            text = paragraph.text.replace(old, new)
            paragraph.text = text
            count += 1
    return count


def patch_si():
    path = finalizer.PAPER / "SI_Final.docx"
    document = Document(path)
    replacements = {
        "S4 event upper": "S4 enhanced gap (+50%)",
        "S4 event-weighted upper": "S4 enhanced gap (+50%)",
    }
    for old, new in replacements.items():
        replace_everywhere(document, old, new)
    for table in document.tables:
        finalizer.submission.apply_three_line_table(table)
    document.save(path)
    shutil.copy2(path, finalizer.ARCHIVE / path.name)


def patch_repository_tag():
    for name in ["Manuscript_Final_MC-SIRC.docx", "SI_Final.docx", "cover letter.docx"]:
        path = finalizer.PAPER / name
        document = Document(path)
        paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            normalized = re.sub(
                r"revision-2026-08-(?:24|25(?:\.\d+)*)",
                finalizer.TAG,
                paragraph.text,
            )
            if normalized != paragraph.text:
                paragraph.text = normalized
        document.save(path)
        shutil.copy2(path, finalizer.ARCHIVE / path.name)


def patch_multipanel_captions():
    targets = {
        "Manuscript_Final_MC-SIRC.docx": [
            ("Figure 3.", finalizer.FIGURE3_CAPTION),
            ("Figure 4.", finalizer.FIGURE4_CAPTION),
        ],
        "SI_Final.docx": [
            ("Figure S1.", finalizer.FIGURES1_CAPTION),
        ],
    }
    for name, replacements in targets.items():
        path = finalizer.PAPER / name
        document = Document(path)
        for prefix, caption in replacements:
            finalizer.replace_prefix(document, prefix, caption)
        document.save(path)
        shutil.copy2(path, finalizer.ARCHIVE / path.name)


def main():
    finalizer.ARCHIVE.mkdir(parents=True, exist_ok=True)
    finalizer.postprocess_manuscript()
    finalizer.build_standalone_abstract()
    finalizer.postprocess_si()
    patch_si()
    finalizer.postprocess_cover_and_ga()
    finalizer.build_exact_response()
    finalizer.prepare_submission_artwork()
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "generate_word_level_highlights_v2.py")], check=True)
    for name in [
        "Manuscript_Final_MC-SIRC.docx", "SI_Final.docx", "Abstract.docx", "Response_Letter.docx",
        "SI_Highlighted_20260824.docx", "Manuscript_Highlighted_20260824.docx",
        "Abstract_Highlighted_20260824.docx", "GA.docx", "cover letter.docx", "highlights.docx",
        "declarationStatement.docx", "Figure_Captions.docx",
    ]:
        path = finalizer.PAPER / name
        if path.exists():
            finalizer.prune_docx_media(path)
            finalizer.prune_empty_docx_comments(path)
    print("Final independent-acceptance repairs complete")


if __name__ == "__main__":
    main()
