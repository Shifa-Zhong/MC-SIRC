#!/usr/bin/env python3
"""
正文与 SI 全文检查 + 润色 + 强化创新点表达
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from pathlib import Path
import shutil

ROOT = Path('D:/shanxi')
MS = ROOT/'paper/Manuscript_Final_MC-SIRC.docx'
SI = ROOT/'paper/SI_Final.docx'

shutil.copy2(MS, MS.with_suffix('.docx.polish_bak'))
shutil.copy2(SI, SI.with_suffix('.docx.polish_bak'))
print("备份: *.docx.polish_bak")

def replace_in_paragraph(para, old, new):
    """Replace text in paragraph while preserving first run formatting."""
    full_text = ''.join(run.text for run in para.runs)
    if old in full_text:
        new_text = full_text.replace(old, new)
        runs = list(para.runs)
        if runs:
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ''
        return True
    return False

def find_and_replace(doc, edits):
    """Apply each edit (old, new) to first matching paragraph."""
    applied = 0
    for old, new in edits:
        for para in doc.paragraphs:
            if replace_in_paragraph(para, old, new):
                applied += 1
                print(f"  ✓ {old[:55]}{'...' if len(old)>55 else ''}")
                break
        else:
            print(f"  ✗ NOT FOUND: {old[:55]}{'...' if len(old)>55 else ''}")
    return applied

# ════════════════════════════════════════════════════════════════════════
# Manuscript edits
# ════════════════════════════════════════════════════════════════════════
print("\n" + "="*80)
print("MANUSCRIPT 编辑")
print("="*80)
ms = docx.Document(str(MS))

# ─── A. 摘要润色（一并改 ten unknowns 与方法描述） ───
print("\n[A] 摘要润色")
abstract_edits = [
    # 1. 方法部分重写: 三段式 + 评级系统
    ('MC-SIRC combines Bayesian MAP estimation with full MCMC posterior sampling—'
     'resolving the severely underdetermined single-section problem (one constraint, '
     'ten unknowns per pollutant)—with a monthly-calibrated spatial distance-decay '
     'model that translates river-entry loads into effective contributions at the outlet.',

     'MC-SIRC integrates three reinforcing components: Bayesian MAP estimation with '
     'full MCMC posterior sampling to resolve the severely underdetermined single-section '
     'problem (one constraint, 9–10 unknowns per pollutant); a five-scenario prior '
     'sensitivity matrix paired with explicit A/B/C reliability ratings that separate '
     'data-driven anomalies from prior-influenced estimates; and a monthly-calibrated '
     'spatial distance-decay model that translates river-entry loads into effective '
     'contributions at the outlet.'),

    # 2. 强化"三线证据收敛"
    ('Bayesian posteriors and five-scenario prior sensitivity analyses convergently '
     'identify systematic overestimation of large-scale livestock coefficients, '
     'a finding robust to prior choice.',

     'Three independent lines of evidence—Bayesian MAP, MCMC posteriors, and prior '
     'sensitivity scenarios—converge to identify systematic overestimation of '
     'large-scale livestock coefficients, with the TP correction factor binding to '
     'the lower bound 0.100 across all five prior scenarios.'),

    # 3. 把"4 倍放大 / 5 倍衰减"的反演显化（在已存在的 centralized facility 那一句之后）
    ('while a centralized treatment facility ranked fourth in COD emissions (6.8%) '
     'rises to first in effective contribution (26.8%).',

     'while a centralized treatment facility ranked fourth in COD emissions (6.8%) '
     'rises to first in effective contribution (26.8%)—a four-fold amplification for '
     'the proximate facility (~7 km from outlet) versus a five-fold attenuation for '
     'distant livestock farms (30–50 km).'),

    # 4. 透明性表达更尖锐
    ('Strong fits for COD (R² = 0.81) and TN (R² = 0.96) contrast with inadequate '
     'performance for NH₃-N and TP, indicating limits of simple exponential decay '
     'for biogeochemically active species—a limitation we report transparently '
     'rather than mask.',

     'Strong fits for COD (R² = 0.81) and TN (R² = 0.96) contrast with inadequate '
     'performance for NH₃-N and TP (R² < 0.45)—a limitation we report transparently '
     'and trace to the unsuitability of simple exponential decay for biogeochemically '
     'active species and particulate-phase transport.'),
]
find_and_replace(ms, abstract_edits)

# ─── B. 引言 P15: 强化"四个互联贡献"的表达 ───
print("\n[B] 引言 P15: 四点贡献强化")
intro_edits = [
    ('The methodology presented here makes four interrelated contributions. '
     'We extend the inventory framework beyond river-entry estimation to achieve '
     'the first quantitative closure with downstream monitoring data, converting '
     'inventories into independently testable load predictions. To resolve '
     'single-section underdetermination, we combine Bayesian MAP estimation with '
     'full MCMC posterior sampling, obtaining transparent uncertainty bounds for '
     'every source-coefficient pair rather than point estimates alone. '
     'We calibrate the spatial decay model against monthly rather than annual loads, '
     'which produces meaningful goodness-of-fit statistics and reveals a systematic '
     'divergence between emission rankings and effective water-quality contributions. '
     'Throughout, we report results with explicit reliability ratings and openly '
     'identify cases of inadequate model fit, enabling readers and managers to '
     'distinguish data-driven findings from prior-influenced ones.',

     'The MC-SIRC framework presented here makes four interrelated methodological '
     'contributions. First, we extend the inventory paradigm beyond river-entry '
     'estimation to achieve—within the inventory-based modelling tradition—the first '
     'quantitative closure with downstream monitoring data, converting static '
     'inventory products into independently testable load predictions. Second, '
     'to resolve the single-section underdetermination problem, we triangulate '
     'three independent lines of evidence (Bayesian MAP, full MCMC posteriors, and '
     'a five-scenario prior sensitivity matrix), enabling each source coefficient '
     'to be classified as data-driven (Rating A), substantially corrected (Rating B), '
     'or prior-influenced (Rating C). Third, we calibrate the spatial decay model '
     'against monthly rather than annual loads, which yields meaningful goodness-of-fit '
     'statistics and—translated through channel attenuation—reveals a systematic '
     'divergence between emission rankings and effective water-quality contributions '
     'that has not previously been quantified within an inventory-based workflow. '
     'Fourth, we report cases of inadequate model fit transparently rather than mask '
     'them, equipping managers and reviewers to discriminate robust findings from '
     'speculative ones and providing a template for accountable inventory-monitoring '
     'integration in data-limited watersheds.'),
]
find_and_replace(ms, intro_edits)

# ─── C. P42 删除孤立的 "Note: E = emission..." (Table 2 脚注遗留) ───
print("\n[C] 删除 Table 2 残留脚注")
for i, p in enumerate(ms.paragraphs):
    text = p.text.strip()
    if text == 'Note: E = emission; R = river-entry load. "—" indicates data not available.' or \
       text == 'Note: E = emission; R = river-entry load. “—” indicates data not available.':
        # Clear paragraph
        runs = list(p.runs)
        for r in runs:
            r.text = ''
        print(f"  ✓ 已清空 Para {i} (孤立脚注)")
        break

# ─── D. P49 (§3.2): unknowns 数量 ───
print("\n[D] §3.2 P49: unknowns 数量")
find_and_replace(ms, [
    ('All prediction deviations were within ±12%, validating Bayesian robustness '
     'under the severely underdetermined conditions (1 equation, 10 unknowns per pollutant).',
     'All prediction deviations were within ±12%, validating Bayesian robustness '
     'under the severely underdetermined conditions (1 equation, 9–10 unknowns per pollutant).'),
])

# ─── E. P63 (§3.3): rainfall-weighted → flow-weighted ───
print("\n[E] §3.3 P63: rainfall → flow weighting")
find_and_replace(ms, [
    ('the exponential decay with rainfall-weighted monthly allocation cannot capture '
     'complex nitrogen cycling dynamics',
     'the exponential decay with flow-weighted monthly allocation cannot capture '
     'complex nitrogen cycling dynamics'),
])

# ─── F. §3.3 P66: 中心发现段 — 强化机制 + 影响 ───
print("\n[F] §3.3 P66: 中心发现段强化")
find_and_replace(ms, [
    ('This pattern — a four-fold amplification for the centralized facility versus '
     'a five-fold attenuation for livestock — is consistent with the distance-decay '
     'framework described in SPARROW applications [18–20] and underscores the importance '
     'of incorporating spatial transport in pollution source management [13].',

     'This pattern — a four-fold amplification for the proximate centralized facility '
     'versus a five-fold attenuation for distant livestock — emerges from two '
     'independent mechanisms acting in concert: a high river-entry coefficient '
     '(α = 1.0 for direct WWTP discharge vs. α = 0.309 for livestock farms whose '
     'effluent must traverse soil and surface pathways) and proximity-driven exponential '
     'survival of in-stream load (exp(−0.073 × 7) = 0.60 for the centralized facility '
     'vs. exp(−0.073 × 40) ≈ 0.05 for distant livestock under the COD calibration). '
     'The result is consistent with the distance-decay framework underlying SPARROW '
     'applications [18–20] but quantifies, for the first time within an inventory-based '
     'workflow, the magnitude by which spatial information can invert the policy-relevant '
     'ranking of pollution sources [13].'),
])

# ─── G. §3.5 P72: ten unknowns ───
print("\n[G] §3.5 P72: ten unknowns")
find_and_replace(ms, [
    ('which yields only one mass-balance equation per pollutant against ten unknowns.',
     'which yields only one mass-balance equation per pollutant against 9–10 unknowns.'),
])

# ─── H. §4 结论 P79–82: 整段重写, 更尖锐 + 行动指向 ───
print("\n[H] §4 Conclusions: 4 段重写")
conclusions_edits = [
    # (1)
    ('(1) The MC-SIRC framework presented here achieves the first quantitative closure '
     'between emission inventories and water quality monitoring within an inventory-based '
     'paradigm. Comprehensive transport efficiency from emission to monitoring section '
     'ranged from 6% (TP) to 55% (TN), with channel attenuation as the primary loss mechanism.',

     '(1) The MC-SIRC framework demonstrates the first quantitative closure of the '
     '"emission → river-entry → monitoring" chain within an inventory-based paradigm, '
     'converting inventory products from open-loop estimates into monitoring-anchored '
     'predictions. Comprehensive transport efficiency from emission to outlet ranges '
     'from 6% (TP, particulate-settling-dominated) to 55% (TN, conservative behavior), '
     'with channel attenuation as the primary loss mechanism — an order-of-magnitude '
     'pollutant-specific signal that any management framework based on emission totals '
     'alone is structurally unable to capture.'),

    # (2)
    ('(2) Bayesian MAP optimization with full MCMC posterior sampling and five-scenario '
     'prior sensitivity analysis provided robust identification of key uncertainty '
     'sources through three-way convergence: MAP z-score = 2.33 for both COD and TP, '
     'MCMC 95% CI: 0.100–0.152 for TP large livestock, and correction factor stable '
     'at 0.100 across S2–S5 for both pollutants and across all five scenarios for '
     'TP large livestock specifically. This constitutes definitive evidence for systematic '
     'coefficient overestimation in large-scale livestock farming.',

     '(2) Triangulation across Bayesian MAP, MCMC posteriors, and a five-scenario prior '
     'sensitivity matrix converges on a single empirical conclusion: large-scale livestock '
     'production-emission coefficients are systematically overestimated. Convergent evidence '
     'includes a MAP z-score of 2.33 for both COD and TP, a tight MCMC 95% CI of 0.100–0.152 '
     'for the TP large-livestock factor, and stable binding at the lower bound 0.100 across '
     'S2–S5 (and across all five scenarios for TP large-livestock specifically). The '
     'corresponding A/B/C reliability rating (4 / 5 / 25 of 34 source-pollutant pairs) '
     'translates Bayesian uncertainty into actionable triage guidance — a feature absent '
     'from existing inventory frameworks.'),

    # (3)
    ('(3) Monthly-calibrated spatial attenuation modeling achieved credible fits for '
     'COD (R² = 0.812) and TN (R² = 0.958), while transparently reporting model inadequacy '
     'for NH₃-N and TP (R² < 0.45). Point sources dominated effective contributions '
     '(53%–88%), revealing that major emitters and major water quality contributors '
     'diverge substantially when spatial transport is considered.',

     '(3) Monthly-calibrated spatial decay modeling achieves credible fits for COD '
     '(R² = 0.81) and TN (R² = 0.96) while transparently reporting model inadequacy for '
     'NH₃-N and TP (R² < 0.45) and tracing it to seasonal nitrification-denitrification '
     'dynamics and particulate settling that a constant first-order rate cannot represent. '
     'For pollutants that the model captures, point sources dominate effective contributions '
     '(53–88% across pollutants); the centralized facility, ranking only fourth in COD '
     'emissions, becomes the largest effective contributor (26.8%) due to proximity (~7 km) '
     'and direct discharge (α = 1.0), while large-livestock farming — the dominant emitter — '
     'falls to fifth in effective contribution due to distance (30–50 km from outlet).'),

    # (4)
    ('(4) The separation between emission-based and impact-based management was '
     'demonstrated quantitatively. A differentiated strategy of "near-source: prioritize '
     'river-entry control; far-source: prioritize emission reduction" is recommended, '
     'supported by explicit reliability ratings for management application.',

     '(4) The quantitative inversion of source rankings between emission and effective-'
     'contribution scales — a four-fold amplification for the centralized facility versus '
     'a five-fold attenuation for large livestock — supports a differentiated control '
     'strategy: near-source priority on river-entry coefficient reduction (wastewater '
     'treatment upgrades, riparian industrial control), far-source priority on emission '
     'reduction (livestock production-emission factor verification and farming-scale '
     'recensus). Explicit reliability ratings let managers act on data-driven anomalies '
     'without conflating them with prior-driven artifacts, advancing the field from total '
     'emission control toward precision contribution control.'),
]
find_and_replace(ms, conclusions_edits)

ms.save(str(MS))
print(f"\n✓ Manuscript 保存: {MS}")

# ════════════════════════════════════════════════════════════════════════
# SI edits
# ════════════════════════════════════════════════════════════════════════
print("\n" + "="*80)
print("SI 编辑")
print("="*80)
si = docx.Document(str(SI))

# ─── SI 概述更新 ───
print("\n[SI 1] 内容概述更新")
find_and_replace(si, [
    ('This Supporting Information contains 26 tables (Tables S1–S26), 1 figure '
     '(Figure S1), and 4 supplementary method descriptions (S1–S4).',
     'This Supporting Information contains 29 tables (Tables S1–S29), 4 figures '
     '(Figures S1–S4), and 4 supplementary method descriptions (S1–S4).'),
])

# ─── SI Table S5 题注: 月度过滤口径 ───
print("\n[SI 2] Table S5 footnote: 月度过滤口径")
find_and_replace(si, [
    ('Note: Aug–Nov had <210 h/month and were excluded from spatial model calibration '
     '(8 months retained).',
     'Note: Aug–Nov had <50% data coverage (records <360 h vs. 720–744 h expected) '
     'and were excluded from spatial-model calibration (8 months retained).'),
])

# ─── SI 标题统一: 与 manuscript 一致 ───
print("\n[SI 3] SI 标题加 MC-SIRC 框架名")
find_and_replace(si, [
    ('Quantitative Full-Chain Analysis of Watershed Water Pollution Based on Emission '
     'Inventory and Water Quality Monitoring: A Case Study of the Nanchuan River Basin',

     'Supporting Information for: MC-SIRC: A Monitoring-Constrained Source Inventory '
     'Reverse Calibration Framework for Full-Chain Watershed Pollution Analysis—'
     'A Case Study of the Nanchuan River Basin, Loess Plateau'),
])

# ─── SI 作者顺序与 manuscript 一致 ───
print("\n[SI 4] SI 作者顺序")
find_and_replace(si, [
    ('WANG Yujie, ZHONG Shifa*, LUO Jinhong, ZHANG Weifeng, ZHANG Cheng',
     'Yujie Wang, Jinhong Luo, Weifeng Zhang, Cheng Zhang, Shifa Zhong*'),
])

si.save(str(SI))
print(f"\n✓ SI 保存: {SI}")

print("\n" + "="*80)
print("全文润色完成。备份: *.docx.polish_bak")
print("="*80)
