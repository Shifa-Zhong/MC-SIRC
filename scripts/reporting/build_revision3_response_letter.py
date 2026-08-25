#!/usr/bin/env python3
"""Populate all Reply: xxx placeholders in the latest reviewer-comments file."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
ARCHIVE_DUP = PAPER / "_archive_non_submission_20260824" / "duplicate_dated_versions"
BLUE = RGBColor(0x00, 0x70, 0xC0)


def insert_before(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if text:
        result.add_run(text)
    return result


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)


COMMON = {
    "abstract": (
        "Revised MS (Abstract):",
        "Study focus: We reformulate Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) as an identifiability-aware workflow. Missing-not-at-random gaps are represented by four annual-load scenarios; Bayesian maximum a posteriori and Markov chain Monte Carlo analyses provide prior-regularized discrepancy allocations; Jacobian-rank, null-space, and k = 0 profile diagnostics determine which quantities the data can support."
    ),
    "rank": (
        "Revised MS (§2.6 and §3.2):",
        "For Eq. (2), the Jacobian with respect to the active source factors and U has one row. Its rank, nullity, and identifiable linear combinations were calculated directly. The annual Jacobian has rank 1 for every pollutant, with nullity 8 for COD and TN and 9 for NH₃-N and TP. Thus, only one aggregate linear combination is observed."
    ),
    "tier": (
        "Revised MS (§2.5):",
        "The standardized prior shift z = |fMAP − μ|/σ is retained as a prior-conflict diagnostic: Tier A, z > 2; Tier B, 1 < z ≤ 2; Tier C, z ≤ 1. These tiers do not measure identifiability or reliability. Tier A means only that the regularized reconciliation is in strong conflict with the stated prior and therefore identifies a field-audit priority."
    ),
    "missing": (
        "Revised MS (§2.2):",
        "The 2022 archive contains 5,928 unique hourly timestamps from a possible 8,760, giving 67.7% coverage and 2,832 absent hours. Missingness is concentrated in July–November and is therefore treated as missing not at random (MNAR). Four annual-load scenarios were evaluated: S1 observed hours only; S2 month-specific coverage scaling (revised default); S3 annual-mean scaling by 8,760/5,928; and S4 an event-weighted upper sensitivity that multiplies only the S2-imputed portion by 1.5."
    ),
    "crossyear": (
        "Revised MS (§2.6 and §3.2):",
        "The same annual-mean protocol was also applied to 2020, 2021, and 2023 monitoring archives with the 2022 inventory fixed. This is labelled a cross-year consistency check rather than validation: the factors necessarily absorb year-specific hydrology and unobserved source changes. The 2020–2021 archives have only 1.0–1.3% paired hourly coverage and are treated as data-poor."
    ),
    "spatial": (
        "Revised MS (§3.3):",
        "Once γ is profiled without the former lower/upper bounds, the distance-decay profiles are essentially flat relative to k = 0. The incremental ΔR² is 0.00003 for COD, 0.00019 for NH₃-N, 0.00333 for TN, and 0.00069 for TP. These are constraint and scaling effects, not independently estimated attenuation rates."
    ),
    "withdraw": (
        "Revised MS (§3.3):",
        "Because k is not identified, the previously reported half-life distances, effective-contribution shares, point-source percentages, and emission-to-effective ranking inversions are withdrawn. Distance and nominal river-entry coefficients remain useful descriptors for designing monitoring locations, but they do not support quantitative intervention rankings in this dataset."
    ),
    "management": (
        "Revised MS (§3.4):",
        "Operational measures should therefore not be re-ranked from the present factors. Any near-source/far-source strategy is treated as a testable field hypothesis, to be evaluated with intervention or multi-section data before policy use."
    ),
    "transfer": (
        "Revised MS (§3.5):",
        "MC-SIRC is therefore presented as a methodological prototype and audit template, not a demonstrated transferable calibration."
    ),
    "repro": (
        "Revised MS (Data availability) and SI Text S8:",
        "The exact public snapshot used as the revision baseline is permanently addressable at https://github.com/Shifa-Zhong/MC-SIRC/tree/126fa0a789751ecf0850b4971244f26cd404a880. Input schemas, parameter settings, fixed random seeds, and commands for regenerating the main tables and figures are listed in SI Text S8 and the repository reproduction guide."
    ),
}


RESPONSES = [
    # Reviewer 1: general assessment, strengths, eight numbered comments, overall.
    (
        "We thank Reviewer #1 for recognizing the importance of connecting source inventories with downstream monitoring. The revision preserves that objective but changes the inferential claim substantially: MC-SIRC is now an identifiability-aware reconciliation and audit workflow, not a source-coefficient validation framework. We added formal rank/null-space and k = 0 diagnostics and removed every conclusion that those diagnostics do not support.",
        [COMMON["abstract"]],
    ),
    (
        "We appreciate this balanced summary of the original strengths. We retained the transparent uncertainty reporting and inventory–monitoring linkage, while revising the terminology so that MAP, MCMC, prior sensitivity, and Monte Carlo are no longer presented as independent validation. The A/B/C system is retained only as prior-conflict screening.",
        [COMMON["tier"]],
    ),
    (
        "We agree. The earlier wording overstated what one section can identify. We now report the Jacobian rank and nullity, construct exactly equivalent source-factor solutions, and explicitly state that individual factors are prior-regularized allocations. Claims of coefficient validation have been removed throughout.",
        [COMMON["rank"], COMMON["tier"]],
    ),
    (
        "We agree and have corrected both the method and the internal inconsistency. The correct reference year contains 8,760 hours, not 8,705; with 5,928 unique records, 2,832 hours are absent. We replaced the former single default with four coherent scenarios. The four original boundary conflicts persist across all scenarios, but other tier counts and some regularized rankings change, so we no longer use the latter for decisions.",
        [COMMON["missing"], ("Revised SI (Table S26 note):", "All four original boundary conflicts remain at 0.100 under S1–S4, whereas other tier assignments and some regularized component rankings change.")],
    ),
    (
        "We agree. The analysis is now called a cross-year consistency check. We emphasize that the fixed 2022 inventory and sparse 2020–2021 archives preclude independent validation and that factors absorb interannual hydrological and source variability.",
        [COMMON["crossyear"]],
    ),
    (
        "We went beyond limiting interpretation to NH₃-N and TP. The new k = 0 profile audit shows that k is not independently identifiable for any pollutant once γ is profiled without artificial bounds. We therefore withdrew all pollutant-specific half-lives and effective-contribution estimates, including the former COD/TN ranking result.",
        [COMMON["spatial"], COMMON["withdraw"]],
    ),
    (
        "We agree. Management statements are now framed as audit and monitoring hypotheses. The four persistent boundary conflicts identify where field verification is informative, but no source is operationally re-ranked and no direct policy prescription is made from the regularized factors.",
        [COMMON["management"]],
    ),
    (
        "We agree. Claims of portability or transferability have been removed. The manuscript now states that one Loess Plateau basin demonstrates an audit prototype only; cross-basin transfer requires new data and a fresh identifiability analysis.",
        [COMMON["transfer"]],
    ),
    (
        "We added a concise algorithmic table immediately after the workflow description. It lists each component's input, output, key assumption, main uncertainty, and supported inferential role, allowing readers to distinguish accounting, regularization, forward uncertainty, and identifiability diagnostics.",
        [("Revised MS (Table 1):", "MC-SIRC components, assumptions, uncertainties, and supported inferential roles. Supported outputs are limited to the bottom-up baseline, observed aggregate constraint, prior-conflict screening, inventory-side uncertainty, and formal estimability diagnostics.")],
    ),
    (
        "We strengthened reproducibility with a stable commit-specific URL, explicit file structure, parameter-table references, fixed seeds, executable commands for the new diagnostics and figures, and an anonymized example schema for restricted inputs.",
        [COMMON["repro"]],
    ),
    (
        "We appreciate the summary and have addressed each listed focus directly. Most importantly, the revised paper does not preserve unsupported conclusions by changing their wording: it formally demonstrates the limitations and withdraws the spatial ranking and coefficient-validation claims.",
        [("Revised MS (Conclusion):", "The principal contribution is an identifiability-aware workflow that converts disagreement between inventory and monitoring into transparent aggregate diagnostics and a concrete monitoring roadmap. Multi-section, seasonally resolved, and process-specific data are prerequisites for source-level calibration and transfer beyond this case basin.")],
    ),

    # Reviewer 2: recommendation and 11 comments.
    (
        "We thank Reviewer #2 for identifying the structural issue at the center of the former manuscript. We agree that the previously reported k, half-life distances, effective-contribution shares, and management re-ranking were unsupported. We have removed those conclusions and rebuilt the paper around formal identifiability auditing. The revised contribution is therefore different and narrower: it reconciles aggregate inventory and monitoring scales, identifies prior conflicts for field audit, and diagnoses what additional observations are required. To keep the revised account self-contained without restoring unsupported inference, we retain the names of the inventory categories and their directly calculated emission composition.",
        [COMMON["withdraw"], COMMON["rank"], ("Revised MS (§2.2 and §3.1):", "The non-point categories are rural domestic, agricultural cultivation, aquaculture, urban NPS, and dispersed urban domestic emissions; point sources include household-livestock farms, large-scale livestock farms, industrial sources, and one centralized treatment facility. At the source-emission scale, large-scale livestock accounts for 56.6% of COD and 65.1% of TP emissions, whereas the centralized facility accounts for 37.1% of TN emissions. These are inventory-composition shares, not estimates of source contributions at the outlet.")],
    ),
    (
        "We agree with the algebraic criticism. The former manuscript stated that a common flow allocation was applied uniformly. A code audit found assumed agricultural and aquaculture calendars, but all point sources and most inventory mass share the discharge profile; those assumed calendars are not independent source observations. We now give the exact reduction L̂m = γwmA(k), analytically profile γ, and demonstrate that the apparent optimum disappears. All k-dependent outputs are withdrawn.",
        [("Revised MS (§2.8):", "If wi,m = wm for all sources, Eq. (3) reduces exactly to L̂m = γwmA(k), with A(k) = ΣEiαi exp(−kdi). Only the product γA(k) is identifiable, so any k can be offset by γ."), COMMON["spatial"]],
    ),
    (
        "We agree. We added the requested k = 0 flow-only null. With γ profiled without its former bounds, COD and TN improve by only 0.00003 and 0.00333 in R², respectively; NH₃-N and TP show similarly negligible increments. The former high R² values are no longer described as validation of distance decay.",
        [COMMON["spatial"]],
    ),
    (
        "We agree. The causal attribution of f = 0.10 to upstream production-emission coefficients has been withdrawn. f is now defined as a discrepancy allocation that conflates inventory bias, channel transformation, omitted sources, and compensation. Persistent boundary conflict justifies a field audit, not a causal conclusion.",
        [("Revised MS (§3.4):", "This is a hypothesis-generation step; low regularized factors are equally compatible with inventory bias, unrepresented channel loss, and compensating model terms.")],
    ),
    (
        "We agree. MAP, MCMC, and prior sensitivity are now explicitly described as complementary summaries of one likelihood–prior system, not independent evidence. We also state that the narrow TP interval is piled against the 0.10 bound and cannot be interpreted as strong data information.",
        [("Revised MS (§3.2):", "The former narrow MCMC interval for large-livestock TP (0.100–0.152 under S3) is piled against the 0.10 truncation bound. Its narrowness is consequently evidence of boundary conflict within the stated model, not strong data information.")],
    ),
    (
        "We agree and have retired the interpretation of z as identifiability. A/B/C are now prior-conflict tiers only. We added formal diagnostics: rank 1, nullity 8–9, and constructive alternative factor vectors that yield exactly the same aggregate load.",
        [COMMON["tier"], COMMON["rank"], ("Revised SI (Table S33):", "The paired allocations preserve the regularized source sum to numerical precision while changing individual factors substantially; non-uniqueness is structural.")],
    ),
    (
        "We agree. The ±12% language has been changed from validation/robustness to constructional reconciliation or bookkeeping residual. The manuscript explains that small residuals follow from prior-regularized flexibility and are not predictive performance.",
        [("Revised MS (§3.1):", "The prior-regularized reconciliation remains close to each target by construction. The residual deviations of +2.0% to +12.0% describe the chosen observation-error and prior trade-off, not predictive validation.")],
    ),
    (
        "We agree. We corrected the 8,705/8,760 discrepancy and replaced the inconsistent descriptions with one explicit scenario framework. S2 now uses month-specific coverage; S1, S3, and S4 bound the sensitivity. The four original boundary conflicts persist, but other tier assignments and some rankings change, which is now reported rather than hidden.",
        [COMMON["missing"], ("Revised SI (Text S1):", "The corrected 2022 factor is 8,760/5,928 = 1.4777; the former 8,705-hour denominator was erroneous.")],
    ),
    (
        "We agree. The section, tables, and captions now use 'consistency check' rather than validation or reproducibility confirmation. The fixed-inventory and coverage limitations are stated in Methods, Results, SI, and Conclusions.",
        [COMMON["crossyear"]],
    ),
    (
        "We agree. The Monte Carlo analysis is now named forward inventory uncertainty. It is explicitly pre-transport and assumption-dependent; P(entry > outlet) is retained only as a scale comparison and is not used as evidence of inventory overestimation. We report the leading elasticities as supported inventory-side sensitivity results while explicitly separating them from outlet source contributions.",
        [("Revised MS (§2.7 and §3.2):", "The output is a distribution of nominal river-entry load before in-stream transport. Inventory-side elasticities identify large-scale livestock as the most influential input for nominal COD (0.391), NH₃-N (0.364), and TP (0.528) river-entry totals, and the centralized facility for TN (0.704). These elasticities quantify sensitivity of the pre-transport inventory total, not source contributions at the outlet.")],
    ),
    (
        "We agree. The 18-parameter urban-NPS secondary optimization and its 75% concentration reductions have been withdrawn from the main manuscript and supporting results. Functional-zone weights remain only as descriptive inventory inputs.",
        [("Revised SI (Text S2):", "Urban NPS functional-zone weights are retained as descriptive inventory inputs. The former 18-parameter secondary optimization is withdrawn because it introduced no new observation and could not identify zone-specific concentrations.")],
    ),
    (
        "We agree. Because k is not identified, every quantitative effective-contribution percentage and near/far control recommendation has been removed, including the former COD/TN claims. The paper now proposes only field-audit and monitoring-design hypotheses.",
        [COMMON["withdraw"], COMMON["management"]],
    ),

    # Reviewer 3: general, four main comments, three SI comments.
    (
        "We thank Reviewer #3 for the constructive assessment. We revised the scientific positioning, literature synthesis, figures, process discussion, conclusion length, and SI formatting. The new identifiability diagnostics also lead us to a more conservative interpretation than the prior version.",
        [COMMON["abstract"]],
    ),
    (
        "We expanded the Introduction from a general inventory–monitoring disconnect to the specific identifiability and MNAR bottlenecks. We added recent representative work on spatial source/transport modelling, Bayesian multi-tributary phosphorus export, upper-Yangtze phosphorus tracing, and adaptive dynamic apportionment. We also restored foundational references for inventory construction, river-network transport, Bayesian environmental inference, and uncertainty assessment, while keeping references tied only to withdrawn conclusions out of the paper. The research gap and four objectives are now stated explicitly.",
        [("Revised MS (Introduction):", "Export-coefficient and spatially integrated inventory approaches provide the source-category and spatial detail needed for refined management (Johnes, 1996; Chen et al., 2014; Shen et al., 2012; Strokal et al., 2016). Bayesian methods have been increasingly applied to water-environment parameter estimation because they incorporate prior knowledge, quantify parameter uncertainty, and remain computationally stable in underdetermined settings (Qian et al., 2003; Borsuk et al., 2004; Reckhow, 2003; Gelman et al., 2013)."), ("Revised MS (§2.2):", "The inventory compilation and spatial-accounting protocol are described by Wang et al. (2026).")],
    ),
    (
        "We completely redrew Figure 2 with a larger layout and non-overlapping labels; panel c now separates scenarios with short multiline labels and distinct markers. Figure 3 was enlarged to a 2 × 2 layout with substantially larger axes, titles, and legends. Its caption was also corrected to state that the distributions are pre-transport.",
        [("Revised MS (Figure 2 caption):", "Identifiability-aware Bayesian reconciliation. A/B/C tiers are redefined as strong/moderate/no resolved prior conflict; they do not denote data identifiability."), ("Revised MS (Figure 3 caption):", "Forward Monte Carlo uncertainty in nominal river-entry loads before in-stream transport.")],
    ),
    (
        "We added a process-focused Loess Plateau discussion covering particulate TP settling and storm remobilization, nitrification, denitrification, oxygen and residence-time controls, and the contrast between flashy-event transport and low-flow retention. Foundational river-retention and integrated modelling references have been restored. We use this discussion to explain why a constant k is physically incomplete while emphasizing that the current outlet record cannot estimate a more complex process model.",
        [("Revised MS (§4.1):", "Flashy hydrographs shorten residence time during events but increase particulate transport, whereas low-flow periods favor contact with sediments and microbial processing (Alexander et al., 2000; Behrendt and Opitz, 2000; Wollheim et al., 2006; Wollheim et al., 2008). Matching process detail to available observations is a recurring challenge in integrated water-quality modelling (Rode et al., 2010).")],
    ),
    (
        "We condensed the Conclusions to two paragraphs. Detailed numerical results remain in Results and the SI; the Conclusions now contain only the supported aggregate finding, formal limitation, withdrawal of unsupported spatial outputs, and monitoring implication.",
        [("Revised MS (Conclusion):", "The spatial audit reaches the same boundary: after profiling γ, a k = 0 flow-only model is effectively indistinguishable from the best decay profile (ΔR² ≤ 0.0033). Accordingly, half-life distances, effective-contribution shares, ranking inversions, and direct policy prescriptions have been removed.")],
    ),
    (
        "We bolded the Table S1 title. Table S17 headers were shortened and set not to split within words. All long tables now use repeated header rows on continuation pages and rows are prevented from splitting across pages, implementing a consistent continued-table convention. We also added a concise main-text navigation sentence that cites every supplementary text, figure, and table in ascending ranges before later item-specific references.",
        [("Revised MS (§2.3):", "Supplementary methods, diagnostic figures, and detailed tables are organized in Texts S1–S8, Figures S1–S2, and Tables S1–S35, respectively."), ("Revised SI formatting:", "Table S1. Overview of research data (data type, content, resolution, and volume). Table S17 headers: Rural; Agri. NPS; HH livest.; Aqua.; Urban NPS; Disp. urban; Large livest.; Industrial; Central fac.; Unknown.")],
    ),
    (
        "We standardized every dagger marker programmatically. The dagger is now attached directly to the number with no intervening or trailing spaces in all SI tables, including the former Tables S6 and S33.",
        [("Revised SI formatting:", "Boundary-marker convention: 5.000† (no space before or after †).")],
    ),
    (
        "We added an abbreviation line before the first SI method or table occurrence, spelling out MAP and MCMC as requested and also defining NPS and the four pollutant abbreviations.",
        [("Revised SI (front matter):", "Abbreviations: maximum a posteriori (MAP); Markov chain Monte Carlo (MCMC); non-point source (NPS); chemical oxygen demand (COD); ammonia nitrogen (NH₃-N); total nitrogen (TN); total phosphorus (TP).")],
    ),

    # AE.
    (
        "We reformatted the manuscript to the Journal of Hydrology: Regional Studies convention. The abstract now has the required three labelled sections (Study region, Study focus, and New hydrological insights for the region); citations use author–year style; the reference list is alphabetical; and the manuscript uses double spacing and continuous line numbering. Figure and table captions and the separate structured Abstract file were updated consistently.",
        [("Revised MS (Abstract headings):", "Study region: …\nStudy focus: …\nNew hydrological insights for the region: …"), ("Revised MS format:", "Author–year citations; alphabetical references; double-spaced text; continuous line numbering.")],
    ),
]


def populate_reply(paragraph, response, excerpts):
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    lead = paragraph.add_run("Reply: ")
    lead.bold = True
    paragraph.add_run(response)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    for label, excerpt in excerpts:
        paragraph.add_run().add_break()
        paragraph.add_run().add_break()
        run = paragraph.add_run(label)
        run.bold = True
        run.font.color.rgb = BLUE
        paragraph.add_run().add_break()
        run = paragraph.add_run(excerpt)
        run.font.color.rgb = BLUE


def main():
    source = PAPER / "_archive_non_submission_20260824" / "review_source" / "comments.docx"
    doc = Document(source)
    placeholders = [p for p in doc.paragraphs if p.text.strip().lower() == "reply: xxx"]
    if len(placeholders) != len(RESPONSES):
        raise RuntimeError(f"Expected {len(RESPONSES)} placeholders, found {len(placeholders)}")

    first = doc.paragraphs[0]
    p = insert_before(first, "Response to the Editor and Reviewers")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    p = insert_before(first, "EJRH-D-26-02074")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = insert_before(first, "Manuscript: MC-SIRC: An Identifiability-Aware Workflow for Reconciling Watershed Source Inventories with Outlet Monitoring—A Case Study of the Nanchuan River Basin, Loess Plateau")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = insert_before(first, "Dear Editor and Reviewers,")
    p = insert_before(first, "We thank you for the careful and technically incisive reviews. The central revision is substantive rather than rhetorical: we formally audited identifiability, corrected the missing-data scenarios, withdrew unsupported k-dependent and source-validation claims, and repositioned MC-SIRC as an aggregate reconciliation and monitoring-design workflow. All revised MS/SI excerpts are reproduced in blue below each reply.")

    for paragraph, (response, excerpts) in zip(placeholders, RESPONSES):
        populate_reply(paragraph, response, excerpts)

    style_document(doc)
    out = PAPER / "Response_Letter.docx"
    doc.save(out)
    shutil.copy2(out, ARCHIVE_DUP / "Response_Letter_Revised_20260824.docx")
    print(out)


if __name__ == "__main__":
    main()
