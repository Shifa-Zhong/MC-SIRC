#!/usr/bin/env python3
"""Validate clean/highlighted Word identity, highlight granularity, and package integrity."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
AUDIT_JSON = ROOT / "output" / "results" / "revision4_consistency_audit.json"
OUT_JSON = ROOT / "output" / "results" / "revision4_deliverable_validation.json"
OUT_MD = PAPER / "Highlight_Validation_20260824.md"


PAIRS = [
    ("MS", PAPER / "Manuscript_Final_MC-SIRC.docx", PAPER / "Manuscript_Highlighted_20260824.docx"),
    ("SI", PAPER / "SI_Final.docx", PAPER / "SI_Highlighted_20260824.docx"),
    ("Abstract", PAPER / "Abstract.docx", PAPER / "Abstract_Highlighted_20260824.docx"),
]


def sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def document_text(doc):
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]
    return paragraphs, tables


def zip_stats(path):
    with ZipFile(path) as archive:
        bad = archive.testzip()
        xml = archive.read("word/document.xml")
        return {
            "zip_ok": bad is None,
            "yellow": xml.count(b'<w:highlight w:val="yellow"'),
            "all_highlight": xml.count(b"<w:highlight"),
            "tracked_insertions": xml.count(b'<w:ins ') + xml.count(b'<w:ins>'),
            "tracked_deletions": xml.count(b'<w:del ') + xml.count(b'<w:del>'),
            "paragraph_shading": xml.count(b'<w:shd w:fill="yellow"'),
        }


def run_highlight_stats(doc):
    highlighted_runs = []
    paragraph_counts = []
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        marked = []
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                highlighted_runs.append(run.text)
                marked.append(run.text)
        if marked:
            paragraph_counts.append(len(marked))
    return {
        "highlighted_runs": len(highlighted_runs),
        "highlighted_paragraphs_or_cells": len(paragraph_counts),
        "multi_segment_paragraphs_or_cells": sum(value > 1 for value in paragraph_counts),
        "whitespace_only_highlights": sum(not text.strip() for text in highlighted_runs),
        "edge_whitespace_highlights": sum(text != text.strip() for text in highlighted_runs),
    }


def validate_pair(scope, clean_path, highlighted_path):
    clean = Document(clean_path)
    highlighted = Document(highlighted_path)
    clean_text = document_text(clean)
    highlighted_text = document_text(highlighted)
    stats = zip_stats(highlighted_path)
    run_stats = run_highlight_stats(highlighted)
    checks = {
        "text_identity": clean_text == highlighted_text,
        "paragraph_count_identity": len(clean.paragraphs) == len(highlighted.paragraphs),
        "table_count_identity": len(clean.tables) == len(highlighted.tables),
        "image_count_identity": len(clean.inline_shapes) == len(highlighted.inline_shapes),
        "zip_ok": stats["zip_ok"],
        "all_highlights_are_yellow": stats["yellow"] == stats["all_highlight"] and stats["all_highlight"] > 0,
        "no_paragraph_level_yellow_shading": stats["paragraph_shading"] == 0,
        "no_tracked_changes": stats["tracked_insertions"] == 0 and stats["tracked_deletions"] == 0,
        "no_whitespace_only_highlights": run_stats["whitespace_only_highlights"] == 0,
        "no_edge_whitespace_in_highlights": run_stats["edge_whitespace_highlights"] == 0,
        "word_level_segmentation_present": run_stats["multi_segment_paragraphs_or_cells"] > 0,
    }
    return {
        "scope": scope,
        "clean": str(clean_path),
        "highlighted": str(highlighted_path),
        "clean_sha256": sha256(clean_path),
        "highlighted_sha256": sha256(highlighted_path),
        "checks": checks,
        "xml": stats,
        "runs": run_stats,
    }


def main():
    pairs = [validate_pair(*pair) for pair in PAIRS]
    audit = json.loads(AUDIT_JSON.read_text(encoding="utf-8"))
    audit_pass = all(check["status"] == "PASS" for check in audit["checks"])
    response = Document(PAPER / "Response_Letter.docx")
    replies = [p for p in response.paragraphs if p.text.strip().startswith("Reply:")]
    blue_replies = 0
    for paragraph in replies:
        blue = [run for run in paragraph.runs if run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "0070C0" and run.text.strip()]
        if len(blue) >= 2:
            blue_replies += 1
    expected = [
        ROOT / "REPRODUCIBILITY.md",
        ROOT / "config" / "revision3_parameters.json",
        ROOT / "data" / "example" / "outlet_monthly_example.csv",
        ROOT / "data" / "example" / "source_inventory_example.csv",
        ROOT / "scripts" / "analysis" / "revision3_diagnostics.py",
        ROOT / "scripts" / "analysis" / "spatial_identifiability_unbounded.py",
        ROOT / "scripts" / "analysis" / "revision4_s3_sensitivity.py",
    ]
    package_checks = {
        "consistency_audit_all_pass": audit_pass,
        "response_replies": len(replies) == 32,
        "response_blue_excerpts": blue_replies == 32,
        "repository_core_files_present": all(path.exists() for path in expected),
        "no_code_supplement_zip": not (PAPER / "Code_Supplement_20260824.zip").exists(),
        "graphical_abstract_one_image": len(Document(PAPER / "GA.docx").inline_shapes) == 1,
    }
    result = {"pairs": pairs, "package_checks": package_checks}
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Highlight and deliverable validation — 2026-08-24",
        "",
        "| Scope | Result | Text identical | Yellow run highlights | Multi-segment paragraphs/cells |",
        "|---|---|---:|---:|---:|",
    ]
    for pair in pairs:
        passed = all(pair["checks"].values())
        lines.append(
            f"| {pair['scope']} | {'PASS' if passed else 'CHECK'} | {pair['checks']['text_identity']} | "
            f"{pair['xml']['all_highlight']} | {pair['runs']['multi_segment_paragraphs_or_cells']} |"
        )
    lines.extend(["", "## Package checks", ""])
    for name, passed in package_checks.items():
        lines.append(f"- {'PASS' if passed else 'CHECK'} — {name}")
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    failures = []
    for pair in pairs:
        failures.extend(f"{pair['scope']}: {name}" for name, passed in pair["checks"].items() if not passed)
    failures.extend(f"package: {name}" for name, passed in package_checks.items() if not passed)
    print("PASS" if not failures else "CHECK")
    for failure in failures:
        print(failure)
    print(OUT_MD)


if __name__ == "__main__":
    main()
