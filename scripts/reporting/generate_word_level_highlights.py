#!/usr/bin/env python3
"""Generate word-level highlighted Word files against the locked submitted originals."""

from __future__ import annotations

import json
import re
from collections import Counter
from copy import deepcopy
from difflib import SequenceMatcher
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
BACKUP = PAPER / "_archive_non_submission_20260824" / "historical_backups" / "backup_before_revision_20260824"
REPORT = ROOT / "output" / "results" / "word_level_highlight_report.json"
TOKEN_RE = re.compile(r"\w+(?:[‑–—-]\w+)*|[^\w\s]", flags=re.UNICODE)


JOBS = [
    (
        "MS",
        BACKUP / "Manuscript_Final_MC-SIRC.docx",
        PAPER / "Manuscript_Final_MC-SIRC.docx",
        PAPER / "Manuscript_Highlighted_20260824.docx",
    ),
    (
        "SI",
        BACKUP / "SI_Final.docx",
        PAPER / "SI_Final.docx",
        PAPER / "SI_Highlighted_20260824.docx",
    ),
    (
        "Abstract",
        BACKUP / "Abstract.docx",
        PAPER / "Abstract.docx",
        PAPER / "Abstract_Highlighted_20260824.docx",
    ),
]


def normalized_tokens(text):
    return [match.group(0).lower() for match in TOKEN_RE.finditer(text)]


def paragraph_similarity(old_paragraph, new_paragraph):
    old = normalized_tokens(old_paragraph.text)
    new = normalized_tokens(new_paragraph.text)
    if not old or not new:
        return 0.0
    if old == new:
        return 1.0
    sequence = SequenceMatcher(None, old, new, autojunk=False).ratio()
    old_set, new_set = set(old), set(new)
    jaccard = len(old_set & new_set) / max(1, len(old_set | new_set))
    score = 0.72 * sequence + 0.28 * jaccard
    old_heading = old_paragraph.style.name.startswith("Heading") or old_paragraph.style.name == "Title"
    new_heading = new_paragraph.style.name.startswith("Heading") or new_paragraph.style.name == "Title"
    if old_heading != new_heading:
        score *= 0.25
    elif old_heading and new_heading:
        score = min(1.0, score + 0.08)
    return score


def align_paragraphs(old_doc, new_doc):
    old = [(index, paragraph) for index, paragraph in enumerate(old_doc.paragraphs) if paragraph.text.strip()]
    new = [(index, paragraph) for index, paragraph in enumerate(new_doc.paragraphs) if paragraph.text.strip()]
    n_old, n_new = len(old), len(new)
    gap = -0.12
    scores = [[0.0] * (n_new + 1) for _ in range(n_old + 1)]
    trace = [[None] * (n_new + 1) for _ in range(n_old + 1)]
    for i in range(1, n_old + 1):
        scores[i][0] = i * gap
        trace[i][0] = "old-gap"
    for j in range(1, n_new + 1):
        scores[0][j] = j * gap
        trace[0][j] = "new-gap"
    similarity_cache = {}
    for i in range(1, n_old + 1):
        for j in range(1, n_new + 1):
            similarity = paragraph_similarity(old[i - 1][1], new[j - 1][1])
            similarity_cache[(i, j)] = similarity
            match = scores[i - 1][j - 1] + 2.0 * similarity - 0.70
            skip_old = scores[i - 1][j] + gap
            skip_new = scores[i][j - 1] + gap
            best = max(match, skip_old, skip_new)
            scores[i][j] = best
            trace[i][j] = "match" if best == match else ("old-gap" if best == skip_old else "new-gap")
    mapping = {}
    details = []
    i, j = n_old, n_new
    while i > 0 or j > 0:
        action = trace[i][j]
        if action == "match":
            similarity = similarity_cache[(i, j)]
            if similarity >= 0.23:
                mapping[new[j - 1][0]] = old[i - 1][0]
                details.append({"new": new[j - 1][0], "old": old[i - 1][0], "similarity": similarity})
            i -= 1
            j -= 1
        elif action == "old-gap":
            i -= 1
        else:
            j -= 1
    return mapping, list(reversed(details))


def changed_mask(old_text, new_text):
    old_matches = list(TOKEN_RE.finditer(old_text))
    new_matches = list(TOKEN_RE.finditer(new_text))
    old_tokens = [match.group(0) for match in old_matches]
    new_tokens = [match.group(0) for match in new_matches]
    mask = [False] * len(new_text)
    matcher = SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    for tag, _, _, start, end in matcher.get_opcodes():
        if tag == "equal":
            continue
        for match in new_matches[start:end]:
            for position in range(match.start(), match.end()):
                mask[position] = True
    return mask, len(new_tokens), sum(1 for match in new_matches if any(mask[match.start():match.end()]))


def fully_changed_mask(text):
    mask = [False] * len(text)
    matches = list(TOKEN_RE.finditer(text))
    for match in matches:
        for position in range(match.start(), match.end()):
            mask[position] = True
    return mask, len(matches), len(matches)


def replace_runs_with_highlight(paragraph, mask):
    if not any(mask) or paragraph._p.xpath(".//w:drawing"):
        return 0
    runs = list(paragraph.runs)
    snapshots = [(run.text, deepcopy(run._r.rPr) if run._r.rPr is not None else None) for run in runs]
    combined = "".join(text for text, _ in snapshots)
    if combined != paragraph.text or len(combined) != len(mask):
        return 0
    for run in runs:
        run._r.getparent().remove(run._r)
    offset = 0
    highlighted_segments = 0
    for text, properties in snapshots:
        local = mask[offset:offset + len(text)]
        start = 0
        while start < len(text):
            state = local[start]
            end = start + 1
            while end < len(text) and local[end] == state:
                end += 1
            new_run = paragraph.add_run(text[start:end])
            if properties is not None:
                existing = new_run._r.rPr
                if existing is not None:
                    new_run._r.remove(existing)
                new_run._r.insert(0, deepcopy(properties))
            if state:
                existing_highlight = new_run._r.get_or_add_rPr().find(qn("w:highlight"))
                if existing_highlight is not None:
                    new_run._r.get_or_add_rPr().remove(existing_highlight)
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                highlighted_segments += 1
            start = end
        offset += len(text)
    return highlighted_segments


def highlight_paragraph(old_text, paragraph):
    if not paragraph.text.strip():
        return {"tokens": 0, "changed": 0, "segments": 0, "mode": "empty"}
    if old_text is None:
        mask, tokens, changed = fully_changed_mask(paragraph.text)
        mode = "new"
    else:
        mask, tokens, changed = changed_mask(old_text, paragraph.text)
        mode = "unchanged" if changed == 0 else ("full" if changed == tokens else "partial")
    segments = replace_runs_with_highlight(paragraph, mask)
    return {"tokens": tokens, "changed": changed, "segments": segments, "mode": mode}


def table_mapping(scope, old_doc, new_doc):
    if scope == "MS":
        return {0: None, 1: 0, 2: 1}
    return {index: index if index < len(old_doc.tables) else None for index in range(len(new_doc.tables))}


def highlight_document(scope, original_path, clean_path, output_path):
    original = Document(original_path)
    highlighted = Document(clean_path)
    clean_text_snapshot = [paragraph.text for paragraph in highlighted.paragraphs]
    mapping, mapping_details = align_paragraphs(original, highlighted)
    counts = Counter()
    paragraph_details = []
    for new_index, paragraph in enumerate(highlighted.paragraphs):
        old_index = mapping.get(new_index)
        old_text = original.paragraphs[old_index].text if old_index is not None else None
        result = highlight_paragraph(old_text, paragraph)
        counts[result["mode"]] += 1
        counts["tokens"] += result["tokens"]
        counts["changed_tokens"] += result["changed"]
        counts["highlight_segments"] += result["segments"]
        if result["changed"]:
            paragraph_details.append({"new": new_index, "old": old_index, **result})
    mappings = table_mapping(scope, original, highlighted)
    table_rows = []
    for new_table_index, new_table in enumerate(highlighted.tables):
        old_table_index = mappings.get(new_table_index)
        old_table = original.tables[old_table_index] if old_table_index is not None and old_table_index < len(original.tables) else None
        for row_index, row in enumerate(new_table.rows):
            for column_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    old_text = None
                    if old_table is not None and row_index < len(old_table.rows) and column_index < len(old_table.columns):
                        old_cell = old_table.cell(row_index, column_index)
                        if paragraph_index < len(old_cell.paragraphs):
                            old_text = old_cell.paragraphs[paragraph_index].text
                    result = highlight_paragraph(old_text, paragraph)
                    counts[f"table_{result['mode']}"] += 1
                    counts["table_tokens"] += result["tokens"]
                    counts["table_changed_tokens"] += result["changed"]
                    counts["table_highlight_segments"] += result["segments"]
                    if result["changed"]:
                        table_rows.append({
                            "table": new_table_index,
                            "row": row_index,
                            "column": column_index,
                            "paragraph": paragraph_index,
                            **result,
                        })
    if [paragraph.text for paragraph in highlighted.paragraphs] != clean_text_snapshot:
        raise RuntimeError(f"Text changed while highlighting {scope}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    highlighted.save(output_path)
    with ZipFile(output_path) as archive:
        if archive.testzip() is not None:
            raise RuntimeError(f"Corrupt highlighted package: {output_path}")
        highlight_count = archive.read("word/document.xml").count(b"<w:highlight")
    return {
        "scope": scope,
        "original": str(original_path),
        "clean": str(clean_path),
        "output": str(output_path),
        "counts": dict(counts),
        "xml_highlight_elements": highlight_count,
        "paragraph_mapping": mapping_details,
        "changed_paragraphs": paragraph_details,
        "changed_table_cells": table_rows,
    }


def main():
    reports = [highlight_document(*job) for job in JOBS]
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8")
    for report in reports:
        print(report["scope"], report["counts"], "xml_highlights=", report["xml_highlight_elements"])
        print(report["output"])
    print(REPORT)


if __name__ == "__main__":
    main()
