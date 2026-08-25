#!/usr/bin/env python3
"""Independent final audit of the 2026-08-25 submission package."""

from __future__ import annotations

from collections import Counter
from pathlib import Path
import hashlib
import json
import re
import zipfile

import pandas as pd
from PIL import Image
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
OUT_JSON = ROOT / "output" / "results" / "final_submission_audit_20260825.json"
OUT_MD = ROOT / "output" / "results" / "final_submission_audit_20260825.md"
BLUE = "0070C0"
TOKEN_RE = re.compile(r"\b[\w–—-]+\b", re.UNICODE)


def all_text(document):
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def table_signature(table):
    return tuple(tuple(cell.text for cell in row.cells) for row in table.rows)


def paragraph_and_table_text(document):
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    tables = [[cell.text for row in table.rows for cell in row.cells] for table in document.tables]
    return paragraphs, tables


def three_line(table):
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is None:
        return False
    values = {child.tag.rsplit("}", 1)[-1]: child.get(qn("w:val")) for child in borders}
    return values.get("top") == "single" and values.get("bottom") == "single" and values.get("insideH") == "nil" and values.get("insideV") == "nil"


def check_docx(path):
    with zipfile.ZipFile(path) as archive:
        corrupt = archive.testzip()
        document_xml = archive.read("word/document.xml")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    return corrupt is None, b"w:ins" not in document_xml and b"w:del" not in document_xml, len(media)


def add(results, name, passed, detail):
    results.append({"status": "PASS" if passed else "FAIL", "check": name, "detail": detail})


def main():
    ms = Document(PAPER / "Manuscript_Final_MC-SIRC.docx")
    si = Document(PAPER / "SI_Final.docx")
    response = Document(PAPER / "Response_Letter.docx")
    abstract = Document(PAPER / "Abstract.docx")
    results = []

    required = [
        "Manuscript_Final_MC-SIRC.docx", "Manuscript_Highlighted_20260824.docx",
        "SI_Final.docx", "SI_Highlighted_20260824.docx", "Abstract.docx",
        "Abstract_Highlighted_20260824.docx", "Response_Letter.docx", "GA.docx",
        "cover letter.docx", "highlights.docx", "declarationStatement.docx",
        "Nanchuan_River_Basin_Study_Area.kml", "Nanchuan_River_Basin_Study_Area.kmz",
    ]
    add(results, "required submission files", all((PAPER / name).exists() for name in required), str(required))

    for name in [item for item in required if item.endswith(".docx")]:
        valid, clean_tracking, media = check_docx(PAPER / name)
        add(results, f"valid Word package: {name}", valid, f"media={media}")
        add(results, f"no tracked insertions/deletions: {name}", clean_tracking, "document.xml")

    labels = [paragraph.text.split(":", 1)[0] for paragraph in ms.paragraphs if paragraph.text.startswith(("Study Region:", "Study Focus:", "New Hydrological Insights for the Region:"))]
    abstract_text = " ".join(paragraph.text.split(":", 1)[1] for paragraph in ms.paragraphs if paragraph.text.startswith(("Study Region:", "Study Focus:", "New Hydrological Insights for the Region:")))
    words = len(TOKEN_RE.findall(abstract_text))
    add(results, "exact abstract labels", labels == ["Study Region", "Study Focus", "New Hydrological Insights for the Region"], str(labels))
    add(results, "abstract within 225 words", words <= 225, f"words={words}")

    ms_text = all_text(ms)
    si_text = all_text(si)
    bad_patterns = [
        r"53,155 non-point-source grids", r"53,155 1 km × 1 km cells", r"Previously reported k",
        r"excluded from spatial calibration", r"removing the k", r"Monitor/Entry", r"95% CI",
        r"\bContrib\.\b", r"event-weighted upper", r"S4 event upper", r"Code_Supplement",
        r"four objectives", r"revision-2026-08-24", r"†",
    ]
    hits = [pattern for pattern in bad_patterns if re.search(pattern, ms_text + "\n" + si_text, re.I)]
    add(results, "no known stale or ambiguous terminology in MS/SI", not hits, str(hits))
    add(results, "correct inventory record/grid distinction", "1,588 unique grid IDs" in ms_text and "53,155 polygon-intersection records" in ms_text, "MS and SI")
    add(results, "industrial TN is unavailable, not zero", si.tables[7].cell(1, 3).text == "—", si.tables[7].cell(1, 3).text)
    add(results, "Table S30 exact-bound count corrected", si.tables[32].cell(3, 5).text.startswith("1 of 4"), si.tables[32].cell(3, 5).text)
    add(results, "Table S25 has no contribution column", "Contrib." not in [cell.text for cell in si.tables[27].rows[0].cells], str([cell.text for cell in si.tables[27].rows[0].cells]))
    add(results, "MCMC tables use CrI/Rhat/ESS", all("95% CrI low" in [cell.text for cell in si.tables[index].rows[0].cells] and "Split R-hat" in [cell.text for cell in si.tables[index].rows[0].cells] and "ESS" in [cell.text for cell in si.tables[index].rows[0].cells] for index in range(12, 16)), "Tables S13–S16")
    add(results, "all MS/SI tables use three-line borders", all(three_line(table) for table in list(ms.tables) + list(si.tables)), f"MS={len(ms.tables)}, SI={len(si.tables)}")

    empty_si = [index for index, paragraph in enumerate(si.paragraphs) if not paragraph.text.strip() and not paragraph._p.xpath(".//w:drawing | .//w:pict | .//w:br | .//w:sectPr")]
    add(results, "no empty SI spacer paragraphs", not empty_si, str(empty_si))
    text_s8 = next((paragraph for paragraph in si.paragraphs if paragraph.text.startswith("Text S8.")), None)
    add(results, "Text S8 uses heading style", text_s8 is not None and text_s8.style.name.startswith("Heading"), text_s8.style.name if text_s8 else "missing")

    references_heading = next(i for i, paragraph in enumerate(ms.paragraphs) if paragraph.text.strip() == "References")
    references = [paragraph.text for paragraph in ms.paragraphs[references_heading + 1:] if paragraph.text.strip()]
    add(results, "29 references retained", len(references) == 29, f"count={len(references)}")
    add(results, "references alphabetical", references == sorted(references, key=str.lower), "first authors=" + ", ".join(item.split(",", 1)[0] for item in references[:5]))
    reference_checks = [
        "Ongley, E.D., Zhang, X., Yu, T.", "Ecological Modelling 159, 269–277", "The MARINA model (Model to Assess River Inputs of Nutrients to seAs)",
        "10.3969/j.issn.1000-5641.2026.01.012",
    ]
    add(results, "corrected reference metadata present", all(item in ms_text for item in reference_checks), str(reference_checks))

    replies = [paragraph for paragraph in response.paragraphs if paragraph.text.strip().startswith("Reply:")]
    add(results, "all 32 reviewer replies", len(replies) == 32, f"count={len(replies)}")
    add(results, "no reply placeholders", "reply: xxx" not in all_text(response).lower(), "Response_Letter.docx")
    blue_runs = [run for paragraph in response.paragraphs for run in paragraph.runs if run.font.color.rgb is not None and str(run.font.color.rgb) == BLUE]
    add(results, "response contains blue revision excerpts", len(blue_runs) >= 60, f"blue_runs={len(blue_runs)}")
    labels_in_response = [paragraph for paragraph in response.paragraphs if paragraph.text.startswith("Revised ") and paragraph.text.rstrip().endswith(":")]
    blue_labels = [paragraph for paragraph in labels_in_response if all(run.font.color.rgb is not None and str(run.font.color.rgb) == BLUE for run in paragraph.runs if run.text)]
    add(results, "all response figure/table labels are blue", len(labels_in_response) == len(blue_labels), f"labels={len(labels_in_response)}, blue={len(blue_labels)}")
    source_tables = Counter(table_signature(table) for table in list(ms.tables) + list(si.tables))
    unmatched_tables = [index for index, table in enumerate(response.tables) if table_signature(table) not in source_tables]
    add(results, "all response tables are exact MS/SI tables", not unmatched_tables, str(unmatched_tables))

    for clean_name, highlighted_name in [
        ("Manuscript_Final_MC-SIRC.docx", "Manuscript_Highlighted_20260824.docx"),
        ("SI_Final.docx", "SI_Highlighted_20260824.docx"),
        ("Abstract.docx", "Abstract_Highlighted_20260824.docx"),
    ]:
        clean = Document(PAPER / clean_name)
        highlighted = Document(PAPER / highlighted_name)
        add(results, f"highlighted text equals clean text: {clean_name}", paragraph_and_table_text(clean) == paragraph_and_table_text(highlighted), highlighted_name)
    report = json.loads((ROOT / "output" / "results" / "word_level_highlight_report.json").read_text(encoding="utf-8"))
    ms_report = next(item for item in report if item["scope"] == "MS")
    rematches = [item for item in ms_report["paragraph_mapping"] if item.get("reference_rematch")]
    add(results, "references matched independent of order for word-level diff", len(rematches) >= 20, f"reference_rematches={len(rematches)}")

    metadata = pd.read_excel(ROOT / "output" / "results" / "revision5_mcmc_diagnostics.xlsx", sheet_name="Metadata")
    add(results, "MCMC split-Rhat ≤1.05", float(metadata.Max_split_Rhat.max()) <= 1.05, f"max={metadata.Max_split_Rhat.max():.4f}")
    add(results, "MCMC ESS ≥1000", float(metadata.Min_ESS.min()) >= 1000, f"min={metadata.Min_ESS.min():.1f}")
    add(results, "MCMC acceptance fractions 0.2–0.5", bool(((metadata.Mean_acceptance_fraction >= 0.2) & (metadata.Mean_acceptance_fraction <= 0.5)).all()), str(metadata.Mean_acceptance_fraction.round(3).tolist()))

    with zipfile.ZipFile(PAPER / "Nanchuan_River_Basin_Study_Area.kmz") as archive:
        kml_names = archive.namelist()
        kml_valid = archive.testzip() is None and "doc.kml" in kml_names
    add(results, "valid study-area KMZ", kml_valid, str(kml_names))
    kml_text = (PAPER / "Nanchuan_River_Basin_Study_Area.kml").read_text(encoding="utf-8")
    add(results, "KML area description matches basin", "1438.64 km²" in kml_text, "1438.64 km²")
    with Image.open(ROOT / "output" / "figures" / "revision4" / "graphical_abstract_identifiability.png") as image:
        size = image.size
    add(results, "graphical abstract resolution", size[0] >= 1328 and size[1] >= 531, str(size))

    repo_files = [
        "scripts/reporting/generate_revision3_figure1.py", "scripts/reporting/generate_revision3_figures.py",
        "scripts/reporting/finalize_revision5_submission.py", "scripts/reporting/generate_word_level_highlights_v2.py",
        "scripts/analysis/revision5_mcmc_diagnostics.py", "REPRODUCIBILITY.md",
    ]
    add(results, "repository reproduction files exist", all((ROOT / item).exists() for item in repo_files), str(repo_files))
    repo_text = (ROOT / "REPRODUCIBILITY.md").read_text(encoding="utf-8")
    add(results, "repository instructions cite existing scripts", all(item in repo_text for item in ["revision3_diagnostics.py", "spatial_identifiability_unbounded.py", "revision5_mcmc_diagnostics.py", "generate_revision3_figures.py"]), "REPRODUCIBILITY.md")

    counts = Counter(item["status"] for item in results)
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps({"summary": counts, "checks": results}, ensure_ascii=False, indent=2, default=dict), encoding="utf-8")
    lines = ["# Final submission audit — 2026-08-25", "", f"Summary: {dict(counts)}", "", "| Status | Check | Detail |", "|---|---|---|"]
    for item in results:
        detail = str(item["detail"]).replace("|", "\\|").replace("\n", " ")
        lines.append(f"| {item['status']} | {item['check']} | {detail} |")
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(dict(counts))
    for item in results:
        if item["status"] == "FAIL":
            print("FAIL", item["check"], item["detail"])
    print(OUT_MD)
    raise SystemExit(0 if counts["FAIL"] == 0 else 1)


if __name__ == "__main__":
    main()
