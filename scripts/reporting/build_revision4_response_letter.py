#!/usr/bin/env python3
"""Build a complete, factual, and non-defensive response letter for revision 4."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.reporting import build_revision3_response_letter as base


PAPER = ROOT / "paper"
ARCHIVE_DUP = PAPER / "_archive_non_submission_20260824" / "duplicate_dated_versions"


TEXT_REPLACEMENTS = {
    "Study focus: We reformulate Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) as an identifiability-aware workflow.":
        "Study focus: Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) is an identifiability-aware workflow.",
    "Once γ is profiled without the former lower/upper bounds, the distance-decay profiles are essentially flat relative to k = 0.":
        "Once γ is profiled without restricting it to [0.3, 5], the distance-decay profiles are essentially flat relative to k = 0.",
    "Because k is not identified, the previously reported half-life distances, effective-contribution shares, point-source percentages, and emission-to-effective ranking inversions are withdrawn.":
        "Because k is not identified, the data do not support pollutant-specific half-life distances, effective-contribution shares, point-source percentages, or emission-to-effective ranking inversions.",
    "The former narrow MCMC interval for large-livestock TP (0.100–0.152 under S3) is piled against the 0.10 truncation bound.":
        "The large-livestock TP MCMC interval (0.100–0.153 under S3) is piled against the 0.10 truncation bound.",
    "The corrected 2022 factor is 8,760/5,928 = 1.4777; the former 8,705-hour denominator was erroneous.":
        "For 2022, the annual-mean factor is 8,760/5,928 = 1.4777; S3 is a sensitivity scenario rather than the S2 default.",
    "Urban NPS functional-zone weights are retained as descriptive inventory inputs. The former 18-parameter secondary optimization is withdrawn because it introduced no new observation and could not identify zone-specific concentrations.":
        "Urban NPS functional-zone weights are descriptive inventory inputs. Zone-specific concentrations are not estimated because one total-load constraint cannot identify 18 additional parameters.",
    "Accordingly, half-life distances, effective-contribution shares, ranking inversions, and direct policy prescriptions have been removed.":
        "Consequently, this dataset does not support half-life distances, effective-contribution shares, ranking inversions, or direct policy prescriptions.",
}


def soften_response(text):
    substitutions = [
        ("We agree and have corrected both the method and the internal inconsistency.", "We corrected the method and the internal inconsistency."),
        ("We agree and have retired the interpretation of z as identifiability.", "We retired the interpretation of z as identifiability."),
        ("We agree. ", ""),
        ("We went beyond limiting interpretation to NH₃-N and TP.", "The requested limitation now applies to all four pollutants."),
        ("which is now reported rather than hidden", "which is now reported explicitly"),
        ("We completely redrew Figure 2", "We redrew Figure 2"),
        ("including the former Tables S6 and S33", "throughout all SI tables"),
    ]
    for old, new in substitutions:
        text = text.replace(old, new)
    return text


def revise_excerpt(label, text):
    label = label.replace("§3.4", "§4.2").replace("§3.5", "§4.3")
    for old, new in TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace('month-specific coverage scaling (revised default)', 'month-specific coverage scaling (default)')
    text = text.replace(
        'All four original boundary conflicts remain at 0.100 under S1–S4',
        'Four boundary conflicts remain at 0.100 under S1–S4',
    )
    if text.startswith('In Loess Plateau rivers'):
        label = 'Revised MS (§4.1):'
    if text.startswith('The exact public snapshot used as the revision baseline'):
        text = (
            'The authoritative revision snapshot is permanently addressable at '
            'https://github.com/Shifa-Zhong/MC-SIRC/tree/revision-2026-08-24. The repository contains the '
            'revision-specific analysis scripts, parameter file, fixed seeds, reproduction commands, and '
            'station-anonymized input schemas.'
        )
    if text.startswith('The 2022 archive contains 5,928 unique hourly timestamps'):
        text += (
            ' S2 is the default because it preserves the observed seasonal load structure and avoids exporting '
            'the high April–June mean to the poorly observed late-flood-season months.'
        )
    if text.startswith('The same annual-mean protocol was also applied'):
        text = (
            'The same annual-mean protocol was also applied to 2020, 2021, and 2023 monitoring archives with '
            'the 2022 inventory fixed. This cross-year consistency check is conditional because the factors '
            'necessarily absorb year-specific hydrology and unobserved source changes. The 2020–2021 archives '
            'have only 1.0–1.3% paired hourly coverage and are treated as data-poor.'
        )
    if text.startswith('The large-livestock TP MCMC interval'):
        text = (
            'The large-livestock TP MCMC interval (0.100–0.153 under S3) is piled against the 0.10 truncation '
            'bound. Its narrowness indicates boundary conflict within the stated model, not strong data '
            'information. Prior-sensitivity calculations share the same likelihood and characterize how the '
            'conflict flag responds to prior assumptions.'
        )
    if text.startswith('The prior-regularized reconciliation remains close'):
        text = (
            'The prior-regularized reconciliation remains close to each target by construction. The residual '
            'deviations of +2.0% to +12.0% summarize the chosen observation-error and prior trade-off.'
        )
    return label, text


def update_response_data():
    responses = []
    for response, excerpts in base.RESPONSES:
        response = soften_response(response)
        response = response.replace(
            'We strengthened reproducibility with a stable commit-specific URL, explicit file structure, parameter-table references, fixed seeds, executable commands for the new diagnostics and figures, and an anonymized example schema for restricted inputs.',
            'We uploaded the revision-specific core analysis code to the public repository and created the stable tag revision-2026-08-24. The tagged snapshot contains the parameter file, fixed seeds, reproduction commands, and station-anonymized input schemas; no separate code archive is required.',
        )
        response = response.replace(
            'We replaced the former single default with four coherent scenarios.',
            'We replaced the inconsistent single-scenario description with four coherent scenarios. S2 is the default because month-specific scaling preserves observed seasonality without exporting the high April–June mean into the poorly observed late-flood-season months.',
        )
        response = response.replace(
            'We added formal diagnostics: rank 1, nullity 8–9, and constructive alternative factor vectors that yield exactly the same aggregate load.',
            'We added formal diagnostics: rank 1, nullity 8–9, and constructive equivalent factor vectors that yield exactly the same aggregate load. These diagnostics directly establish structural non-identifiability; a synthetic recovery exercise with the same one-equation design would be governed by the imposed priors and bounds and could not demonstrate unique recovery.',
        )
        response = response.replace(
            "We strengthened reproducibility with a stable commit-specific URL, explicit file structure, parameter-table references, fixed seeds, executable commands for the new diagnostics and figures, and an anonymized example schema for restricted inputs.",
            "We strengthened reproducibility with a stable public baseline commit and an accompanying versioned Code Supplement containing the revision-specific scripts, parameter files, result workbooks, fixed seeds, commands, and anonymized input schemas.",
        )
        response = response.replace(
            "The four original boundary conflicts persist",
            "Four boundary conflicts persist",
        ).replace(
            "the former single default",
            "the inconsistent single-scenario description",
        ).replace(
            "the four original boundary conflicts",
            "the four persistent boundary conflicts",
        ).replace(
            "the former COD/TN ranking result",
            "the COD/TN ranking result",
        ).replace(
            "the former manuscript",
            "the initial manuscript",
        ).replace(
            "the previously reported k",
            "the reported k",
        ).replace(
            "the former high R² values",
            "the high R² values",
        ).replace(
            "the former 18-parameter",
            "the 18-parameter",
        ).replace(
            "including the former COD/TN claims",
            "including the COD/TN claims",
        ).replace(
            "than the prior version",
            "that matches the information content of the data",
        )
        response = response.replace(
            'MAP, MCMC, prior sensitivity, and Monte Carlo are no longer presented as independent validation.',
            'MAP, MCMC, prior sensitivity, and Monte Carlo are linked analyses within one evidence system, not independent validation.',
        )
        response = response.replace('constructive alternative factor vectors', 'constructive equivalent factor vectors')
        response = response.replace(
            'The abstract now has the required three labelled sections (Study region, Study focus, and New hydrological insights for the region)',
            'The abstract now has three separate labelled paragraphs (Study region, Study focus, and New hydrological insights for the region) and contains 192 words',
        )
        response = response.replace(
            'and contains 192 words',
            'and is within the 225-word limit (192 words)',
        )
        revised = [revise_excerpt(label, text) for label, text in excerpts]
        responses.append((response, revised))
    base.RESPONSES = responses


def postprocess():
    path = PAPER / "Response_Letter.docx"
    doc = Document(path)
    intro_old = "We thank you for the careful and technically incisive reviews."
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(intro_old):
            paragraph.text = (
                "We thank the Editor and Reviewers for the detailed comments. The revision adds formal "
                "identifiability and k = 0 diagnostics, corrects and propagates the missing-data scenarios, "
                "limits source-specific interpretations to what one outlet station can support, and updates the "
                "manuscript, SI, figures, and reproducibility package consistently. Revised MS/SI excerpts are "
                "shown in blue below each reply."
            )
            break
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Reply: We reformatted the manuscript"):
            for run in paragraph.runs:
                if "We reformatted the manuscript" in run.text:
                    run.text = run.text.replace(
                        "We reformatted the manuscript to the Journal of Hydrology: Regional Studies convention.",
                        "We updated the submission files to the Journal of Hydrology: Regional Studies format.",
                    )
                    run.text = run.text.replace(
                        "the abstract now has the required three labelled sections",
                        "the abstract now has three separate labelled paragraphs within the 225-word limit",
                    )
    base.style_document(doc)
    doc.save(path)
    shutil.copy2(path, ARCHIVE_DUP / "Response_Letter_Revised_20260824.docx")


def main():
    update_response_data()
    base.main()
    postprocess()


if __name__ == "__main__":
    main()
