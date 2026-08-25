#!/usr/bin/env python3
"""Paragraph-, table-, and reviewer-item-level audit for the revision-4 package."""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
OUT_MD = PAPER / "Submission_Consistency_Audit_20260824.md"
OUT_JSON = ROOT / "output" / "results" / "revision4_consistency_audit.json"


FILES = {
    "MS": PAPER / "Manuscript_Final_MC-SIRC.docx",
    "SI": PAPER / "SI_Final.docx",
    "Abstract": PAPER / "Abstract.docx",
    "Response": PAPER / "Response_Letter.docx",
    "Highlights": PAPER / "highlights.docx",
    "Cover": PAPER / "cover letter.docx",
    "GA": PAPER / "GA.docx",
}


REVIEW_ITEMS = [
    ("R1-General", [["identifiability-aware"], ["rank", "k = 0"]]),
    ("R1-Strengths", [["prior-conflict"], ["not independent validation", "same likelihood"]]),
    ("R1-1", [["rank", "nullity"], ["source-coefficient validation", "prior-regularized allocations"]]),
    ("R1-2", [["8,760", "5,928", "2,832"], ["S1", "S2", "S3", "S4"], ["rankings", "tier"], ["preserves observed seasonality", "preserves the observed seasonal"]]),
    ("R1-3", [["consistency check"], ["2022 inventory fixed"], ["1.0–1.3%", "data-poor"]]),
    ("R1-4", [["k is not independently identifiable", "k is not identified"], ["all four pollutants", "any pollutant"], ["half-lives", "half-life"]]),
    ("R1-5", [["field verification", "field hypothesis"], ["no direct policy", "not be re-ranked"]]),
    ("R1-6", [["prototype"], ["cross-basin", "transferability"]]),
    ("R1-7", [["algorithmic table"], ["input", "output", "assumption"]]),
    ("R1-8", [["revision-2026-08-24"], ["public repository", "tagged snapshot"], ["fixed seeds"], ["anonymized"]]),
    ("R1-Overall", [["identifiability-aware"], ["monitoring roadmap"]]),
    ("R2-General", [["structural issue"], ["effective-contribution"], ["additional observations"]]),
    ("R2-1", [["γwmA(k)"], ["profile γ"], ["k-dependent outputs", "k is not identified"]]),
    ("R2-2", [["k = 0"], ["0.00003"], ["0.00333"]]),
    ("R2-3", [["conflates"], ["field audit"], ["causal"]]),
    ("R2-4", [["same likelihood", "one likelihood"], ["0.10 bound"], ["boundary conflict"]]),
    ("R2-5", [["prior-conflict tiers"], ["rank 1"], ["equivalent"], ["synthetic recovery"], ["structural non-identifiability"]]),
    ("R2-6", [["constructional", "by construction"], ["observation-error and prior trade-off", "reconciliation residual"]]),
    ("R2-7", [["8,705", "8,760"], ["month-specific"], ["S1", "S4"]]),
    ("R2-8", [["consistency check"], ["fixed-inventory", "fixed 2022 inventory"]]),
    ("R2-9", [["pre-transport"], ["cannot attribute", "not used as evidence"]]),
    ("R2-10", [["18-parameter"], ["descriptive inventory inputs"]]),
    ("R2-11", [["effective-contribution percentage"], ["near/far"], ["monitoring-design"]]),
    ("R3-General", [["literature"], ["figures"], ["SI formatting"]]),
    ("R3-1", [["Introduction"], ["recent representative work"], ["research gap"]]),
    ("R3-2", [["redrew Figure 2"], ["Figure 3 was enlarged"]]),
    ("R3-3", [["particulate"], ["nitrification"], ["denitrification"]]),
    ("R3-4", [["two paragraphs"], ["Conclusions"]]),
    ("R3-SI-1", [["Table S1"], ["Table S17"], ["repeated header"]]),
    ("R3-SI-2", [["dagger"], ["no intervening", "no space"]]),
    ("R3-SI-3", [["maximum a posteriori"], ["Markov chain Monte Carlo"]]),
    ("AE", [["three"], ["225-word"], ["author–year"], ["line numbering"]]),
]


def normalize(text):
    return " ".join(text.split())


def all_text(doc):
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def word_count(text):
    return len(re.findall(r"\b[\w–-]+\b", text, flags=re.UNICODE))


def package_ok(path):
    try:
        with ZipFile(path) as archive:
            return archive.testzip() is None and "word/document.xml" in archive.namelist()
    except Exception:
        return False


def paragraph_flags(scope, text):
    lowered = text.lower()
    flags = []
    if re.search(r"\s{2,}", text):
        flags.append("repeated whitespace")
    if scope in {"MS", "SI", "Abstract", "Highlights", "Cover"}:
        for phrase in ["review context", "revised manuscript", "previously reported", "former manuscript"]:
            if phrase in lowered:
                flags.append(f"submission-context phrase: {phrase}")
        for pattern in [r"8,?705", r"68\.1%", r"26\.8%", r"10\.3%", r"four-fold", r"five-fold", r"\b75%\b"]:
            if re.search(pattern, lowered):
                flags.append(f"obsolete claim/number: {pattern}")
    if scope in {'MS', 'SI', 'Abstract', 'Highlights'}:
        for phrase in ['former ', 'previously ', 'withdrawn', 'revised ', 'reviewer ', 'review context']:
            if phrase in lowered:
                flags.append(f'submission-context phrase: {phrase.strip()}')
        for phrase in ['corrected s3', 'corrected annual-mean', 'earlier interval', 'retained only', 'revision-specific', 'revision files', 'are not reported']:
            if phrase in lowered:
                flags.append(f'editing-process phrase: {phrase}')
    return flags


def audit_paragraphs(docs):
    rows = []
    for scope in ["MS", "SI", "Abstract", "Highlights", "Cover"]:
        doc = docs[scope]
        for index, paragraph in enumerate(doc.paragraphs):
            text = normalize(paragraph.text)
            if not text:
                continue
            flags = paragraph_flags(scope, text)
            rows.append({
                "scope": scope,
                "id": f"P{index:03d}",
                "style": paragraph.style.name,
                "status": "CHECK" if flags else "PASS",
                "flags": flags,
                "text": text,
            })
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                text = " | ".join(normalize(cell.text) for cell in row.cells)
                flags = paragraph_flags(scope, text)
                rows.append({
                    "scope": scope,
                    "id": f"T{table_index:02d}R{row_index:03d}",
                    "style": "table row",
                    "status": "CHECK" if flags else "PASS",
                    "flags": flags,
                    "text": text,
                })
    return rows


def extract_abstract(doc):
    paragraphs = []
    active = False
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text == "Abstract":
            active = True
            continue
        if active and text.startswith("Keywords:"):
            break
        if active and text:
            paragraphs.append(text)
    return paragraphs


def reference_audit(doc):
    body = []
    references = []
    active = False
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text == "References":
            active = True
            continue
        (references if active else body).append(text)
    body_text = "\n".join(body)
    body_pairs = set()
    for match in re.finditer(r"([A-Z][A-Za-zÀ-ÖØ-öø-ÿ'’-]+)(?:\s+et al\.|\s+and\s+[A-Z][A-Za-zÀ-ÖØ-öø-ÿ'’-]+)?(?:,\s*|\s+\()(\d{4})\)?", body_text):
        body_pairs.add((match.group(1), match.group(2)))
    ref_pairs = []
    for reference in references:
        if not reference:
            continue
        surname = reference.split(",", 1)[0]
        year_match = re.search(r"\b(19|20)\d{2}\b", reference)
        if year_match:
            ref_pairs.append((surname, year_match.group(0), reference))
    missing = sorted(pair for pair in body_pairs if not any(pair[0] == r[0] and pair[1] == r[1] for r in ref_pairs))
    unused = [reference for surname, year, reference in ref_pairs if (surname, year) not in body_pairs]
    return {"body_pairs": sorted(body_pairs), "missing": missing, "unused": unused, "reference_count": len(ref_pairs)}


def supplement_cross_reference_audit(ms_doc, si_doc):
    """Check SI item numbering, MS coverage, and first-citation order by item type."""
    report = {}
    for kind in ["Text", "Figure", "Table"]:
        item_pattern = re.compile(rf"^{kind}\s+S(\d+)\.", flags=re.I)
        items = [
            int(match.group(1))
            for paragraph in si_doc.paragraphs
            for match in [item_pattern.match(paragraph.text.strip())]
            if match
        ]
        citation_pattern = re.compile(
            rf"\b{kind}s?\s+S(\d+)(?:\s*[-–—]\s*S?(\d+))?",
            flags=re.I,
        )
        cited = []
        events = []
        for paragraph_index, paragraph in enumerate(ms_doc.paragraphs):
            text = normalize(paragraph.text)
            for match in citation_pattern.finditer(text):
                start = int(match.group(1))
                end = int(match.group(2)) if match.group(2) else start
                expanded = list(range(start, end + 1)) if start <= end else [start, end]
                cited.extend(expanded)
                events.append({"paragraph": paragraph_index, "raw": match.group(0), "ids": expanded})
        first_order = []
        for item in cited:
            if item not in first_order:
                first_order.append(item)
        report[kind] = {
            "si_items": items,
            "si_sequential": bool(items) and items == list(range(1, max(items) + 1)),
            "ms_events": events,
            "ms_first_order": first_order,
            "ms_first_order_ascending": first_order == sorted(first_order),
            "missing_from_ms": [item for item in items if item not in cited],
            "nonexistent_in_ms": [item for item in first_order if item not in items],
        }
    return report


def response_audit(doc):
    replies = [p for p in doc.paragraphs if p.text.strip().startswith("Reply:")]
    matrix = []
    for index, (item_id, groups) in enumerate(REVIEW_ITEMS):
        text = normalize(replies[index].text) if index < len(replies) else ""
        group_results = [any(term.lower() in text.lower() for term in alternatives) for alternatives in groups]
        blue_runs = 0
        if index < len(replies):
            for run in replies[index].runs:
                rgb = run.font.color.rgb
                if rgb is not None and str(rgb).upper() == "0070C0" and run.text.strip():
                    blue_runs += 1
        matrix.append({
            "id": item_id,
            "status": "PASS" if all(group_results) and blue_runs >= 2 else "CHECK",
            "coverage_groups": group_results,
            "blue_runs": blue_runs,
            "text": text,
        })
    full_text = "\n".join(normalize(p.text) for p in replies)
    tone = {
        "we_agree_count": len(re.findall(r"\bWe agree\b", full_text, flags=re.I)),
        "defensive_phrase_count": sum(full_text.lower().count(x) for x in ["rather than hidden", "substantive rather than rhetorical", "obvious question"]),
        "xxx_count": full_text.lower().count("xxx"),
    }
    return {"reply_count": len(replies), "matrix": matrix, "tone": tone}


def table_properties(doc):
    issues = []
    for index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            issues.append(f"table {index}: header does not repeat")
        for row_index, row in enumerate(table.rows):
            if row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is None:
                issues.append(f"table {index} row {row_index}: row may split")
    return issues


def three_line_table_issues(doc):
    issues = []
    for index, table in enumerate(doc.tables):
        borders = table._tbl.tblPr.find(qn('w:tblBorders'))
        if borders is None:
            issues.append(f'table {index}: missing table borders')
            continue
        values = {}
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = borders.find(qn(f'w:{edge}'))
            values[edge] = None if element is None else element.get(qn('w:val'))
        if values.get('top') != 'single' or values.get('bottom') != 'single':
            issues.append(f'table {index}: top/bottom rules are not single')
        if any(values.get(edge) != 'nil' for edge in ['left', 'right', 'insideH', 'insideV']):
            issues.append(f'table {index}: vertical/internal rules remain')
        for cell in table.rows[0].cells:
            tc_borders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
            bottom = None if tc_borders is None else tc_borders.find(qn('w:bottom'))
            if bottom is None or bottom.get(qn('w:val')) != 'single':
                issues.append(f'table {index}: header bottom rule missing')
                break
        if any(cell._tc.get_or_add_tcPr().find(qn('w:shd')) is not None for row in table.rows for cell in row.cells):
            issues.append(f'table {index}: cell shading remains')
    return issues


def main():
    docs = {name: Document(path) for name, path in FILES.items()}
    checks = []

    def add(name, passed, detail):
        checks.append({"name": name, "status": "PASS" if passed else "CHECK", "detail": detail})

    for name, path in FILES.items():
        add(f"Word package: {name}", package_ok(path), str(path))
    abstract = extract_abstract(docs["MS"])
    abstract_words = word_count(" ".join(abstract))
    add("MS abstract has three paragraphs", len(abstract) == 3, f"paragraphs={len(abstract)}")
    add("MS abstract ≤225 words", abstract_words <= 225, f"words={abstract_words}")
    add(
        "MS abstract labels",
        [p.split(":", 1)[0] for p in abstract] == ["Study region", "Study focus", "New hydrological insights for the region"],
        str([p.split(":", 1)[0] for p in abstract]),
    )
    ms_text = all_text(docs["MS"])
    si_text = all_text(docs["SI"])
    add("MS key missingness counts", all(x in ms_text for x in ["5,928", "8,760", "2,832", "67.7%"]), "5,928/8,760/2,832/67.7%")
    add("MS key identifiability results", all(x in ms_text for x in ["rank 1", "nullity 8–9", "0.00333"]), "rank/nullity/ΔR²")
    add("SI corrected S3 factor", "1.4777" in si_text and "8,705" not in si_text, "8,760/5,928")
    add("SI corrected TP MCMC upper interval", "0.153" in si_text, "expected 0.100–0.153")
    add("No unsupported old percentages in MS/SI", not re.search(r"26\.8%|10\.3%|four-fold|five-fold|\b75%\b", ms_text + "\n" + si_text, flags=re.I), "old effective-ranking/urban-optimization claims")
    add("Supported inventory composition retained", all(x in ms_text for x in ["56.6%", "65.1%", "37.1%", "inventory-composition shares"]), "descriptive emission shares")
    add("Forward elasticity results retained with scope", all(x in ms_text for x in ["0.391", "0.364", "0.528", "0.704", "pre-transport inventory total"]), "inventory sensitivity, not outlet contribution")
    add("MS figure captions 1–4", all(ms_text.count(f"Figure {i}.") == 1 for i in range(1, 5)), str([ms_text.count(f"Figure {i}.") for i in range(1, 5)]))
    add("MS table captions 1–3", all(ms_text.count(f"Table {i}.") == 1 for i in range(1, 4)), str([ms_text.count(f"Table {i}.") for i in range(1, 4)]))
    si_captions = [int(m.group(1)) for p in docs["SI"].paragraphs for m in [re.match(r"Table S(\d+)\.", p.text.strip())] if m]
    add("SI captions S1–S35 sequential", si_captions == list(range(1, 36)), str(si_captions))
    supplement_refs = supplement_cross_reference_audit(docs["MS"], docs["SI"])
    for kind, expected_count in [("Text", 8), ("Figure", 2), ("Table", 35)]:
        row = supplement_refs[kind]
        add(
            f"SI {kind} numbering is sequential",
            row["si_sequential"] and len(row["si_items"]) == expected_count,
            str(row["si_items"]),
        )
        add(
            f"MS cites every SI {kind}",
            not row["missing_from_ms"],
            f"missing={row['missing_from_ms']}",
        )
        add(
            f"MS first SI {kind} citations are ascending",
            row["ms_first_order_ascending"],
            str(row["ms_first_order"]),
        )
        add(
            f"MS has no nonexistent SI {kind} citations",
            not row["nonexistent_in_ms"],
            f"extra={row['nonexistent_in_ms']}",
        )
    add("SI long-table continuation settings", not table_properties(docs["SI"]), "; ".join(table_properties(docs["SI"])) or "all tables")
    for scope in ["MS", "SI", "Response"]:
        issues = three_line_table_issues(docs[scope])
        add(f"{scope} tables use three-line format", not issues, "; ".join(issues) or f"tables={len(docs[scope].tables)}")
    dagger_bad = bool(re.search(r"\s+†|†\s+", "\n".join(cell.text for t in docs["SI"].tables for row in t.rows for cell in row.cells)))
    add("SI dagger attachment", not dagger_bad, "no spaces around †")
    refs = reference_audit(docs["MS"])
    add("All in-text citations have references", not refs["missing"], str(refs["missing"]))
    add("All references are cited", not refs["unused"], f"unused={len(refs['unused'])}")
    add("Reference list restores cited foundational breadth", refs["reference_count"] == 29, f"references={refs['reference_count']}")
    foundational_citations = [
        "Johnes, 1996", "Chen et al., 2014", "Shen et al., 2012",
        "Behrendt and Opitz, 2000", "Preston et al., 2011",
        "Streeter and Phelps, 1925", "Chapra, 1997",
        "Qian et al., 2003", "Borsuk et al., 2004", "Reckhow, 2003",
        "Wang et al. (2026)", "Rode et al., 2010",
    ]
    add(
        "Restored foundational citations appear in substantive text",
        all(citation in ms_text for citation in foundational_citations),
        str([citation for citation in foundational_citations if citation not in ms_text]),
    )
    supported_after_r2 = [
        "aggregate inventory–outlet discrepancy",
        "inventory-composition shares",
        "pre-transport inventory total",
        "field-verification priorities",
        "conditional comparison across archive years",
    ]
    add(
        "Reviewer 2 correction retains supported non-causal results",
        all(phrase in ms_text for phrase in supported_after_r2),
        str([phrase for phrase in supported_after_r2 if phrase not in ms_text]),
    )
    response = response_audit(docs["Response"])
    add("All 32 reviewer items have replies", response["reply_count"] == 32, f"replies={response['reply_count']}")
    add("All reviewer replies pass coverage and blue-excerpt checks", all(r["status"] == "PASS" for r in response["matrix"]), f"checks={Counter(r['status'] for r in response['matrix'])}")
    response_labels = [
        p for p in docs["Response"].paragraphs
        if p.text.startswith(("Revised MS Figure", "Revised SI Figure", "Revised MS Table", "Revised SI Table"))
    ]
    blue_response_labels = sum(
        any(run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "0070C0" for run in p.runs)
        for p in response_labels
    )
    add("Response embeds five revised figures", len(docs["Response"].inline_shapes) == 5, f"figures={len(docs['Response'].inline_shapes)}")
    add("Response embeds sixteen revised table blocks", len(docs["Response"].tables) == 16, f"tables={len(docs['Response'].tables)}")
    add("Response figure/table labels are blue", len(response_labels) == 21 and blue_response_labels == 21, f"labels={len(response_labels)}, blue={blue_response_labels}")
    add("No response placeholders", response["tone"]["xxx_count"] == 0, str(response["tone"]))
    add("Response tone is not over-defensive", response["tone"]["we_agree_count"] <= 2 and response["tone"]["defensive_phrase_count"] == 0, str(response["tone"]))
    highlights = [normalize(p.text).removeprefix("• ") for p in docs["Highlights"].paragraphs if normalize(p.text).startswith("•")]
    add("Highlights count 3–5", 3 <= len(highlights) <= 5, f"count={len(highlights)}")
    add("Each highlight ≤85 characters", all(len(text) <= 85 for text in highlights), str([len(text) for text in highlights]))
    cover_text = all_text(docs["Cover"])
    add("Cover letter targets EJRH", "Journal of Hydrology: Regional Studies" in cover_text and "Water Research readers" not in cover_text, "journal title and audience")
    submission_text = "\n".join(all_text(docs[name]) for name in ["MS", "SI", "Response", "Cover"])
    add("Submission cites tagged public repository", "revision-2026-08-24" in submission_text, "stable revision tag")
    add("Submission does not promise a code ZIP", "Code_Supplement" not in submission_text and "code supplement" not in submission_text.lower(), "repository-only distribution")
    add("GA contains one revised image", len(docs["GA"].inline_shapes) == 1, f"images={len(docs['GA'].inline_shapes)}")

    paragraph_rows = audit_paragraphs(docs)
    add("Paragraph/table-row scan has no automatic flags", all(r["status"] == "PASS" for r in paragraph_rows), f"checks={Counter(r['status'] for r in paragraph_rows)}")
    result = {"checks": checks, "references": refs, "supplement_crossrefs": supplement_refs, "response": response, "paragraph_rows": paragraph_rows}
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Submission consistency audit — 2026-08-24",
        "",
        "This report records paragraph-, table-row-, and reviewer-item-level checks for the final revision package.",
        "",
        "## Executive checks",
        "",
        "| Status | Check | Detail |",
        "|---|---|---|",
    ]
    for check in checks:
        detail = str(check["detail"]).replace("|", "\\|")
        lines.append(f"| {check['status']} | {check['name']} | {detail} |")
    lines.extend(["", "## Reviewer coverage matrix", "", "| Status | Item | Coverage groups | Blue runs |", "|---|---|---:|---:|"])
    for row in response["matrix"]:
        lines.append(f"| {row['status']} | {row['id']} | {sum(row['coverage_groups'])}/{len(row['coverage_groups'])} | {row['blue_runs']} |")
    lines.extend(["", "## Paragraph and table-row review", "", "| Status | Scope | ID | Style | Flags | Text |", "|---|---|---|---|---|---|"])
    for row in paragraph_rows:
        text = row["text"].replace("|", "\\|")
        flags = "; ".join(row["flags"]).replace("|", "\\|")
        lines.append(f"| {row['status']} | {row['scope']} | {row['id']} | {row['style']} | {flags} | {text} |")
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps(Counter(check["status"] for check in checks), ensure_ascii=False))
    for check in checks:
        if check["status"] != "PASS":
            print(f"CHECK: {check['name']}: {check['detail']}")
    print(OUT_MD)


if __name__ == "__main__":
    from docx.oxml.ns import qn
    main()
