#!/usr/bin/env python3
"""
按 Setup A 重新运算结果, 更新 Manuscript_Final_MC-SIRC.docx 与 SI_Final.docx
直接替换错误的数据。
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import pandas as pd
from pathlib import Path
from copy import deepcopy
import docx
from docx.shared import Pt
import shutil

ROOT = Path(r'D:/shanxi')
SETUP_A_FILE = ROOT / 'output' / 'results' / 'setup_A_完整结果.xlsx'

MS_DOCX = ROOT / 'paper' / 'Manuscript_Final_MC-SIRC.docx'
SI_DOCX = ROOT / 'paper' / 'SI_Final.docx'

# 备份
shutil.copy2(MS_DOCX, MS_DOCX.with_suffix('.docx.bak'))
shutil.copy2(SI_DOCX, SI_DOCX.with_suffix('.docx.bak'))
print(f"备份: {MS_DOCX}.bak, {SI_DOCX}.bak")

# ────────────── 工具函数: 替换单元格内容（保留格式）──────────────
def set_cell(cell, text):
    """替换 cell 文本, 保留第一个 run 的字体格式。
    注意: cell.paragraphs[i] 每次返回新 wrapper, 用 idx 比较而非对象。
    """
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        return
    # 第一个段落: 写入新文本到第一个 run, 清空其他 run
    first_para = paragraphs[0]
    first_runs = list(first_para.runs)
    if first_runs:
        first_runs[0].text = str(text)
        for r in first_runs[1:]:
            r.text = ''
    else:
        first_para.add_run(str(text))
    # 其他段落: 清空所有 run
    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ''

def replace_table_data(table, data_rows, start_row=1):
    """data_rows: list of list of strings, replaces from start_row onward"""
    for ri, row_data in enumerate(data_rows):
        target_row_idx = start_row + ri
        if target_row_idx >= len(table.rows):
            break
        row = table.rows[target_row_idx]
        for ci, val in enumerate(row_data):
            if ci < len(row.cells):
                set_cell(row.cells[ci], val)

# ────────────── 读取 Setup A 数据 ──────────────
xls = pd.ExcelFile(SETUP_A_FILE)
mcmc = {}
for p_en, p_zh in [('COD','COD'), ('NH3N','氨氮'), ('TN','总氮'), ('TP','总磷')]:
    mcmc[p_en] = pd.read_excel(SETUP_A_FILE, f'MCMC_{p_zh}').set_index('参数')

sens = {}
for p_en, p_zh in [('COD','COD'), ('NH3N','氨氮'), ('TN','总氮'), ('TP','总磷')]:
    sens[p_en] = pd.read_excel(SETUP_A_FILE, f'先验敏感性_{p_zh}').set_index('情景')

sigma_obs = pd.read_excel(SETUP_A_FILE, 'σ_obs敏感性')
ratings   = pd.read_excel(SETUP_A_FILE, '可靠性评级')

# 工具: 数字格式化
def f3(x):
    if pd.isna(x): return '—'
    return f'{x:.3f}'
def f2(x):
    if pd.isna(x): return '—'
    return f'{x:.2f}'
def fpct(x, sign=True):
    if pd.isna(x): return '—'
    s = '+' if (x >= 0 and sign) else ''
    return f'{s}{x:.1f}%'
def fkg(x):
    if pd.isna(x): return '—'
    return f'{int(round(x)):,}'

# ────────────── 计算实际原始月度负荷（用于 SI Table S5）──────────────
print("\n计算实际月度负荷...")
df_mon = pd.read_excel(ROOT/'data'/'processed'/'monitor_2022_cleaned_v2.xlsx')
df_mon['月'] = pd.to_datetime(df_mon['监测时间']).dt.month
def hr_load(g, col): return (g[col]*g['瞬时流量(m³/s)']*3.6).sum()

monthly = []
for m in range(1, 13):
    g = df_mon[df_mon['月']==m]
    monthly.append({
        'Month': ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'][m-1],
        'COD': hr_load(g,'化学需氧量(mg/L)')/1000,
        'NH3N': hr_load(g,'氨氮(mg/L)')/1000,
        'TN': hr_load(g,'总氮(mg/L)')/1000,
        'TP': hr_load(g,'总磷(mg/L)')/1000,
        'Records': len(g)
    })
df_monthly = pd.DataFrame(monthly)
print(df_monthly.to_string())

# ────────────── 修改 SI 文档 ──────────────
print("\n" + "="*80)
print("更新 SI_Final.docx")
print("="*80)
si = docx.Document(str(SI_DOCX))

# SI Table 5 (索引 4): 月度负荷 — 使用原始数据
print("\n[SI Table S5] 月度负荷 (原始值, 不缩放)")
t = si.tables[4]
data = []
for _, r in df_monthly.iterrows():
    data.append([r['Month'], f"{r['COD']:.2f}", f"{r['NH3N']:.3f}",
                 f"{r['TN']:.2f}", f"{r['TP']:.4f}", str(r['Records'])])
# Annual row
total_cod  = df_monthly['COD'].sum()
total_nh   = df_monthly['NH3N'].sum()
total_tn   = df_monthly['TN'].sum()
total_tp   = df_monthly['TP'].sum()
total_rec  = df_monthly['Records'].sum()
data.append(['Annual', f"{total_cod:.2f}", f"{total_nh:.2f}",
             f"{total_tn:.2f}", f"{total_tp:.3f}", str(total_rec)])
replace_table_data(t, data)
print(f"  Annual = COD {total_cod:.2f}, NH3N {total_nh:.3f}, TN {total_tn:.2f}, TP {total_tp:.4f}, Records {total_rec}")

# SI Table 16 (索引 15): MCMC COD
print("\n[SI Table S16a] MCMC posteriors – COD")
t = si.tables[15]
src_order = ['面-农村生活污染源','畜禽散养','面-城市面源','面-城镇散排','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
en_names  = ['Rural domestic','Household livestock','Urban NPS','Dispersed urban','Large livestock','Industrial','Centralized fac.']
data = []
df = mcmc['COD']
for src, en in zip(src_order, en_names):
    if src in df.index:
        r = df.loc[src]
        data.append([en, f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
# Unknown
r = df.loc['Unknown']
data.append(['Unknown (t)', f2(r['均值']), f2(r['中位数']), f2(r['标准差']), f2(r['95%CI下限']), f2(r['95%CI上限'])])
replace_table_data(t, data)

# SI Table 17 (索引 16): MCMC NH3N
print("[SI Table S16b] MCMC posteriors – NH3-N")
t = si.tables[16]
src_order = ['面-农村生活污染源','畜禽散养','面-城市面源','面-城镇散排','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
data = []
df = mcmc['NH3N']
for src, en in zip(src_order, en_names):
    if src in df.index:
        r = df.loc[src]
        data.append([en, f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
r = df.loc['Unknown']
data.append(['Unknown (t)', f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
replace_table_data(t, data)

# SI Table 18 (索引 17): MCMC TN
print("[SI Table S16c] MCMC posteriors – TN")
t = si.tables[17]
src_order = ['面-农村生活污染源','面-农业面源','畜禽散养','面-城市面源','规模畜禽养殖','点-集中式污染治理设施']
en_names_tn = ['Rural domestic','Agricultural NPS','Household livestock','Urban NPS','Large livestock','Centralized fac.']
data = []
df = mcmc['TN']
for src, en in zip(src_order, en_names_tn):
    if src in df.index:
        r = df.loc[src]
        data.append([en, f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
r = df.loc['Unknown']
data.append(['Unknown (t)', f2(r['均值']), f2(r['中位数']), f2(r['标准差']), f2(r['95%CI下限']), f2(r['95%CI上限'])])
replace_table_data(t, data)

# SI Table 19 (索引 18): MCMC TP
print("[SI Table S16d] MCMC posteriors – TP")
t = si.tables[18]
src_order = ['面-农村生活污染源','面-农业面源','畜禽散养','面-城市面源','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
en_names_tp = ['Rural domestic','Agricultural NPS','Household livestock','Urban NPS','Large livestock','Industrial','Centralized fac.']
data = []
df = mcmc['TP']
for src, en in zip(src_order, en_names_tp):
    if src in df.index:
        r = df.loc[src]
        data.append([en, f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
r = df.loc['Unknown']
data.append(['Unknown (t)', f3(r['均值']), f3(r['中位数']), f3(r['标准差']), f3(r['95%CI下限']), f3(r['95%CI上限'])])
replace_table_data(t, data)

# SI Tables 20-23 (索引 19-22): 先验敏感性
def make_sens_data(p_key, src_order_keys):
    df = sens[p_key]
    data = []
    for sc in ['S1 Low','S2 Default','S3 High','S4 Weak','S5 Uninf.']:
        if sc not in df.index: continue
        r = df.loc[sc]
        row = [sc, fpct(r['偏差%'])]
        for k in src_order_keys:
            if k in df.columns:
                row.append(f3(r[k]))
            else:
                row.append('—')
        data.append(row)
    return data

# Table 20 (idx 19): S20 COD prior sensitivity (cols: Rural, Livest-h, Urban, Disp., Large liv., Industrial, Central)
print("\n[SI Table S21 (S20 in docx)] Prior sensitivity – COD")
t = si.tables[19]
keys = ['面-农村生活污染源','畜禽散养','面-城市面源','面-城镇散排','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
replace_table_data(t, make_sens_data('COD', keys))

# Table 21 (idx 20): S21 NH3N prior sensitivity (cols: Rural, Livest-h, Urban, Disp., Large liv., Industrial, Central)
print("[SI Table S22] Prior sensitivity – NH3-N")
t = si.tables[20]
replace_table_data(t, make_sens_data('NH3N', keys))

# Table 22 (idx 21): S22 TN prior sensitivity (cols: Rural, Agri., Livest-h, Urban, Large liv., Central)
print("[SI Table S23] Prior sensitivity – TN")
t = si.tables[21]
keys_tn = ['面-农村生活污染源','面-农业面源','畜禽散养','面-城市面源','规模畜禽养殖','点-集中式污染治理设施']
replace_table_data(t, make_sens_data('TN', keys_tn))

# Table 23 (idx 22): S23 TP prior sensitivity (cols: Rural, Agri., Livest-h, Urban, Large liv., Industrial, Central)
print("[SI Table S24] Prior sensitivity – TP")
t = si.tables[22]
keys_tp = ['面-农村生活污染源','面-农业面源','畜禽散养','面-城市面源','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
replace_table_data(t, make_sens_data('TP', keys_tp))

# Table 24 (idx 23): Reliability ratings (use z>2 rule for A; 1<z<=2 for B; z<=1 for C)
print("\n[SI Table S25 (S24 in docx)] Reliability ratings (z>2 rule)")

# 重新评级 (z>2 = A, 1<z≤2 = B, z≤1 = C)
def get_rating_z(z):
    if z > 2: return 'A'
    elif z > 1: return 'B'
    else: return 'C'

ratings_new = ratings.copy()
ratings_new['评级'] = ratings_new['shift'].apply(get_rating_z)
counts = ratings_new['评级'].value_counts()
print(f"  新评级 (z>2 rule): A={counts.get('A',0)}, B={counts.get('B',0)}, C={counts.get('C',0)}")

# 排序: A 优先, B 次之, 同评级内按 shift 降序
rating_order = {'A':0, 'B':1, 'C':2}
ratings_sorted = ratings_new.copy()
ratings_sorted['_order'] = ratings_sorted['评级'].map(rating_order)
ratings_sorted = ratings_sorted.sort_values(['_order', 'shift'], ascending=[True, False])

# Top 12 行
top12 = ratings_sorted.head(12)

src_to_en = {
    '面-农村生活污染源':'Rural domestic',
    '面-农业面源':'Agricultural NPS',
    '畜禽散养':'Household livestock',
    '面-水产养殖':'Aquaculture',
    '面-城市面源':'Urban NPS',
    '面-城镇散排':'Dispersed urban',
    '规模畜禽养殖':'Large livestock',
    '点-工业源':'Industrial',
    '点-集中式污染治理设施':'Centralized fac.',
}
pol_to_en = {'COD':'COD', '氨氮':'NH₃-N', '总氮':'TN', '总磷':'TP'}

t = si.tables[23]
data = []
for _, r in top12.iterrows():
    data.append([
        src_to_en.get(r['污染源'], r['污染源']),
        pol_to_en.get(r['污染物'], r['污染物']),
        f"{r['先验μ']:.2f}",
        f3(r['MAP f']),
        f"{r['shift']:.2f}",
        r['评级']
    ])
replace_table_data(t, data)

# Table 25 (idx 24): σ_obs sensitivity
print("\n[SI Table S26 (S25 in docx)] σ_obs sensitivity")
t = si.tables[24]
data = []
sig_keys = ['面-农村生活污染源','面-城市面源','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
sig_en   = ['Rural','Urban NPS','Large livest.','Industrial','Central fac.']
for p_zh, p_en in [('COD','COD'), ('总磷','TP')]:
    sub = sigma_obs[sigma_obs['污染物']==p_zh]
    for _, r in sub.iterrows():
        row = [p_en, r['σ_obs']]
        for k in sig_keys:
            if k in sub.columns:
                row.append(f3(r[k]))
            else:
                row.append('—')
        data.append(row)
replace_table_data(t, data)

# 保存 SI
si_out = SI_DOCX
si.save(str(si_out))
print(f"\n✓ SI 已保存到: {si_out}")

# ────────────── 修改 Manuscript ──────────────
print("\n" + "="*80)
print("更新 Manuscript_Final_MC-SIRC.docx")
print("="*80)
ms = docx.Document(str(MS_DOCX))

# Manuscript Table 5 (索引 4): MCMC posteriors for key sources
print("\n[Manuscript Table 5] MCMC posteriors")
t = ms.tables[4]
key_src = [
    ('面-农村生活污染源','Rural domestic'),
    ('规模畜禽养殖','Large livestock'),
    ('点-集中式污染治理设施','Central facility'),
    ('点-工业源','Industrial'),
    ('面-城市面源','Urban NPS'),
]
def ci_str(lo, hi, prec=2):
    if pd.isna(lo) or pd.isna(hi): return '—'
    return f"{lo:.{prec}f}–{hi:.{prec}f}"

data = []
for src, en in key_src:
    row = [en]
    for p_en in ['COD','NH3N','TN','TP']:
        df = mcmc[p_en]
        if src in df.index:
            r = df.loc[src]
            row.append(f3(r['均值']))
            row.append(ci_str(r['95%CI下限'], r['95%CI上限']))
        else:
            row.append('—')
            row.append('—')
    data.append(row)
replace_table_data(t, data)

# Manuscript Table 6 (索引 5): Prior sensitivity for COD and TP, S1, S2, S3, S5
print("[Manuscript Table 6] Prior sensitivity – key scenarios")
t = ms.tables[5]
key_keys = ['面-农村生活污染源','面-城市面源','规模畜禽养殖','点-工业源','点-集中式污染治理设施']
data = []
for p_zh, p_en in [('COD','COD'),('总磷','TP')]:
    df = sens[{'COD':'COD','总磷':'TP'}[p_zh]]
    for sc, sc_en in [('S1 Low','S1 Low'),('S2 Default','S2 Default'),('S3 High','S3 High'),('S5 Uninf.','S5 Uninf.')]:
        if sc not in df.index: continue
        r = df.loc[sc]
        row = [p_en, sc_en, fpct(r['偏差%'])]
        for k in key_keys:
            if k in df.columns:
                row.append(f3(r[k]))
            else:
                row.append('—')
        data.append(row)
replace_table_data(t, data)

# ─── 文字修改 ───
print("\n[Manuscript text] §3.2 与结论部分文字修改")

new_a = int(counts.get('A', 0))
new_b = int(counts.get('B', 0))
new_c = int(counts.get('C', 0))

# 找到 TN 强相关 r 值
import numpy as np
tn_corr = pd.read_excel(SETUP_A_FILE, 'MCMC_corr_总氮', index_col=0)
tn_pairs = []
for i in range(len(tn_corr)):
    for j in range(i+1, len(tn_corr)):
        tn_pairs.append((tn_corr.index[i], tn_corr.columns[j], tn_corr.iloc[i,j]))
tn_pairs.sort(key=lambda x: -abs(x[2]))
strongest_tn = tn_pairs[0]
print(f"  TN 最强相关: {strongest_tn[0]} ↔ {strongest_tn[1]}, r = {strongest_tn[2]:+.2f}")

text_replacements = [
    # §3.2 评级数量
    ('classified 7 source-pollutant pairs as Rating A',
     f'classified {new_a} source-pollutant pairs as Rating A'),
    ('6 as Rating B',  f'{new_b} as Rating B'),
    ('14 as Rating C', f'{new_c} as Rating C'),
    # 评级数量 in §3.4
    ('Rating A sources (7 pairs',  f'Rating A sources ({new_a} pairs'),
    ('Rating B sources (6 pairs',  f'Rating B sources ({new_b} pairs'),
    ('Rating C sources (14 pairs', f'Rating C sources ({new_c} pairs'),
    # TN 相关 r 值
    ('strong negative correlation (r = −0.59)',
     f'strong negative correlation (r = {strongest_tn[2]:.2f})'.replace('-','−')),
    ('strong negative correlation (r = -0.59)',
     f'strong negative correlation (r = {strongest_tn[2]:.2f})'),
    # 7 source-pollutant pairs 在结论里
]

for para in ms.paragraphs:
    for old, new in text_replacements:
        if old in para.text:
            # 替换 run-by-run, 因 docx 中文本可能跨多个 run
            # 简化: 直接合并 run 文本, 替换, 写回第一个 run
            full_text = ''.join(run.text for run in para.runs)
            if old in full_text:
                new_text = full_text.replace(old, new)
                # 先清空所有 run
                for run in para.runs:
                    run.text = ''
                # 写到第一个 run
                if para.runs:
                    para.runs[0].text = new_text
                print(f"  文字替换: '{old[:50]}...' → '{new[:50]}...'")
                break  # 一段只替换一次

# 保存 Manuscript
ms.save(str(MS_DOCX))
print(f"\n✓ Manuscript 已保存到: {MS_DOCX}")

print("\n" + "="*80)
print("完成。备份在 .docx.bak")
print("="*80)
print(f"\n新评级数量: A={new_a}, B={new_b}, C={new_c}")
