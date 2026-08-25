#!/usr/bin/env python3
"""Build and finalize the 2026-08-25 submission package from reproducible sources."""

from __future__ import annotations

from calendar import monthrange
from copy import deepcopy
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree as ET
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.reporting import build_revision3_documents as base
from scripts.reporting import build_revision3_response_letter as response_base
from scripts.reporting import build_revision4_response_letter as response_v4
from scripts.reporting import build_revision4_submission as submission
from scripts.reporting import augment_revision4_response_with_figures_tables as augment
from scripts.reporting import generate_revision3_figures as old_figures


PAPER = ROOT / "paper"
FIG = ROOT / "output" / "figures" / "revision3"
FIG_S2 = ROOT / "output" / "figures" / "figureS2_monthly_loads"
GA_IMAGE = ROOT / "output" / "figures" / "revision4" / "graphical_abstract_identifiability.png"
ARCHIVE = PAPER / "_archive_non_submission_20260824" / "finalization_20260825"
MCMC_DIAGNOSTICS = ROOT / "output" / "results" / "revision5_mcmc_diagnostics.xlsx"
TAG = "revision-2026-08-25.1"

POLLUTANT_MAP = {"COD": "COD", "氨氮": "NH₃-N", "总氮": "TN", "总磷": "TP"}
SOURCE_MAP = {
    "面-农村生活污染源": "Rural domestic",
    "面-农业面源": "Agricultural NPS",
    "畜禽散养": "Household livestock",
    "面-水产养殖": "Aquaculture",
    "面-城市面源": "Urban NPS",
    "面-城镇散排": "Dispersed urban",
    "规模畜禽养殖": "Large livestock",
    "点-工业源": "Industrial",
    "点-集中式污染治理设施": "Central facility",
    "Unknown": "Unknown (t)",
}

ABSTRACT_PARTS = [
    (
        "Study Region",
        "The Nanchuan River Basin (1,438 km²) is a hilly-gully Loess Plateau catchment in China. "
        "A 2022 bottom-up inventory resolved over 1,588 unique 1-km grid IDs and 100 georeferenced "
        "source records is paired with 5,928 observed hourly records at one outlet station.",
    ),
    (
        "Study Focus",
        "Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) is an "
        "identifiability-aware workflow. Four scenarios represent missing-not-at-random gaps; Bayesian "
        "maximum a posteriori and Markov chain Monte Carlo analyses allocate aggregate discrepancies under "
        "explicit priors; Jacobian-rank, null-space, and k = 0 profile diagnostics test estimability; and "
        "forward Monte Carlo simulation describes inventory-side uncertainty.",
    ),
    (
        "New Hydrological Insights for the Region",
        "The annual outlet constraint identifies one aggregate linear combination of 9–10 unknowns "
        "(Jacobian rank 1; nullity 8–9), precluding source-specific coefficient validation. Four boundary "
        "prior conflicts persist across missing-data scenarios and define field-audit priorities. After "
        "profiling the global scale, the monthly distance-decay model adds virtually no explanatory power "
        "over a k = 0 flow-only null (ΔR² ≤ 0.0033), leaving half-life distances and effective-contribution "
        "rankings unresolved. The supported regional result is an aggregate inventory–monitoring discrepancy "
        "and a monitoring-design diagnosis: tributary sections and observed source seasonality are needed "
        "for source-specific transport inference.",
    ),
]


def panel_label(ax, label):
    ax.text(-0.12, 1.05, label, transform=ax.transAxes, fontsize=13, fontweight="bold", va="top")


def workflow_box(ax, y, text, color, height=0.13, fontsize=17.5, edgecolor="#444444"):
    patch = FancyBboxPatch(
        (0.04, y), 0.92, height,
        boxstyle="round,pad=0.012,rounding_size=0.015",
        facecolor=color, edgecolor=edgecolor, linewidth=1.35,
    )
    ax.add_patch(patch)
    label = ax.text(
        0.50, y + height / 2, text,
        ha="center", va="center", fontsize=fontsize, fontweight="bold", linespacing=1.08,
    )
    return patch, label


def generate_figure1():
    map_image = Image.open(ROOT / "output" / "figures" / "figure1a_watershed_standalone" / "figure1a_watershed.png")
    fig, axes = plt.subplots(1, 2, figsize=(18, 8.5), gridspec_kw={"width_ratios": [1.30, 1.20]})
    axes[0].imshow(map_image)
    axes[0].axis("off")
    axes[0].set_title("(a) Nanchuan River Basin: inventory sources and outlet station", loc="left", fontsize=18, fontweight="bold")
    ax = axes[1]
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title("(b) Identifiability-aware MC-SIRC workflow", loc="left", fontsize=18, fontweight="bold")
    entries = [
        (0.815, "Inventory and outlet monitoring\n1,588 unique 1-km grid IDs\n100 sources | 5,928 observed hours", "#fdd49e"),
        (0.650, "Missing-data scenarios\nMonthly default and lower-bound\nEnhanced-gap sensitivity", "#c7e9c0"),
        (0.485, "Prior-regularized reconciliation\nMAP/MCMC source allocations\nwithin a common evidence system", "#c6dbef"),
        (0.320, "Formal identifiability audit\nJacobian rank | null-space solutions\nk = 0 profile", "#dadaeb"),
        (0.155, "Supported outputs\nAggregate discrepancy | prior-conflict flags\nMonitoring priorities", "#fcbba1"),
    ]
    framed_text = []
    for y, text, color in entries:
        framed_text.append(workflow_box(ax, y, text, color))
    for current, following in zip(entries[:-1], entries[1:]):
        ax.add_patch(FancyArrowPatch(
            (0.5, current[0]), (0.5, following[0] + 0.13), arrowstyle="-|>", mutation_scale=18,
            linewidth=1.5, color="#444444",
        ))
    framed_text.append(workflow_box(
        ax, 0.010,
        "Not resolved at current monitoring resolution\n"
        "Source-specific attenuation and half-lives\n"
        "Effective outlet shares and policy re-ranking",
        "#fff5f0", height=0.11, fontsize=15.0, edgecolor="#8b0000",
    ))
    fig.tight_layout(pad=0.8)
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    for patch, label in framed_text:
        frame = patch.get_window_extent(renderer)
        text_extent = label.get_window_extent(renderer)
        if not (
            text_extent.x0 >= frame.x0 + 4
            and text_extent.x1 <= frame.x1 - 4
            and text_extent.y0 >= frame.y0 + 3
            and text_extent.y1 <= frame.y1 - 3
        ):
            raise RuntimeError(f"Panel 1b text exceeds its frame: {label.get_text()!r}")
    FIG.mkdir(parents=True, exist_ok=True)
    fig.savefig(FIG / "figure1_revised_framework.png", dpi=400, bbox_inches="tight")
    fig.savefig(FIG / "figure1_revised_framework.pdf", bbox_inches="tight")
    plt.close(fig)


def generate_figure2():
    workbook = ROOT / "output" / "results" / "revision3_identifiability_and_missingness.xlsx"
    detail = pd.read_excel(workbook, sheet_name="Missing_detail")
    summary = pd.read_excel(workbook, sheet_name="Missing_summary")
    flags = pd.read_excel(workbook, sheet_name="Flag_persistence")
    rank = pd.read_excel(workbook, sheet_name="Annual_rank_nullity")
    default = detail[detail.Scenario == "S2 month-specific (default)"].copy()
    sources = list(SOURCE_MAP)[:-1]
    pollutants = ["COD", "氨氮", "总氮", "总磷"]
    matrix = np.full((4, len(sources)), np.nan)
    for i, pollutant in enumerate(pollutants):
        for j, source in enumerate(sources):
            match = default[(default.Pollutant == pollutant) & (default.Source == source)]
            if not match.empty:
                matrix[i, j] = match.iloc[0].MAP_factor

    fig, axes = plt.subplots(2, 2, figsize=(14, 10.2), constrained_layout=True)
    ax = axes[0, 0]
    cmap = plt.get_cmap("Blues_r").copy()
    cmap.set_bad("#d9d9d9")
    image = ax.imshow(matrix, cmap=cmap, vmin=0.1, vmax=1.2, aspect="auto")
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            if np.isfinite(matrix[i, j]):
                ax.text(j, i, f"{matrix[i, j]:.2f}", ha="center", va="center", fontsize=8,
                        color="white" if matrix[i, j] < 0.45 else "black")
            else:
                ax.text(j, i, "NA", ha="center", va="center", fontsize=8, color="#555555")
    ax.set_xticks(range(len(sources)), [SOURCE_MAP[item] for item in sources], rotation=48, ha="right")
    ax.set_yticks(range(4), [POLLUTANT_MAP[item] for item in pollutants])
    ax.set_title("Regularized discrepancy factors (default missing-data scenario)")
    fig.colorbar(image, ax=ax, fraction=0.046, pad=0.03, label="MAP factor (prior-regularized)")
    panel_label(ax, "a")

    ax = axes[0, 1]
    default_summary = summary[summary.Scenario == "S2 month-specific (default)"].copy()
    x = np.arange(4)
    a = default_summary.Tier_A_count.to_numpy()
    b = default_summary.Tier_B_count.to_numpy()
    c = default_summary.Tier_C_count.to_numpy()
    ax.bar(x, a, color="#b2182b", label="A: strong prior conflict")
    ax.bar(x, b, bottom=a, color="#ef8a62", label="B: moderate prior conflict")
    ax.bar(x, c, bottom=a + b, color="#bdbdbd", label="C: no resolved conflict")
    ax.set_xticks(x, [POLLUTANT_MAP[item] for item in default_summary.Pollutant])
    ax.set_ylabel("Number of source–pollutant pairs")
    ax.set_title("Diagnostic tiers describe prior conflict, not identifiability")
    ax.legend(frameon=False, loc="upper left")
    panel_label(ax, "b")

    ax = axes[1, 0]
    scenarios = ["S1 observed-only", "S2 month-specific (default)", "S3 annual-mean", "S4 event-weighted upper"]
    short = ["S1\nobserved", "S2\nmonthly", "S3\nannual mean", "S4\ngap +50%"]
    colors = ["#2166ac", "#4393c3", "#d6604d", "#b2182b"]
    markers = ["o", "s", "^", "D"]
    for idx, (_, row) in enumerate(flags.iterrows()):
        values = [row[item + " factor"] for item in scenarios]
        label = f"{POLLUTANT_MAP[row.Pollutant]}—{SOURCE_MAP[row.Source]}"
        ax.plot(range(4), values, marker=markers[idx], ms=7, lw=1.8, color=colors[idx], label=label)
    ax.axhline(0.1, color="black", lw=0.9, ls="--", label="Truncation bound")
    ax.set_xticks(range(4), short)
    ax.set_ylim(0.06, 0.25)
    ax.set_ylabel("Regularized factor")
    ax.set_title("Persistent boundary conflicts across missing-data treatments")
    ax.legend(frameon=False, loc="upper center", ncol=2)
    panel_label(ax, "c")

    ax = axes[1, 1]
    labels = [POLLUTANT_MAP[item] for item in rank.Pollutant]
    params = rank.Active_source_factors + rank.Unknown_source_terms
    ax.bar(np.arange(4) - 0.18, params, width=0.36, color="#7b3294", label="Unknown parameters")
    ax.bar(np.arange(4) + 0.18, rank.Jacobian_rank, width=0.36, color="#008837", label="Jacobian rank")
    for i, nullity in enumerate(rank.Nullity):
        ax.text(i - 0.18, params.iloc[i] + 0.15, f"nullity={int(nullity)}", ha="center", fontsize=8.5)
    ax.set_xticks(range(4), labels)
    ax.set_ylabel("Dimension")
    ax.set_ylim(0, max(params) + 2)
    ax.set_title("One annual observation identifies one aggregate combination")
    ax.legend(frameon=False, loc="upper right")
    panel_label(ax, "d")
    fig.savefig(FIG / "figure2_identifiability_aware_bayesian.png", dpi=400, bbox_inches="tight")
    fig.savefig(FIG / "figure2_identifiability_aware_bayesian.pdf", bbox_inches="tight")
    plt.close(fig)


def generate_figure4():
    workbook = ROOT / "output" / "results" / "spatial_identifiability_unbounded.xlsx"
    profile = pd.read_excel(workbook, sheet_name="Profile")
    summary = pd.read_excel(workbook, sheet_name="Summary")
    fig, axes = plt.subplots(2, 2, figsize=(14, 8.3), constrained_layout=True)
    for idx, (ax, pollutant) in enumerate(zip(axes.flat, ["COD", "氨氮", "总氮", "总磷"])):
        data = profile[profile.Pollutant == pollutant].copy()
        ratio = data.relative_SSE_free / data.relative_SSE_free.min()
        ax.plot(data["k_km-1"], ratio, color="#2166ac", lw=2.2)
        ax.axvline(0, color="#b2182b", lw=1.5, ls="--", label="k = 0 flow-only null")
        row = summary[summary.Pollutant == pollutant].iloc[0]
        ax.axvline(row.k_best_free, color="#4d4d4d", lw=1.4, ls=":", label="Profile-minimum k")
        ax.text(0.03, 0.94, f"ΔR² vs k=0 = {row.Delta_R2_free:.4f}", transform=ax.transAxes,
                va="top", bbox=dict(facecolor="white", edgecolor="#cccccc", alpha=0.9))
        ax.set_title(POLLUTANT_MAP[pollutant])
        ax.set_xlabel("Attenuation parameter k (km⁻¹)")
        ax.set_ylabel("Profile relative SSE / minimum")
        pad = max((ratio.max() - ratio.min()) * 0.15, 0.00002)
        ax.set_ylim(ratio.min() - pad, ratio.max() + pad)
        ax.ticklabel_format(axis="y", style="plain", useOffset=False)
        ax.legend(frameon=False, loc="lower right")
        panel_label(ax, chr(ord("a") + idx))
    fig.suptitle("Spatial profile audit with the global scale profiled at each k", fontsize=14)
    fig.savefig(FIG / "figure4_spatial_identifiability_profile.png", dpi=400, bbox_inches="tight")
    fig.savefig(FIG / "figure4_spatial_identifiability_profile.pdf", bbox_inches="tight")
    plt.close(fig)


def generate_figure_s2():
    data = pd.read_csv(FIG_S2 / "data_monthly_loads.csv")
    days = np.array([monthrange(2022, month)[1] for month in range(1, 13)])
    coverage = data.records_h.to_numpy() / (days * 24)
    excluded = coverage < 0.50
    fig, axes = plt.subplots(2, 2, figsize=(13, 8.8), constrained_layout=True)
    x = np.arange(12)
    for idx, (ax, (column, label)) in enumerate(zip(axes.flat, [("COD", "COD"), ("NH3N", "NH₃-N"), ("TN", "TN"), ("TP", "TP")])):
        for month_index in np.where(excluded)[0]:
            ax.axvspan(month_index - 0.5, month_index + 0.5, color="#d9d9d9", alpha=0.65, zorder=0)
        ax.plot(x, data[column], marker="o", lw=2.0, color="#2166ac", label="Observed-hour load (S1)")
        ax.set_xticks(x, data.month, rotation=45, ha="right")
        ax.set_ylabel("Load (t month⁻¹)")
        ax.set_title(label)
        ax2 = ax.twinx()
        ax2.plot(x, coverage * 100, marker="s", ms=4.5, lw=1.3, color="#b2182b", label="Timestamp coverage")
        ax2.axhline(50, color="#b2182b", ls="--", lw=1.0)
        ax2.set_ylim(0, 105)
        ax2.set_ylabel("Timestamp coverage (%)")
        if idx == 0:
            first = ax.get_legend_handles_labels()
            second = ax2.get_legend_handles_labels()
            ax.legend(first[0] + second[0], first[1] + second[1], frameon=False, loc="upper left")
        panel_label(ax, chr(ord("a") + idx))
    fig.text(0.5, 0.005,
             "Gray months have <50% timestamp coverage and are excluded from the spatial-profile objective; "
             "all months remain represented in the S1–S4 annual-load scenarios.",
             ha="center", va="bottom", fontsize=10)
    fig.savefig(FIG_S2 / "figureS2_monthly_loads.png", dpi=400, bbox_inches="tight")
    fig.savefig(FIG_S2 / "figureS2_monthly_loads.pdf", bbox_inches="tight")
    plt.close(fig)


def generate_graphical_abstract():
    map_image = Image.open(ROOT / "output" / "figures" / "figure1a_watershed_standalone" / "figure1a_watershed.png")
    fig = plt.figure(figsize=(13.3, 5.4))
    grid = fig.add_gridspec(1, 3, width_ratios=[1.25, 1.0, 1.15], wspace=0.08)
    ax = fig.add_subplot(grid[0, 0])
    ax.imshow(map_image)
    ax.axis("off")
    ax.set_title("Nanchuan River Basin\n1,438 km² Loess Plateau catchment", fontsize=12, fontweight="bold")
    ax.text(0.5, -0.02, "1,588 unique 1-km grid IDs\n100 georeferenced sources • one outlet station",
            transform=ax.transAxes, ha="center", va="top", fontsize=10)

    ax = fig.add_subplot(grid[0, 1])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title("MC-SIRC evidence chain", fontsize=12, fontweight="bold")
    items = [
        (0.72, "Four missing-data\nscenarios", "#d9f0d3"),
        (0.44, "Prior-regularized\naggregate reconciliation", "#c6dbef"),
        (0.16, "Rank/null-space +\nk = 0 profile audit", "#dadaeb"),
    ]
    for y, text, color in items:
        patch = FancyBboxPatch((0.10, y), 0.80, 0.17, boxstyle="round,pad=0.02",
                               facecolor=color, edgecolor="#4d4d4d")
        ax.add_patch(patch)
        ax.text(0.50, y + 0.085, text, ha="center", va="center", fontsize=11, fontweight="bold")
    for start in [0.72, 0.44]:
        ax.add_patch(FancyArrowPatch((0.5, start), (0.5, start - 0.10), arrowstyle="-|>", mutation_scale=15,
                                     color="#4d4d4d"))

    ax = fig.add_subplot(grid[0, 2])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title("Supported regional inference", fontsize=12, fontweight="bold")
    results = [
        (0.77, "Aggregate inventory–outlet\ndiscrepancy quantified", "#fee8c8"),
        (0.52, "Rank 1; nullity 8–9\nsource factors unresolved", "#fdbb84"),
        (0.27, "ΔR² ≤ 0.0033 vs k = 0\ndistance decay unresolved", "#fc8d59"),
        (0.02, "Next measurements:\ntributary sections + source seasonality", "#ef6548"),
    ]
    for y, text, color in results:
        patch = FancyBboxPatch((0.07, y), 0.86, 0.17, boxstyle="round,pad=0.02",
                               facecolor=color, edgecolor="#7f2704")
        ax.add_patch(patch)
        ax.text(0.50, y + 0.085, text, ha="center", va="center", fontsize=10.5, fontweight="bold")
    fig.suptitle("What can one outlet station resolve?", fontsize=17, fontweight="bold", y=0.99)
    GA_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(GA_IMAGE, dpi=300, bbox_inches="tight", facecolor="white")
    fig.savefig(GA_IMAGE.with_suffix(".pdf"), bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_all_figures():
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 10.5})
    generate_figure1()
    generate_figure2()
    old_figures.generate_figure3()
    generate_figure4()
    generate_figure_s2()
    generate_graphical_abstract()


def find_prefix(doc, prefix, optional=False):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    if optional:
        return None
    raise KeyError(prefix)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def replace_prefix(doc, prefix, text, optional=False, body=False):
    paragraph = find_prefix(doc, prefix, optional=optional)
    if paragraph is None:
        return None
    paragraph.text = text
    if body:
        submission.set_body_style(paragraph)
    return paragraph


def set_affiliation_paragraph(paragraph, marker, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(marker)
    run.font.superscript = True
    paragraph.add_run(" " + text)


def set_title_page(doc):
    author = find_prefix(doc, "Yujie Wang", optional=True)
    if author is not None:
        clear_paragraph(author)
        authors = [
            ("Yujie Wang", "a"),
            ("Jinhong Luo", "c"),
            ("Weifeng Zhang", "a"),
            ("Cheng Zhang", "a"),
            ("Shifa Zhong", "b,*"),
        ]
        for index, (name, marker) in enumerate(authors):
            if index:
                author.add_run(", ")
            author.add_run(name)
            run = author.add_run(marker)
            run.font.superscript = True
    affiliations = [
        (("1.", "1 "), "a", "Shanxi Provincial Ecological Environment Planning and Technology Institute, Taiyuan 030009, China"),
        (("2.", "2 "), "b", "College of Environmental Science and Engineering, Tongji University, Shanghai 200092, China"),
        (("3.", "3 "), "c", "Shanxi Provincial Center for Ecological and Environmental Monitoring and Emergency Response, Taiyuan 030024, China"),
    ]
    for prefixes, marker, text in affiliations:
        paragraph = next((p for p in doc.paragraphs if p.text.strip().startswith(prefixes)), None)
        if paragraph is not None:
            set_affiliation_paragraph(paragraph, marker, text)


def rewrite_structured_abstract(doc):
    heading = find_prefix(doc, "Abstract")
    keywords = find_prefix(doc, "Keywords:")
    body = doc._element.body
    children = list(body)
    start = children.index(heading._p)
    end = children.index(keywords._p)
    for child in children[start + 1:end]:
        body.remove(child)
    for label, text in ABSTRACT_PARTS:
        paragraph = submission.paragraph_before(keywords)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(label + ": ").bold = True
        paragraph.add_run(text)


def postprocess_manuscript():
    path = PAPER / "Manuscript_Final_MC-SIRC.docx"
    doc = Document(path)
    set_title_page(doc)
    rewrite_structured_abstract(doc)
    study = find_prefix(doc, "The Nanchuan River Basin is located")
    if "Figure 1a" not in study.text:
        study.text = study.text.rstrip(".") + " (Figure 1a)."
        submission.set_body_style(study)
    workflow = find_prefix(doc, "Each MC-SIRC component")
    workflow.text = workflow.text.replace("(Table 1)", "(Figure 1b; Table 1)")
    submission.set_body_style(workflow)

    replace_prefix(
        doc,
        "The bottom-up 2022 inventory contains",
        "The bottom-up 2022 inventory uses a 1 km × 1 km grid framework for five area-source categories. "
        "Overlaying the grid with land-use and administrative units produces 53,155 polygon-intersection "
        "records assigned to 1,588 unique grid IDs (Table S1); 53,155 is therefore a record count rather than "
        "the number of independent 1-km cells. Four additional categories comprise 100 georeferenced facility "
        "or farm records: 50 household-livestock, 34 large-livestock, 15 industrial, and one centralized-treatment "
        "record. The inventory compilation and spatial-accounting protocol are described by Wang et al. (2026). "
        "Outlet data include hourly COD, NH₃-N, TN, TP, and discharge. The 2022 archive contains 5,928 unique "
        "hourly timestamps from a possible 8,760, giving 67.7% coverage and 2,832 absent hours. Missingness is "
        "concentrated in July–November, when monthly coverage falls to 9.3–62.8%, and is treated as missing not at random (MNAR).",
        body=True,
    )
    replace_prefix(
        doc,
        "Within observed timestamps",
        "Within observed timestamps, concentration gaps were linearly interpolated and implausible discharge "
        "values were replaced using a three-tier hierarchy: the same calendar time in other archive years, "
        "adjacent-month medians, and linear interpolation. Completely absent timestamps were handled by four "
        "annual-load scenarios: S1 observed hours only (lower bound); S2 month-specific coverage scaling (default); "
        "S3 annual-mean scaling by 8,760/5,928; and S4 an enhanced-gap sensitivity that adds 50% to only the "
        "S2-imputed portion. The 1.5 factor is a transparent stress-test choice, not an estimated event multiplier "
        "or confidence bound. S2 preserves the observed seasonal load structure and avoids exporting the high "
        "April–June mean to poorly observed late-flood-season months. The implications of the remaining low-coverage months are examined in the Discussion.",
        body=True,
    )
    replace_prefix(
        doc,
        "where α is the nominal inventory river-entry coefficient",
        "where α is the nominal inventory river-entry coefficient, f is a multiplicative discrepancy allocation, "
        "U is a non-negative unidentified-load term, and ε is observation/model discrepancy. Eight or nine active "
        "source factors plus U are estimated from one annual M. Truncated-normal priors f ~ TN(μ, σ²; [0.1, 2.0]) "
        "and U ~ Gamma(shape = 2, rate = 10/M), truncated at 0.5M, regularize the equation; the MAP calculation "
        "used 0.001M as a numerical lower bound for U. MAP estimates used 30 L-BFGS-B starts with seed 42. MCMC "
        "used the emcee ensemble sampler with 32 walkers, 20,000 steps, 5,000 burn-in steps, and seed 42 "
        "(Foreman-Mackey et al., 2013). The principal MAP analysis uses default S2. MCMC and prior/observation-error "
        "sensitivities use the higher-load S3 scenario as a conservative stress test of boundary pile-up; persistence "
        "across all S1–S4 scenarios is assessed separately. Split-Rhat, effective sample size, autocorrelation time, "
        "and acceptance fractions are reported in Tables S13–S16.",
        body=True,
    )
    replace_prefix(
        doc,
        "The same annual-mean protocol was also applied",
        "The same annual-mean protocol was applied to the 2020, 2021, and 2023 monitoring archives while holding "
        "the 2022 inventory fixed. This cross-year consistency check is conditional because the discrepancy factors "
        "absorb year-specific hydrology and unobserved source changes. Archive timestamp coverage and strict raw "
        "pollutant–flow pair coverage are reported separately: the former determines annual scaling, whereas the latter "
        "is only 1.0–1.3% in 2020–2021 before within-timestamp filling. Those two archives are therefore treated as data-poor.",
        body=True,
    )
    replace_prefix(
        doc,
        "The inventory contains",
        "The inventory contains 1,036.5 t COD, 20.0 t NH₃-N, 131.3 t TN reported across categories with available "
        "TN entries, and 14.0 t TP emissions, corresponding to nominal river-entry totals of 463.9, 6.84, 69.15, "
        "and 5.32 t, respectively (Table S8). Industrial TN is unavailable, so the TN totals are reported-category "
        "sums rather than complete inventory totals. At the source-emission scale, large-scale livestock accounts "
        "for 56.6% of COD and 65.1% of TP emissions, whereas the centralized facility accounts for 37.1% of the "
        "reported TN total. These are inventory-composition shares, not estimates of source contributions at the "
        "outlet. Under default S2, outlet loads are 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. "
        "These differences establish an aggregate inventory–outlet discrepancy but do not separate coefficient error "
        "from channel transformation or omitted sources.",
        body=True,
    )
    replace_prefix(
        doc,
        "The cross-year application provides",
        "The cross-year application provides a conditional comparison across archive years. The four S1–S4-persistent "
        "2022 conflicts reappear exactly or near the 0.10 bound in the two better monitored years, 2022 and 2023, "
        "while the fixed 2022 inventory is paired with monitored COD loads that vary 3.6-fold across 2020–2023. "
        "The discrepancy factors therefore absorb interannual hydrology and source changes. This exercise characterizes "
        "conditional consistency without confirming source coefficients; the 2020–2021 archives are too sparse for replication analysis.",
        body=True,
    )
    replace_prefix(doc, "MC-SIRC now supports two", find_prefix(doc, "MC-SIRC now supports two").text.replace("MC-SIRC now supports", "MC-SIRC supports"), body=True)
    replace_prefix(
        doc,
        "MC-SIRC reconciles a detailed watershed source inventory",
        "MC-SIRC reconciles a detailed watershed source inventory with outlet monitoring while making the limits of "
        "that reconciliation explicit. In the Nanchuan case, default month-specific missing-data scaling gives outlet "
        "loads of 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. Bayesian regularization can allocate the "
        "aggregate discrepancy and flag persistent prior conflicts, but the annual Jacobian rank of 1 and nullity of "
        "8–9 show that individual factors are not identifiable. Four boundary conflicts persist across S1–S4 and "
        "warrant field verification; they do not establish causal coefficient overestimation. Their cross-year recurrence "
        "is conditional on a fixed 2022 inventory and does not constitute validation, especially for the sparse 2020–2021 archives.",
        body=True,
    )
    replace_prefix(
        doc,
        "The spatial audit reaches the same boundary",
        "The spatial audit reaches the same boundary: after profiling γ, a k = 0 flow-only model is effectively "
        "indistinguishable from the best decay profile (ΔR² ≤ 0.0033). The profile therefore leaves half-life distances, "
        "effective-contribution shares, ranking inversions, and direct policy prescriptions unresolved. The principal "
        "regional insight is that a flashy, sediment-active Loess Plateau catchment with flood-season observation gaps "
        "cannot support source-specific transport inference from one outlet series alone. MC-SIRC converts the resulting "
        "inventory–monitoring disagreement into transparent aggregate diagnostics and a monitoring roadmap; tributary "
        "sections, seasonally resolved source inputs, and process-specific nitrogen and phosphorus observations are "
        "prerequisites for source-level calibration and transfer beyond this basin.",
        body=True,
    )
    replace_prefix(
        doc,
        "Figure 1.",
        "Figure 1. (a) Nanchuan River Basin, georeferenced inventory sources, and the outlet monitoring station. "
        "(b) Identifiability-aware MC-SIRC workflow and outputs supported by the present single-station design.",
    )
    replace_prefix(
        doc,
        "Figure 4.",
        "Figure 4. Spatial identifiability profiles with γ solved analytically at each k. Curves show relative-error "
        "SSE divided by its profile minimum; dotted lines mark the profile minimum, not an independently identified "
        "attenuation rate. The negligible ΔR² relative to k = 0 shows that monthly outlet data add virtually no distance information under shared or assumed source calendars.",
    )

    references = find_prefix(doc, "References")
    reference_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is references._p)
    for paragraph in doc.paragraphs[:reference_index]:
        if "modelling" in paragraph.text:
            paragraph.text = paragraph.text.replace("modelling", "modeling")
            if paragraph.style.name == "Normal":
                submission.set_body_style(paragraph)

    reference_updates = {
        "Ongley, E.D.": "Ongley, E.D., Zhang, X., Yu, T., 2010. Current status of agricultural and rural non-point source pollution assessment in China. Environmental Pollution 158, 1159–1168. https://doi.org/10.1016/j.envpol.2009.10.047.",
        "Qian, S.S.": "Qian, S.S., Stow, C.A., Borsuk, M.E., 2003. On Monte Carlo methods for Bayesian inference. Ecological Modelling 159, 269–277. https://doi.org/10.1016/S0304-3800(02)00299-5.",
        "Strokal, M.": "Strokal, M., Kroeze, C., Wang, M., Bai, Z., Ma, L., 2016. The MARINA model (Model to Assess River Inputs of Nutrients to seAs): Model description and results for China. Science of the Total Environment 562, 869–888. https://doi.org/10.1016/j.scitotenv.2016.04.071.",
        "Wang, Y.J.": "Wang, Y.J., Xue, M., Luo, J.H., et al., 2026. A bottom-up high-resolution water pollutant emission inventory accounting method and spatial analysis. Journal of East China Normal University (Natural Science) 2026 (1), 132–139. https://doi.org/10.3969/j.issn.1000-5641.2026.01.012.",
    }
    for prefix, text in reference_updates.items():
        paragraph = replace_prefix(doc, prefix, text)
        paragraph.paragraph_format.left_indent = base.Cm(0.75)
        paragraph.paragraph_format.first_line_indent = base.Cm(-0.75)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(3)

    replace_prefix(
        doc,
        "The authoritative analysis package",
        f"The authoritative analysis package is the tagged public repository snapshot at "
        f"https://github.com/Shifa-Zhong/MC-SIRC/tree/{TAG}. It contains the core analysis and reporting "
        "scripts supporting this study, parameter files, fixed seeds, executable reproduction commands, "
        "MCMC diagnostics, and station-anonymized input schemas. Raw hourly water-quality data are controlled "
        "by the local environmental authority and may be requested from the corresponding author for academic use, subject to approval.",
        body=True,
    )
    for table in doc.tables:
        submission.apply_three_line_table(table)
    base.set_document_format(doc, line_numbers=True)
    doc.save(path)
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, ARCHIVE / path.name)


def build_standalone_abstract():
    path = PAPER / "Abstract.docx"
    doc = Document(path)
    submission.clear_document_body(doc)
    base.set_document_format(doc, line_numbers=False)
    title = doc.add_paragraph(base.TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    heading = doc.add_paragraph("Abstract")
    heading.runs[0].bold = True
    for label, text in ABSTRACT_PARTS:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.add_run(label + ": ").bold = True
        paragraph.add_run(text)
    doc.add_paragraph("Keywords: source inventory; outlet load; structural identifiability; Bayesian regularization; "
                      "missing-not-at-random data; null model; Loess Plateau")
    doc.save(path)
    shutil.copy2(path, ARCHIVE / path.name)


def rebuild_mcmc_tables(doc):
    metadata = pd.read_excel(MCMC_DIAGNOSTICS, sheet_name="Metadata")
    for offset, pollutant in enumerate(["COD", "氨氮", "总氮", "总磷"]):
        key = {"COD": "COD", "氨氮": "NH3N", "总氮": "TN", "总磷": "TP"}[pollutant]
        frame = pd.read_excel(MCMC_DIAGNOSTICS, sheet_name=f"MCMC_{key}")
        rows = []
        for _, row in frame.iterrows():
            rows.append([
                SOURCE_MAP.get(row.Parameter, row.Parameter),
                f"{row.Mean:.3f}", f"{row.Median:.3f}", f"{row.Std:.3f}",
                f"{row['CrI_2.5']:.3f}", f"{row['CrI_97.5']:.3f}",
                f"{row.Split_Rhat:.3f}", f"{row.ESS:.0f}",
            ])
        base.replace_table(
            doc, 12 + offset,
            ["Parameter", "Mean", "Median", "SD", "95% CrI low", "95% CrI high", "Split R-hat", "ESS"],
            rows, font_size=6.9,
        )
        item = metadata[metadata.Pollutant == pollutant].iloc[0]
        prefix = f"Table S{13 + offset}."
        replace_prefix(
            doc,
            prefix,
            f"{prefix} MCMC posterior summaries—{POLLUTANT_MAP[pollutant]}, annual-mean sensitivity scenario S3. "
            f"Mean acceptance fraction = {item.Mean_acceptance_fraction:.3f}; maximum split R-hat = "
            f"{item.Max_split_Rhat:.3f}; minimum effective sample size = {item.Min_ESS:.0f}. "
            "Boundary-piled intervals describe the stated prior-regularized model and do not establish data identifiability.",
        )
    for offset, pollutant in enumerate(["COD", "氨氮", "总氮", "总磷"]):
        table = doc.tables[16 + offset]
        for row in table.rows[1:]:
            for cell in row.cells[1:]:
                try:
                    value = float(cell.text)
                except ValueError:
                    continue
                cell.text = "0.00" if abs(value) < 0.005 else f"{value:.2f}"


def remove_empty_spacers(doc):
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        if paragraph._p.xpath(".//w:drawing | .//w:pict | .//w:br | .//w:sectPr"):
            continue
        paragraph._p.getparent().remove(paragraph._p)
        removed += 1
    return removed


def postprocess_si():
    path = PAPER / "SI_Final.docx"
    doc = Document(path)
    set_title_page(doc)

    table = doc.tables[0]
    table.cell(1, 1).text = "5 area-source inventory categories + 4 georeferenced facility/farm categories"
    table.cell(1, 2).text = "1-km grid framework / individual coordinates"
    table.cell(1, 3).text = "1,588 unique grid IDs; 53,155 polygon-intersection records; 100 georeferenced records"
    table.cell(3, 1).text = "Hourly rainfall at one basin gauge"
    table.cell(3, 3).text = "17,520 records (2022–2023; 8,760 per year)"

    emissions = doc.tables[7]
    emissions.cell(0, 2).text = "NH₃-N-E"
    emissions.cell(0, 6).text = "NH₃-N-R"
    emissions.cell(1, 3).text = "—"

    replace_prefix(doc, "Table S1.", "Table S1. Overview of research data. The 53,155 value is the number of polygon-intersection records assigned to 1,588 unique 1-km grid IDs, not the number of independent 1-km cells.")
    replace_prefix(doc, "Table S5.", "Table S5. Descriptive urban NPS functional-zone weights (not calibrated). Percentages may not sum to exactly 100% because of rounding.")
    replace_prefix(doc, "Table S8.", "Table S8. Source emissions (E, kg) and nominal river-entry loads (R, kg). “—” denotes unavailable inventory data. Industrial TN is unavailable, so TN totals sum reported categories only. The 56.6% large-livestock COD share is recalculated directly from this archived case-study workbook (586,721/1,036,524), rather than copied from the summary percentage in Wang et al. (2026).")
    replace_prefix(doc, "Table S9.", "Table S9. Observed-hour monthly outlet loads (S1) and timestamp counts in 2022.")
    replace_prefix(doc, "Table S10.", "Table S10. Observed-hour (S1) rainfall/non-rainfall load summary for 2022. Rainfall hours have recorded precipitation >0 mm; intensity ratio is mean hourly load during rainfall hours divided by mean hourly load during non-rainfall hours.")
    replace_prefix(doc, "Table S11.", "Table S11. Inventory and outlet discrepancy ratios under default S2. Outlet/Entry is a scale ratio and is not interpreted as a channel-transport coefficient.")

    replace_prefix(
        doc,
        "The 2022 archive contains",
        "The 2022 archive contains 5,928 unique hourly timestamps from a possible 8,760 (67.7% coverage), so "
        "2,832 timestamps are absent. Gaps cluster in July–November and are treated as MNAR. Within observed "
        "timestamps, concentration gaps were linearly interpolated; discharge gaps or invalid values were filled by "
        "same-calendar historical medians, adjacent-month medians, and then interpolation. Completely absent hours "
        "were handled by four scenarios: S1 observed-only; S2 month-specific coverage scaling (default); S3 annual-mean "
        "scaling by 8,760/5,928 = 1.4777; and S4 observed load plus 1.5 times the S2-imputed portion. S4 is an "
        "enhanced-gap stress test that adds 50% to the imputed portion; 1.5 is a scenario-design choice, not an "
        "estimated event multiplier or confidence bound. S2 is preferred because it preserves observed monthly structure.",
    )
    replace_prefix(
        doc,
        "The annual-mean reconciliation was repeated",
        "The annual-mean reconciliation was repeated for 2020–2023 with the 2022 inventory fixed. This design "
        "characterizes conditional consistency because factors absorb interannual hydrology and source changes. "
        "Archive timestamp coverage, used for annual scaling in Table S28, differs from strict raw pollutant–flow pair "
        "coverage before within-timestamp filling. The latter is only 1.0–1.3% in 2020–2021; the two measures therefore "
        "have different denominators and should not be compared as if they were the same completeness statistic.",
    )
    replace_prefix(
        doc,
        "The four persistent boundary conflicts remain near",
        "The four persistent boundary conflicts remain near 0.100 across the tested observation-error scales. This "
        "sensitivity analysis characterizes the response to the assumed likelihood scale; formal source estimability is assessed by rank and null-space diagnostics.",
    )
    replace_prefix(
        doc,
        "“—” denotes unavailable inventory data",
        "“—” denotes unavailable inventory data. Industrial TN has neither an emission entry nor a nominal river-entry "
        "coefficient; it is not encoded as zero, and reported TN totals sum only available categories.",
        optional=True,
    )

    rebuild_mcmc_tables(doc)
    prior_definition = (
        " P1 sets all active-source prior means to 0.5 with default standard deviations; P2 uses the source-specific "
        "defaults in Table S4; P3 sets all means to 1.2 with default standard deviations; P4 uses default means with "
        "σ = 0.5; and P5 uses default means with σ = 1.0. Only sources with positive nominal entry loads are shown."
    )
    for number, pollutant in zip(range(18, 22), ["COD", "NH₃-N", "TN", "TP"]):
        replace_prefix(doc, f"Table S{number}.", f"Table S{number}. Prior sensitivity—{pollutant}, annual-mean sensitivity scenario S3.{prior_definition}")

    elasticity = doc.tables[27]
    rows = [[row.cells[0].text, row.cells[1].text, row.cells[2].text, row.cells[4].text, row.cells[5].text]
            for row in elasticity.rows[1:]]
    base.replace_table(doc, 27, ["Pollutant", "Most sensitive input", "Elasticity", "Second most sensitive input", "Elasticity"], rows, font_size=7.8)
    replace_prefix(doc, "Table S25.", "Table S25. Elasticity of nominal river-entry load to inventory inputs. Elasticity is an inventory-side variance sensitivity and is not an outlet contribution share.")

    crossyear = doc.tables[32]
    crossyear.cell(0, 5).text = "Years exactly at 0.100"
    crossyear.cell(3, 5).text = "1 of 4 (2023 = 0.11)"
    replace_prefix(doc, "Table S30.", "Table S30. Cross-year behavior of the four source–pollutant pairs that are persistent S1–S4 boundary conflicts in 2022; 2020–2021 are data-poor.")
    replace_prefix(
        doc,
        "The 0.100 boundary recurs",
        "All four pairs are at or near the 0.100 bound in 2022. Three are exactly at the bound in 2023; industrial TP is near it at 0.11. This conditional pattern supports field rechecking but does not identify true source coefficients.",
        optional=True,
    )

    labels = {
        "(COD)": "Table S17 (COD panel).",
        "(NH₃-N)": "Table S17 (continued—NH₃-N panel).",
        "(TN)": "Table S17 (continued—TN panel).",
        "(TP)": "Table S17 (continued—TP panel).",
    }
    for prefix, text in labels.items():
        paragraph = replace_prefix(doc, prefix, text, optional=True)
        if paragraph is not None:
            paragraph.runs[0].bold = True
            paragraph.paragraph_format.keep_with_next = True

    image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//a:blip")]
    if image_paragraphs:
        clear_paragraph(image_paragraphs[0])
        image_paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraphs[0].add_run().add_picture(str(FIG / "figure4_spatial_identifiability_profile.png"), width=Inches(6.5))
    if len(image_paragraphs) > 1:
        clear_paragraph(image_paragraphs[1])
        image_paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraphs[1].add_run().add_picture(str(FIG_S2 / "figureS2_monthly_loads.png"), width=Inches(6.5))
    replace_prefix(doc, "Figure S1.", "Figure S1. Spatial-profile identifiability audit with γ solved analytically at each k. Dotted lines mark the profile minimum, not an independently identified attenuation rate; negligible improvement relative to k = 0 leaves k unresolved under shared or assumed source calendars.")
    replace_prefix(doc, "Figure S2.", "Figure S2. Observed-hour (S1) monthly outlet loads and timestamp coverage in 2022. Gray months have <50% coverage and are excluded from the spatial-profile objective; all months remain represented in the S1–S4 annual-load scenarios.")

    text_s8 = find_prefix(doc, "Text S8.")
    text_s8.style = "Heading 2"
    text_s8.paragraph_format.keep_with_next = True
    description_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is text_s8._p) + 1
    if description_index < len(doc.paragraphs):
        description = doc.paragraphs[description_index]
        description.text = (
            f"Tagged repository snapshot: https://github.com/Shifa-Zhong/MC-SIRC/tree/{TAG}. The repository contains "
            "the core analysis and reporting scripts supporting this study, config/revision3_parameters.json, fixed "
            "seeds, executable commands, MCMC convergence diagnostics, and station-anonymized example schemas. "
            "Restricted real-data inputs follow data/raw/{monitor.xlsx, rain.xlsx, data(1).xlsx}; generated files are "
            "placed under output/results and output/figures. MAP, MCMC, Monte Carlo, and optimization seeds are 42."
        )
        description.paragraph_format.line_spacing = 1.0

    remove_empty_spacers(doc)
    for table in doc.tables:
        base.format_table(table, font_size=7.0 if len(table.columns) >= 8 else 8.0)
        submission.apply_three_line_table(table)
        if table.rows:
            base.set_repeat_table_header(table.rows[0])
        for row in table.rows:
            base.keep_row_together(row)
    base.set_document_format(doc, line_numbers=False)
    doc.save(path)
    shutil.copy2(path, ARCHIVE / path.name)


def postprocess_cover_and_ga():
    cover = PAPER / "cover letter.docx"
    doc = Document(cover)
    for paragraph in doc.paragraphs:
        normalized = re.sub(r"revision-2026-08-(?:24|25(?:\.1)*)", TAG, paragraph.text)
        if normalized != paragraph.text:
            paragraph.text = normalized
        paragraph.text = paragraph.text.replace(
            "Updated Highlights, Graphical Abstract, Supporting Information and response letter accompany the manuscript;",
            "Updated Highlights, a code-generated Graphical Abstract, Supporting Information, the study-area KML/KMZ, and the response letter accompany the manuscript;",
        )
    doc.save(cover)
    shutil.copy2(cover, ARCHIVE / cover.name)
    ga = PAPER / "GA.docx"
    doc = Document(ga)
    submission.clear_document_body(doc)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(GA_IMAGE), width=Inches(10.8))
    for section in doc.sections:
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
    doc.save(ga)
    shutil.copy2(ga, ARCHIVE / ga.name)


def exact_paragraph(doc, prefix):
    return find_prefix(doc, prefix).text.strip()


def build_exact_response():
    response_v4.update_response_data()
    ms = Document(PAPER / "Manuscript_Final_MC-SIRC.docx")
    si = Document(PAPER / "SI_Final.docx")
    responses = []
    for response, _ in response_base.RESPONSES:
        response = response.replace("revision-2026-08-24", TAG)
        response = response.replace("and is within the 225-word limit (192 words)", "and is within the 225-word limit")
        response = response.replace("The research gap and four objectives are now stated explicitly.", "The research gap and five methodological contributions are now stated explicitly.")
        responses.append([response, []])
    responses[9][0] = (
        f"We uploaded the complete core analysis and reporting code to the public repository and created the stable "
        f"tag {TAG}. The tagged snapshot contains the parameter files, fixed seeds, executable commands, MCMC "
        "diagnostics, and station-anonymized input schemas; no separate code archive is required."
    )
    responses[11][0] += (
        " A final audit also corrected the inventory-resolution terminology: 53,155 denotes polygon-intersection "
        "records assigned to 1,588 unique 1-km grid IDs, not 53,155 independent grid cells."
    )
    responses[28][0] = (
        "We bolded the Table S1 title, added explicit continued-panel labels to Table S17, retained repeated header "
        "rows and non-splitting rows, and set compact headers not to break within words. The main text cites all "
        "supplementary texts, figures, and tables in ascending ranges before item-specific references."
    )
    responses[29][0] = (
        "The revised SI no longer uses dagger symbols to denote parameter bounds; boundary status is stated "
        "explicitly in captions and table columns. Undefined dagger marks were also removed from the author line. "
        "This eliminates the ambiguous symbol-spacing problem rather than replacing it with another marker."
    )
    responses[31][0] = (
        "We updated the submission files to the Journal of Hydrology: Regional Studies format. The abstract has "
        "three separate title-case labelled paragraphs within the 225-word limit; citations use author–year style; "
        "the reference list is alphabetical; the manuscript uses double spacing and continuous line numbering; "
        "and a study-area KML/KMZ is included. Figure and table captions and the standalone Abstract were updated consistently."
    )

    ms_abs = exact_paragraph(ms, "Study Focus:")
    ms_rank = exact_paragraph(ms, "For Eq. (2), the Jacobian")
    ms_tier = exact_paragraph(ms, "The standardized prior shift")
    ms_missing = exact_paragraph(ms, "Within observed timestamps")
    ms_cross = exact_paragraph(ms, "The same annual-mean protocol")
    ms_spatial = exact_paragraph(ms, "Profiling γ analytically")
    ms_withdraw = exact_paragraph(ms, "Because k is not identified")
    ms_management = exact_paragraph(ms, "Operational measures should therefore")
    ms_transfer = exact_paragraph(ms, "The primary limitation is structural")
    ms_inventory = exact_paragraph(ms, "The bottom-up 2022 inventory uses")
    ms_results_inventory = exact_paragraph(ms, "The inventory contains")
    ms_mcmc = exact_paragraph(ms, "The large-livestock TP MCMC interval")
    ms_residual = exact_paragraph(ms, "The prior-regularized reconciliation remains close")
    ms_forward = exact_paragraph(ms, "Forward Monte Carlo distributions")
    ms_equation = exact_paragraph(ms, "where wᵢ,ₘ")
    ms_process = exact_paragraph(ms, "The absence of an identifiable k")
    ms_conclusion1 = exact_paragraph(ms, "MC-SIRC reconciles a detailed")
    ms_conclusion2 = exact_paragraph(ms, "The spatial audit reaches the same boundary")
    ms_nav = exact_paragraph(ms, "Each MC-SIRC component")
    ms_data = exact_paragraph(ms, "The authoritative analysis package")
    si_text1 = exact_paragraph(si, "The 2022 archive contains")
    si_cross = exact_paragraph(si, "The annual-mean reconciliation was repeated")
    si_urban = exact_paragraph(si, "Urban NPS functional-zone weights")
    si_abbr = exact_paragraph(si, "Abbreviations:")

    def item(label, text):
        return (label, text)

    excerpts = [
        [item("Revised MS (Abstract):", ms_abs)],
        [item("Revised MS (§2.5):", ms_tier)],
        [item("Revised MS (§2.6):", ms_rank), item("Revised MS (§2.5):", ms_tier)],
        [item("Revised MS (§2.2):", ms_missing), item("Revised SI (Text S1):", si_text1)],
        [item("Revised MS (§2.6):", ms_cross), item("Revised MS (Conclusion):", ms_conclusion1)],
        [item("Revised MS (§3.3):", ms_spatial), item("Revised MS (§4.1):", ms_withdraw)],
        [item("Revised MS (§4.2):", ms_management)],
        [item("Revised MS (§4.3):", ms_transfer)],
        [item("Revised MS (Table 1 caption):", exact_paragraph(ms, "Table 1."))],
        [item("Revised MS (Data availability):", ms_data), item("Revised SI (Text S8):", exact_paragraph(si, f"Tagged repository snapshot:"))],
        [item("Revised MS (Conclusion):", ms_conclusion2)],
        [item("Revised MS (§2.2):", ms_inventory), item("Revised MS (§3.1):", ms_results_inventory)],
        [item("Revised MS (§2.8):", ms_equation), item("Revised MS (§3.3):", ms_spatial)],
        [item("Revised MS (§3.3):", ms_spatial)],
        [item("Revised MS (§4.2):", exact_paragraph(ms, "MC-SIRC supports two"))],
        [item("Revised MS (§3.2):", ms_mcmc)],
        [item("Revised MS (§2.5–2.6):", ms_tier + " " + ms_rank), item("Revised SI (Table S33 caption):", exact_paragraph(si, "Table S33."))],
        [item("Revised MS (§3.1):", ms_residual)],
        [item("Revised MS (§2.2):", ms_missing), item("Revised SI (Text S1):", si_text1)],
        [item("Revised MS (§2.6):", ms_cross), item("Revised MS (Conclusion):", ms_conclusion1)],
        [item("Revised MS (§3.2):", ms_forward)],
        [item("Revised SI (Text S2):", si_urban)],
        [item("Revised MS (§4.1–4.2):", ms_withdraw + " " + ms_management)],
        [item("Revised MS (Abstract):", ms_abs)],
        [item("Revised MS (Introduction):", exact_paragraph(ms, "Identifying and quantifying pollution sources")), item("Revised MS (contributions):", exact_paragraph(ms, "MC-SIRC makes five"))],
        [item("Revised MS (Figure 2 caption):", exact_paragraph(ms, "Figure 2.")), item("Revised MS (Figure 3 caption):", exact_paragraph(ms, "Figure 3."))],
        [item("Revised MS (§4.1):", ms_process)],
        [item("Revised MS (Conclusion):", ms_conclusion2)],
        [item("Revised MS (§2.3):", ms_nav), item("Revised SI (Table S1 caption):", exact_paragraph(si, "Table S1.")), item("Revised SI (Table S17 caption):", exact_paragraph(si, "Table S17."))],
        [item("Revised SI (boundary presentation):", exact_paragraph(si, "Table S6.") + " " + exact_paragraph(si, "Table S33."))],
        [item("Revised SI (front matter):", si_abbr)],
        [item("Revised MS (Abstract headings and text):", "\n".join(exact_paragraph(ms, label + ":") for label, _ in ABSTRACT_PARTS))],
    ]
    if len(excerpts) != 32:
        raise RuntimeError("Response excerpt map must contain 32 replies")
    response_base.RESPONSES = [(responses[index][0], excerpts[index]) for index in range(32)]
    response_base.main()

    path = PAPER / "Response_Letter.docx"
    response = Document(path)
    for paragraph in response.paragraphs:
        if paragraph.text.startswith("We thank you for the careful"):
            paragraph.text = (
                "We thank the Editor and Reviewers for the detailed comments. The revision adds formal identifiability "
                "and k = 0 diagnostics, corrects and propagates the missing-data and inventory-resolution descriptions, "
                "limits source-specific interpretations to what one outlet station can support, and updates the manuscript, "
                "SI, figures, tables, repository, and submission files consistently. Exact final MS/SI excerpts are shown in blue below each reply."
            )
            break
    augment.add_materials(response, ms, si)
    replies = [p for p in response.paragraphs if p.text.strip().startswith("Reply:")]

    def add_table(reply_index, label, table_index, caption_prefix):
        augment.table_block(replies[reply_index]._p, response, label, si.tables[table_index], exact_paragraph(si, caption_prefix))

    add_table(3, "Revised SI Table S9 (full table):", 8, "Table S9.")
    add_table(3, "Revised SI Table S10 (full table):", 9, "Table S10.")
    add_table(4, "Revised SI Table S30 (full table):", 32, "Table S30.")
    add_table(11, "Revised SI Table S8 (full table):", 7, "Table S8.")
    add_table(11, "Revised SI Table S11 (full table):", 10, "Table S11.")
    for table_index, number in zip(range(12, 16), range(13, 17)):
        add_table(15, f"Revised SI Table S{number} (full table):", table_index, f"Table S{number}.")
    for table_index, label in zip(range(16, 20), ["COD", "NH₃-N", "TN", "TP"]):
        add_table(15, f"Revised SI Table S17 ({label} panel):", table_index, "Table S17.")
    for table_index, number in zip(range(20, 24), range(18, 22)):
        add_table(15, f"Revised SI Table S{number} (full table):", table_index, f"Table S{number}.")
    add_table(20, "Revised SI Table S25 (full table):", 27, "Table S25.")
    for table_index, label in zip(range(17, 20), ["NH₃-N", "TN", "TP"]):
        add_table(28, f"Revised SI Table S17 ({label} continued panel):", table_index, "Table S17.")
    response_v4.base.style_document(response)
    response.save(path)
    shutil.copy2(path, ARCHIVE / path.name)


def prune_docx_media(path: Path):
    """Remove orphan image relationships and media parts without changing document content."""
    with zipfile.ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}
    updated_rels = {}
    referenced_media = set()
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    for rel_name, data in list(entries.items()):
        if not rel_name.endswith(".rels") or "/_rels/" not in rel_name:
            continue
        rel_path = PurePosixPath(rel_name)
        source_part = rel_path.parent.parent / rel_path.stem
        source_xml = entries.get(source_part.as_posix(), b"")
        tree = ET.fromstring(data)
        changed = False
        for relationship in list(tree):
            rel_type = relationship.attrib.get("Type", "")
            rel_id = relationship.attrib.get("Id", "")
            target = relationship.attrib.get("Target", "")
            if not rel_type.endswith("/image"):
                continue
            if rel_id.encode("utf-8") not in source_xml:
                tree.remove(relationship)
                changed = True
                continue
            resolved = (source_part.parent / target).as_posix()
            resolved = str(PurePosixPath(resolved))
            while "/../" in resolved:
                resolved = re.sub(r"[^/]+/\.\./", "", resolved, count=1)
            referenced_media.add(resolved)
        if changed:
            updated_rels[rel_name] = ET.tostring(tree, encoding="utf-8", xml_declaration=True)
    entries.update(updated_rels)
    for name in list(entries):
        if name.startswith("word/media/") and name not in referenced_media:
            del entries[name]
    temp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED) as destination:
        for name, data in entries.items():
            destination.writestr(name, data)
    os.replace(temp, path)
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            raise RuntimeError(f"Corrupt DOCX after media pruning: {path}")


def run_external_steps():
    subprocess.run([sys.executable, str(ROOT / "scripts" / "analysis" / "revision5_mcmc_diagnostics.py")], check=True)
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "generate_study_area_kml.py")], check=True)
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "generate_word_level_highlights_v2.py")], check=True)


def archive_working_files():
    destination = PAPER / "_archive_non_submission_20260824" / "working_files_20260825"
    destination.mkdir(parents=True, exist_ok=True)
    names = [
        "facts_library.md", "manuscript_text.txt", "response_letter_outline.md", "REVISIONS.md", "si_text.txt",
        "Submission_Consistency_Audit_20260824.md", "Highlight_Validation_20260824.md",
        "New Microsoft PowerPoint Presentation.pptx",
    ]
    for name in names:
        source = PAPER / name
        if source.exists():
            shutil.move(str(source), str(destination / name))


def main():
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    generate_all_figures()
    submission.main()
    subprocess.run([sys.executable, str(ROOT / "scripts" / "analysis" / "revision5_mcmc_diagnostics.py")], check=True)
    postprocess_manuscript()
    build_standalone_abstract()
    postprocess_si()
    postprocess_cover_and_ga()
    build_exact_response()
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "generate_study_area_kml.py")], check=True)
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "generate_word_level_highlights_v2.py")], check=True)
    for name in [
        "Manuscript_Final_MC-SIRC.docx", "SI_Final.docx", "Abstract.docx", "Response_Letter.docx",
        "Manuscript_Highlighted_20260824.docx", "SI_Highlighted_20260824.docx",
        "Abstract_Highlighted_20260824.docx", "GA.docx", "cover letter.docx", "highlights.docx",
        "declarationStatement.docx",
    ]:
        path = PAPER / name
        if path.exists():
            prune_docx_media(path)
    archive_working_files()
    print("Finalized submission package:", PAPER)


if __name__ == "__main__":
    main()
