#!/usr/bin/env python3
'''Rebuild the response letter and insert revised figures/tables under relevant replies.'''

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.reporting import build_revision4_response_letter as response_builder


PAPER = ROOT / 'paper'
ARCHIVE_DUP = PAPER / '_archive_non_submission_20260824' / 'duplicate_dated_versions'
FIG = ROOT / 'output' / 'figures'
BLUE = RGBColor(0x00, 0x70, 0xC0)


def paragraph_after(cursor, document, text='', blue=False, bold=False, size=9):
    element = OxmlElement('w:p')
    cursor.addnext(element)
    paragraph = Paragraph(element, document._body)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(size)
        if blue:
            run.font.color.rgb = BLUE
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def figure_block(cursor, document, label, image_path, caption):
    heading = paragraph_after(cursor, document, label, blue=True, bold=True, size=10)
    picture = paragraph_after(heading._p, document)
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    caption_paragraph = paragraph_after(picture._p, document, caption, blue=True, size=8.5)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return caption_paragraph._p


def style_blue_table(table, font_size):
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn('w:cantSplit')) is None:
            tr_pr.append(OxmlElement('w:cantSplit'))
        if row_index == 0 and tr_pr.find(qn('w:tblHeader')) is None:
            tr_pr.append(OxmlElement('w:tblHeader'))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = BLUE


def table_block(cursor, document, label, source_table, caption, row_limit=None):
    heading = paragraph_after(cursor, document, label, blue=True, bold=True, size=10)
    table_xml = deepcopy(source_table._tbl)
    if row_limit is not None:
        for row in list(table_xml.tr_lst)[row_limit:]:
            table_xml.remove(row)
    heading._p.addnext(table_xml)
    table = Table(table_xml, document._body)
    font_size = 5.8 if len(table.columns) >= 9 else 6.5 if len(table.columns) >= 7 else 7.2
    style_blue_table(table, font_size)
    caption_paragraph = paragraph_after(table_xml, document, caption, blue=True, size=8.5)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return caption_paragraph._p


def caption(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph.text.strip()
    raise KeyError(prefix)


def add_materials(response, ms, si):
    replies = [p for p in response.paragraphs if p.text.strip().startswith('Reply:')]
    if len(replies) != 32:
        raise RuntimeError(f'Expected 32 replies, found {len(replies)}')

    intro = next(p for p in response.paragraphs if p.text.startswith('We thank the Editor and Reviewers'))
    paragraph_after(
        intro._p,
        response,
        'For rapid review, all revised main-text figures and tables, together with the SI figures/tables directly tied to reviewer comments, are reproduced below the corresponding replies. Figure/table labels and table text are shown in blue.',
        blue=False,
        size=10,
    )

    cursor = replies[0]._p
    figure_block(cursor, response, 'Revised MS Figure 1 (full figure):', FIG / 'revision3' / 'figure1_revised_framework.png', caption(ms, 'Figure 1.'))

    cursor = replies[3]._p
    cursor = figure_block(cursor, response, 'Revised SI Figure S2 (full figure):', FIG / 'figureS2_monthly_loads' / 'figureS2_monthly_loads.png', caption(si, 'Figure S2.'))
    cursor = table_block(cursor, response, 'Revised SI Table S3 (full table):', si.tables[2], caption(si, 'Table S3.'))
    cursor = table_block(cursor, response, 'Revised SI Table S26 (full table):', si.tables[28], caption(si, 'Table S26.'))
    table_block(cursor, response, 'Revised SI Table S34 (full table):', si.tables[36], caption(si, 'Table S34.'))

    table_block(replies[4]._p, response, 'Revised SI Table S28 (full table):', si.tables[30], caption(si, 'Table S28.'))

    table_block(replies[8]._p, response, 'Revised MS Table 1 (full table):', ms.tables[0], caption(ms, 'Table 1.'))

    cursor = replies[13]._p
    cursor = figure_block(
        cursor,
        response,
        'Revised MS Figure 4 / SI Figure S1 (shared full diagnostic figure):',
        FIG / 'revision3' / 'figure4_spatial_identifiability_profile.png',
        caption(ms, 'Figure 4.') + ' ' + caption(si, 'Figure S1.'),
    )
    table_block(cursor, response, 'Revised MS Table 3 (full table):', ms.tables[2], caption(ms, 'Table 3.'))

    table_block(replies[15]._p, response, 'Revised SI Table S16 (full table):', si.tables[15], caption(si, 'Table S16.'))

    cursor = replies[16]._p
    cursor = table_block(cursor, response, 'Revised SI Table S32 (full table):', si.tables[34], caption(si, 'Table S32.'))
    table_block(cursor, response, 'Revised SI Table S33 (full table):', si.tables[35], caption(si, 'Table S33.'))

    table_block(replies[17]._p, response, 'Revised MS Table 2 (full table):', ms.tables[1], caption(ms, 'Table 2.'))
    table_block(replies[20]._p, response, 'Revised SI Table S24 (full table):', si.tables[26], caption(si, 'Table S24.'))
    table_block(replies[21]._p, response, 'Revised SI Table S5 (full table):', si.tables[4], caption(si, 'Table S5.'))

    cursor = replies[25]._p
    cursor = figure_block(cursor, response, 'Revised MS Figure 2 (full figure):', FIG / 'revision3' / 'figure2_identifiability_aware_bayesian.png', caption(ms, 'Figure 2.'))
    figure_block(cursor, response, 'Revised MS Figure 3 (full figure):', FIG / 'revision3' / 'figure3_forward_uncertainty.png', caption(ms, 'Figure 3.'))

    cursor = replies[28]._p
    cursor = table_block(cursor, response, 'Revised SI Table S1 (full table):', si.tables[0], caption(si, 'Table S1.'))
    table_block(
        cursor,
        response,
        'Revised SI Table S17 (COD panel shown; all four panels use the same layout):',
        si.tables[16],
        caption(si, 'Table S17.') + ' Representative COD panel; the NH₃-N, TN, and TP panels are formatted identically.',
    )

    cursor = replies[29]._p
    cursor = table_block(cursor, response, 'Revised SI Table S6 (full table):', si.tables[5], caption(si, 'Table S6.'))
    table_block(cursor, response, 'Revised SI Table S33 (boundary-marker formatting shown in full table):', si.tables[35], caption(si, 'Table S33.'))


def main():
    response_builder.main()
    response_path = PAPER / 'Response_Letter.docx'
    response = Document(response_path)
    ms = Document(PAPER / 'Manuscript_Final_MC-SIRC.docx')
    si = Document(PAPER / 'SI_Final.docx')
    add_materials(response, ms, si)
    response_builder.base.style_document(response)
    response.save(response_path)
    shutil.copy2(response_path, ARCHIVE_DUP / 'Response_Letter_Revised_20260824.docx')
    print(response_path)


if __name__ == '__main__':
    main()
