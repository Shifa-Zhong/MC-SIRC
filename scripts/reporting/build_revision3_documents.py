#!/usr/bin/env python3
"""Build the revised manuscript, supporting information, and abstract files."""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
ARCHIVE_DUP = PAPER / "_archive_non_submission_20260824" / "duplicate_dated_versions"
FIG = ROOT / "output" / "figures" / "revision3"
RESULT_XLSX = ROOT / "output" / "results" / "revision3_identifiability_and_missingness.xlsx"
SPATIAL_XLSX = ROOT / "output" / "results" / "spatial_identifiability_unbounded.xlsx"

TITLE = (
    "MC-SIRC: An Identifiability-Aware Workflow for Reconciling Watershed Source Inventories "
    "with Outlet Monitoring—A Case Study of the Nanchuan River Basin, Loess Plateau"
)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_nowrap(cell, enabled=True):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:noWrap"))
    if enabled and existing is None:
        tc_pr.append(OxmlElement("w:noWrap"))
    elif not enabled and existing is not None:
        tc_pr.remove(existing)


def format_table(table, font_size=8.5, header_fill="D9EAF7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
                    if r_idx == 0:
                        run.bold = True


def set_document_format(doc, line_numbers=True):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if "Caption" in doc.styles:
        doc.styles["Caption"].font.size = Pt(10)
        doc.styles["Caption"].font.italic = False
        doc.styles["Caption"].paragraph_format.line_spacing = 1.0
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        if line_numbers:
            sect_pr = section._sectPr
            old = sect_pr.find(qn("w:lnNumType"))
            if old is not None:
                sect_pr.remove(old)
            ln = OxmlElement("w:lnNumType")
            ln.set(qn("w:countBy"), "1")
            ln.set(qn("w:start"), "1")
            ln.set(qn("w:restart"), "continuous")
            ln.set(qn("w:distance"), "360")
            sect_pr.append(ln)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def clear_body_from_paragraph(doc, paragraph_index):
    body = doc._element.body
    marker = doc.paragraphs[paragraph_index]._p
    started = False
    for child in list(body):
        if child is marker:
            started = True
        if started and child.tag != qn("w:sectPr"):
            body.remove(child)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.italic = True
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_with_next = False
    return p


def add_picture(doc, path, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_table(doc, headers, rows, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, value in enumerate(headers):
        table.cell(0, j).text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    format_table(table, font_size=font_size)
    return table


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.addnext(new_paragraph._p)
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.getparent().replace(new_p, new_paragraph._p)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def simple_insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def replace_table(doc, table_index, headers, rows, font_size=8.0):
    old = doc.tables[table_index]
    new = doc.add_table(rows=1, cols=len(headers))
    new.style = "Table Grid"
    for j, value in enumerate(headers):
        new.cell(0, j).text = str(value)
    for row in rows:
        cells = new.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    format_table(new, font_size=font_size)
    old._tbl.addnext(new._tbl)
    old._tbl.getparent().remove(old._tbl)
    return new


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


REFERENCES = [
    "Alexander, R.B., Böhlke, J.K., Boyer, E.W., et al., 2009. Dynamic modeling of nitrogen losses in river networks unravels the coupled effects of hydrological and biogeochemical processes. Biogeochemistry 93, 91–116.",
    "Alexander, R.B., Smith, R.A., Schwarz, G.E., 2000. Effect of stream channel size on the delivery of nitrogen to the Gulf of Mexico. Nature 403, 758–761.",
    "Arnold, J.G., Srinivasan, R., Muttiah, R.S., et al., 1998. Large area hydrologic modeling and assessment part I: Model development. Journal of the American Water Resources Association 34, 73–89.",
    "Beaulieu, J.J., Tank, J.L., Hamilton, S.K., et al., 2011. Nitrous oxide emission from denitrification in stream and river networks. Proceedings of the National Academy of Sciences 108, 214–219.",
    "Behrendt, H., Opitz, D., 2000. Retention of nutrients in river systems: Dependence on specific runoff and hydraulic load. Hydrobiologia 410, 111–122.",
    "Birgand, F., Skaggs, R.W., Chescheir, G.M., et al., 2007. Nitrogen removal in streams of agricultural catchments: A literature review. Critical Reviews in Environmental Science and Technology 37, 381–487.",
    "Borsuk, M.E., Stow, C.A., Reckhow, K.H., 2004. A Bayesian network of eutrophication models for synthesis, prediction, and uncertainty analysis. Ecological Modelling 173, 219–239.",
    "Carpenter, S.R., Caraco, N.F., Correll, D.L., et al., 1998. Nonpoint pollution of surface waters with phosphorus and nitrogen. Ecological Applications 8, 559–568.",
    "Chapra, S.C., 1997. Surface Water-Quality Modeling. McGraw-Hill, New York.",
    "Chen, L., Zhong, Y., Wei, G., Cai, Y., Shen, Z., 2014. Development of an integrated modeling approach for identifying multilevel non-point-source priority management areas at the watershed scale. Water Resources Research 50, 4095–4109.",
    "Foreman-Mackey, D., Hogg, D.W., Lang, D., et al., 2013. emcee: The MCMC Hammer. Publications of the Astronomical Society of the Pacific 125, 306–312.",
    "Gelman, A., Carlin, J.B., Stern, H.S., et al., 2013. Bayesian Data Analysis, third ed. CRC Press, Boca Raton.",
    "Grizzetti, B., Bouraoui, F., De Marsily, G., 2008. Assessing nitrogen pressures on European surface water. Global Biogeochemical Cycles 22, GB4023.",
    "Hirsch, R.M., Moyer, D.L., Archfield, S.A., 2010. Weighted regressions on time, discharge, and season, with an application to Chesapeake Bay river inputs. Journal of the American Water Resources Association 46, 857–880.",
    "Johnes, P.J., 1996. Evaluation and management of the impact of land use change on the nitrogen and phosphorus load delivered to surface waters: The export coefficient modelling approach. Journal of Hydrology 183, 323–349.",
    "Li, Q., Yang, Z., Yin, W., Yu, M., Bai, F., Yue, Y., Ren, Y., 2023. Tracing the sources and transport of the total phosphorus in the upper Yangtze River. Ecological Informatics 77, 102230. https://doi.org/10.1016/j.ecoinf.2023.102230.",
    "Liu, J., Yan, T., Bai, J., Shen, Z., 2023. Integrating source apportionment and landscape patterns to capture nutrient variability across a typical urbanized watershed. Journal of Environmental Management 325, 116559. https://doi.org/10.1016/j.jenvman.2022.116559.",
    "Liu, R., Wang, J., Shi, J., et al., 2014. Runoff characteristics and nutrient loss mechanism from plain farmland under simulated rainfall conditions. Science of the Total Environment 468–469, 1069–1077.",
    "Malve, O., Qian, S.S., 2006. Estimating nutrients and chlorophyll a relationships in Finnish lakes. Environmental Science & Technology 40, 7848–7853.",
    "Meals, D.W., Dressing, S.A., Davenport, T.E., 2010. Lag time in water quality response to best management practices: A review. Journal of Environmental Quality 39, 85–96.",
    "Montefiore, L.R., Nelson, N.G., 2022. Can a simple water quality model effectively estimate runoff-driven nutrient loads to estuarine systems? A national-scale comparison of STEPLgrid and SPARROW. Environmental Modelling & Software 150, 105344. https://doi.org/10.1016/j.envsoft.2022.105344.",
    "Neumann, A., Blukacz-Richards, E.A., Saha, R., Alberto Arnillas, C., Arhonditsis, G.B., 2023. A Bayesian hierarchical spatially explicit modelling framework to examine phosphorus export between contrasting flow regimes. Journal of Great Lakes Research 49, 190–208. https://doi.org/10.1016/j.jglr.2022.10.003.",
    "Ongley, E.D., Xiaolan, Z., Tao, Y., 2010. Current status of agricultural and rural non-point source pollution assessment in China. Environmental Pollution 158, 1159–1168.",
    "Preston, S.D., Alexander, R.B., Schwarz, G.E., et al., 2011. Factors affecting stream nutrient loads: A synthesis of regional SPARROW model results. Journal of the American Water Resources Association 47, 891–915.",
    "Qian, S.S., Stow, C.A., Borsuk, M.E., 2003. On Monte Carlo methods for Bayesian inference. Ecological Modelling 159, 267–277.",
    "Reckhow, K.H., 2003. On the need for uncertainty assessment in TMDL modeling and implementation. Journal of Water Resources Planning and Management 129, 245–246.",
    "Rode, M., Arhonditsis, G., Balin, D., et al., 2010. New challenges in integrated water quality modelling. Hydrological Processes 24, 3447–3461.",
    "Schwarz, G.E., Hoos, A.B., Alexander, R.B., et al., 2006. The SPARROW surface water-quality model: Theory, application, and user documentation. US Geological Survey Techniques and Methods, Book 6, Chapter B3.",
    "Shen, Z., Chen, L., Hong, Q., et al., 2013. Assessment of nitrogen and phosphorus loads and causal factors from different land use and soil types in the Three Gorges Reservoir Area. Science of the Total Environment 454–455, 383–392.",
    "Shen, Z., Liao, Q., Hong, Q., et al., 2012. An overview of research on agricultural non-point source pollution modelling in China. Separation and Purification Technology 84, 104–111.",
    "Smith, R.A., Schwarz, G.E., Alexander, R.B., 1997. Regional interpretation of water-quality monitoring data. Water Resources Research 33, 2781–2798.",
    "Stow, C.A., Borsuk, M.E., Stanley, D.W., 2001. Long-term changes in watershed nutrient inputs and riverine exports in the Neuse River, North Carolina. Water Research 35, 1489–1499.",
    "Streeter, H.W., Phelps, E.B., 1925. A study of the pollution and natural purification of the Ohio River. US Public Health Service Bulletin 146.",
    "Strokal, M., Kroeze, C., Wang, M., Bai, Z., Ma, L., 2016. The MARINA model: Model description and results for China. Science of the Total Environment 562, 869–888.",
    "Vörösmarty, C.J., McIntyre, P.B., Gessner, M.O., et al., 2010. Global threats to human water security and river biodiversity. Nature 467, 555–561.",
    "Wang, W., Liu, G., Zhang, Y., et al., 2024. Enhancing watershed management through adaptive source apportionment under a changing environment. npj Clean Water 7, 29. https://doi.org/10.1038/s41545-024-00325-6.",
    "Wang, Y.J., Xue, M., Luo, J.H., et al., 2026. Bottom-up high-resolution water pollution source emission inventory accounting method and spatial analysis. Journal of East China Normal University (Natural Science) 1, 132–139.",
    "Withers, P.J.A., Jarvie, H.P., 2008. Delivery and cycling of phosphorus in rivers: A review. Science of the Total Environment 400, 379–395.",
    "Wollheim, W.M., Vörösmarty, C.J., Bouwman, A.F., et al., 2008. Global N removal by freshwater aquatic systems using a spatially distributed, within-basin approach. Global Biogeochemical Cycles 22, GB2026.",
    "Wollheim, W.M., Vörösmarty, C.J., Peterson, B.J., et al., 2006. Relationship between river size and nutrient removal. Geophysical Research Letters 33, L06410.",
]


def build_manuscript():
    source = PAPER / "Manuscript_Final_MC-SIRC.docx"
    doc = Document(source)
    clear_body_from_paragraph(doc, 10)
    set_document_format(doc, line_numbers=True)

    # Front matter retained from the submitted file.
    doc.paragraphs[0].text = TITLE
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in doc.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.paragraphs[7].text = "Abstract"
    doc.paragraphs[7].runs[0].bold = True
    clear_paragraph(doc.paragraphs[8])
    abstract = doc.paragraphs[8]
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.line_spacing = 1.5
    parts = [
        ("Study region: ", "The Nanchuan River Basin (1,438 km²) is a flashy hilly-gully catchment on the Loess Plateau, China. The case study combines a 2022 bottom-up inventory (53,155 non-point-source grids and 100 georeferenced point sources) with 5,928 observed hourly records at one outlet station."),
        ("Study focus: ", "We reformulate Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) as an identifiability-aware workflow. Missing-not-at-random gaps are represented by four annual-load scenarios; Bayesian maximum a posteriori and Markov chain Monte Carlo analyses provide prior-regularized discrepancy allocations; Jacobian-rank, null-space, and k = 0 profile diagnostics determine which quantities the data can support; and forward Monte Carlo simulation describes inventory-side uncertainty without treating it as outlet validation."),
        ("New hydrological insights for the region: ", "The single annual outlet constraint identifies only one aggregate linear combination of 9–10 unknowns (Jacobian rank 1; nullity 8–9), so source-specific factors cannot be interpreted as validated coefficients. Four boundary prior conflicts persist under all missing-data scenarios and are retained only as field-audit priorities. After analytically profiling the global scale parameter, the monthly distance-decay model adds virtually no explanatory power over a k = 0 flow-only null (ΔR² ≤ 0.0033); half-life distances, effective-contribution shares, and source-ranking inversions are therefore withdrawn. The defensible regional result is a quantified inventory–monitoring discrepancy and a monitoring-design diagnosis: tributary sections and independently observed source seasonality are required before source-specific transport or intervention effects can be estimated."),
    ]
    for idx, (label, text) in enumerate(parts):
        run = abstract.add_run(label)
        run.bold = True
        abstract.add_run(text)
        if idx < len(parts) - 1:
            abstract.add_run().add_break()
    doc.paragraphs[9].text = (
        "Keywords: source inventory; outlet load; structural identifiability; Bayesian regularization; "
        "missing-not-at-random data; null model; Loess Plateau"
    )

    add_heading(doc, "1 Introduction", 1)
    add_body(doc, "Linking pollutant sources to downstream water-quality loads is central to watershed management, yet bottom-up emission inventories and instream monitoring are usually developed as separate evidence systems. Inventories resolve activities and source categories but typically terminate at source emission or nominal river-entry loads; monitoring records the integrated response of source inputs, hydrological delivery, channel processing, and unrepresented loads. This scale mismatch can produce large discrepancies without revealing which process is responsible (Carpenter et al., 1998; Ongley et al., 2010; Strokal et al., 2016).")
    add_body(doc, "Recent watershed source-apportionment studies increasingly combine source data, landscape attributes, and river-network transport. SPARROW applications in the Beiyun and upper Yangtze basins quantify nutrient sources and retention using spatially distributed monitoring information (Liu et al., 2023; Li et al., 2023), while Bayesian hierarchical work across contrasting flow regimes demonstrates the value of multiple tributary stations for separating export processes (Neumann et al., 2023). Dynamic inventory–transport frameworks also show that source priorities can change with environmental conditions (Wang et al., 2024). These advances share an important design feature: spatial or temporal replication supplies contrasts that permit source and transport parameters to be estimated.")
    add_body(doc, "Data-limited basins often provide only one outlet station. In that configuration, an annual mass balance supplies one equation for many source factors, and monthly calibration does not automatically create spatial information when all or most source inputs share the same temporal allocation. Bayesian priors can regularize such problems, but regularization is not identifiability: a narrow posterior at a hard bound may reflect the bound and prior geometry rather than information in the observations (Gelman et al., 2013). Likewise, matching one annual load is internal reconciliation, not independent validation.")
    add_body(doc, "A second challenge is incomplete monitoring. The Nanchuan archive contains long gaps concentrated in the July–November flood season, so the missingness is not plausibly random. Event-driven pollutant loads may consequently depend on the gap treatment. A credible reconciliation must therefore expose load sensitivity to alternative missing-hour assumptions instead of conditioning all conclusions on a single imputed total.")
    add_body(doc, "This study develops MC-SIRC as an identifiability-aware reconciliation and classification workflow rather than a source-coefficient validation method. The objectives are to: (1) quantify the aggregate discrepancy between the inventory-derived nominal river-entry load and the outlet load under explicit missing-data scenarios; (2) use prior-regularized Bayesian allocations only to flag strong prior–data conflicts for field audit; (3) formally diagnose the annual inversion and spatial-decay model using Jacobian rank, exactly equivalent null-space solutions, and a k = 0 flow-only null; and (4) identify the additional monitoring contrasts required for source-specific estimation. The study is a methodological case analysis for one Loess Plateau basin; cross-basin transfer is not tested.")

    add_heading(doc, "2 Materials and methods", 1)
    add_heading(doc, "2.1 Study area", 2)
    add_body(doc, "The Nanchuan River Basin is located in Zhongyang County, Lüliang City, Shanxi Province, in the middle Yellow River Basin. The 1,438 km² catchment has a 47.8 km main channel, multi-year mean precipitation of 553 mm, and perennial baseflow of approximately 0.5 m³ s⁻¹. Its hilly-gully Loess Plateau terrain produces a flashy monsoon response, with rainfall and sediment transport concentrated in summer. One automated station at the watershed outlet records water quality and discharge.")
    add_picture(doc, FIG / "figure1_revised_framework.png", 6.6)
    add_caption(doc, "Figure 1. (a) Nanchuan River Basin, georeferenced point sources, and the outlet monitoring station. (b) Revised identifiability-aware MC-SIRC workflow. Source-specific attenuation rates and effective-contribution rankings are explicitly excluded from supported outputs under the present single-station design.")

    add_heading(doc, "2.2 Data sources and missingness", 2)
    add_body(doc, "The bottom-up 2022 inventory contains five gridded non-point categories across 53,155 1 km × 1 km cells and four point-source categories comprising 100 georeferenced facilities or farms (Table S1). The inventory compilation and spatial-accounting protocol are described by Wang et al. (2026). The non-point categories are rural domestic, agricultural cultivation, aquaculture, urban NPS, and dispersed urban domestic emissions; point sources include 50 household-livestock farms, 34 large-scale livestock farms, 15 industrial sources, and one centralized treatment facility. Outlet data include hourly COD, NH₃-N, TN, TP, and discharge. The 2022 archive contains 5,928 unique hourly timestamps from a possible 8,760, giving 67.7% coverage and 2,832 absent hours. Missingness is concentrated in July–November, when monthly coverage falls to 9.3–62.8%, and is therefore treated as missing not at random (MNAR).")
    add_body(doc, "Within observed timestamps, concentration gaps were linearly interpolated and implausible discharge values were replaced using a three-tier hierarchy: the same calendar time in other archive years, adjacent-month medians, and linear interpolation. Completely absent timestamps were not silently reconstructed. Instead, four annual-load scenarios were evaluated: S1 observed hours only (lower bound); S2 month-specific coverage scaling (revised default); S3 annual-mean scaling by 8,760/5,928; and S4 an event-weighted upper sensitivity that multiplies only the S2-imputed portion by 1.5. S2 preserves the observed seasonal load structure and avoids exporting the high April–June mean to the poorly observed late-flood-season months. Its limitations are retained in the Discussion.")

    add_heading(doc, "2.3 Workflow and inferential roles", 2)
    add_body(doc, "Each MC-SIRC component has a distinct inferential role (Table 1). The workflow separates descriptive accounting, prior-regularized allocation, forward uncertainty propagation, and formal identifiability auditing. None of the latter three is described as an independent validation of source coefficients because they share either the same outlet load, the same inventory, or both. Supplementary methods, diagnostic figures, and detailed tables are organized in Texts S1–S8, Figures S1–S2, and Tables S1–S35, respectively.")
    add_caption(doc, "Table 1. MC-SIRC components, assumptions, uncertainties, and supported inferential roles.")
    add_table(
        doc,
        ["Component", "Input", "Output", "Key assumption", "Main uncertainty", "Supported role"],
        [
            ["Inventory accounting", "Activity data; emission and river-entry coefficients", "Emission and nominal river-entry loads", "Inventory definitions are internally consistent", "Activity and coefficient error", "Bottom-up baseline"],
            ["Outlet load calculation", "Concentration; discharge; timestamp coverage", "S1–S4 annual loads", "Scenario-specific missing-hour behavior", "Flood-season MNAR gaps", "Observed aggregate constraint"],
            ["Bayesian reconciliation", "Nominal source loads; one annual outlet load; priors", "MAP/MCMC regularized allocations", "Priors regularize a rank-deficient equation", "Prior choice; compensation; bounds", "Prior-conflict screening only"],
            ["Forward Monte Carlo", "Inventory distributions", "Nominal river-entry-load envelope", "Specified input distributions", "No in-stream transport represented", "Inventory-side uncertainty"],
            ["Annual identifiability audit", "Mass-balance Jacobian", "Rank, nullity, equivalent solutions", "Local linear structure reflects estimability", "Structural rather than sampling uncertainty", "Defines supported parameter combinations"],
            ["Spatial profile audit", "Monthly allocations; distances; outlet loads", "Profile objective and k = 0 comparison", "Assumed source calendars", "k–γ ridge; bounds; one station", "Tests whether distance adds information"],
        ],
        font_size=7.5,
    )

    add_heading(doc, "2.4 Outlet load and discrepancy ratios", 2)
    add_body(doc, "For each observed hour, pollutant load was calculated as the product of concentration and discharge and integrated over time:")
    add_equation(doc, "L = 3.6 Σₕ CₕQₕΔt,                                                        (1)")
    add_body(doc, "where L is load (kg), Cₕ is concentration (mg L⁻¹), Qₕ is discharge (m³ s⁻¹), Δt = 1 h, and 3.6 is the unit-conversion factor. The inventory defines emissions E and nominal river-entry loads R. Ratios R/E, M/R, and M/E are reported descriptively. Because M/R combines inventory bias, missing sources, missing-data treatment, and in-channel processing, it is termed an outlet-to-nominal-entry discrepancy ratio, not a measured channel-transport coefficient.")

    add_heading(doc, "2.5 Prior-regularized aggregate reconciliation", 2)
    add_body(doc, "For pollutant p, the one-station annual reconciliation is")
    add_equation(doc, "Mₚ = Σᵢ Eᵢ,ₚ αᵢ,ₚ fᵢ,ₚ + Uₚ + εₚ,                                      (2)")
    add_body(doc, "where α is the nominal inventory river-entry coefficient, f is a multiplicative discrepancy allocation, U is a non-negative unidentified-load term, and ε is observation/model discrepancy. Eight or nine active source factors plus U are estimated from one annual M. Truncated-normal priors f ~ TN(μ, σ²; [0.1, 2.0]) and U ~ Gamma(2, 10/M) regularize the equation. MAP estimates used 30 L-BFGS-B starts with seed 42. MCMC used the emcee ensemble sampler with 32 walkers, 20,000 steps, 5,000 burn-in steps, and seed 42 (Foreman-Mackey et al., 2013).")
    add_body(doc, "The standardized prior shift z = |fMAP − μ|/σ is retained as a prior-conflict diagnostic: Tier A, z > 2; Tier B, 1 < z ≤ 2; Tier C, z ≤ 1. These tiers do not measure identifiability or reliability. Tier A means only that the regularized reconciliation is in strong conflict with the stated prior and therefore identifies a field-audit priority. MAP, MCMC, and prior sensitivity are complementary summaries of the same likelihood–prior system, not independent evidence lines.")

    add_heading(doc, "2.6 Formal identifiability and cross-year consistency checks", 2)
    add_body(doc, "For Eq. (2), the Jacobian with respect to the active source factors and U has one row. Its rank, nullity, and identifiable linear combinations were calculated directly. To demonstrate non-uniqueness constructively, pairwise null-space perturbations were applied to MAP allocations while preserving all bounds and the aggregate source sum to numerical precision. Missing-data robustness was assessed by repeating MAP reconciliation under S1–S4 and comparing prior-conflict flags and rank correlations of the regularized component ordering.")
    add_body(doc, "The same annual-mean protocol was also applied to 2020, 2021, and 2023 monitoring archives with the 2022 inventory fixed. This is labelled a cross-year consistency check rather than validation: the factors necessarily absorb year-specific hydrology and unobserved source changes. The 2020–2021 archives have only 1.0–1.3% paired hourly coverage and are treated as data-poor.")

    add_heading(doc, "2.7 Forward inventory uncertainty", 2)
    add_body(doc, "A forward Monte Carlo analysis (10,000 iterations; seed 42) sampled source emissions from normal distributions with CV = 20% and nominal river-entry coefficients from uniform [0.5, 1.5]× ranges. The output is a distribution of nominal river-entry load before in-stream transport. Comparing that distribution with the outlet load visualizes the scale discrepancy but cannot attribute it to inventory bias because channel loss is omitted. Elasticity identifies which inventory inputs dominate forward variance; it is not a source-contribution estimate at the outlet.")

    add_heading(doc, "2.8 Spatial distance-decay identifiability audit", 2)
    add_body(doc, "The previously calibrated spatial model is retained only as an identifiability case study:")
    add_equation(doc, "L̂ₚ,ₘ = γₚ Σᵢ wᵢ,ₘ Eᵢ,ₚ αᵢ,ₚ exp(−kₚdᵢ),                              (3)")
    add_body(doc, "where wᵢ,ₘ allocates annual source load to month m, dᵢ is distance to the outlet, k is attenuation, and γ is a global scale. Most source types, including all point sources, use the same discharge-derived monthly weights; agricultural and aquaculture categories use assumed calendars multiplied by discharge. If wᵢ,ₘ = wₘ for all sources, Eq. (3) reduces exactly to L̂ₘ = γwₘA(k), with A(k) = ΣEᵢαᵢexp(−kdᵢ). Only the product γA(k) is identifiable, so any k can be offset by γ. The small departures created by assumed agricultural/aquaculture calendars do not constitute independent observations of source seasonality.")
    add_body(doc, "We therefore profiled γ analytically at each k ∈ [0, 0.3] km⁻¹ without the former γ ∈ [0.3, 5] bound, compared the resulting fit with the k = 0 flow-only null, and repeated the calculation with the original bounds to diagnose boundary effects. The profile, rather than a point estimate, determines whether distance contributes information. No half-life distance or effective-contribution share is reported unless k is independently identifiable.")

    add_heading(doc, "3 Results and discussion", 1)
    add_heading(doc, "3.1 Inventory scale and aggregate reconciliation", 2)
    add_body(doc, "The inventory contains 1,036.5 t COD, 20.0 t NH₃-N, 131.3 t TN, and 14.0 t TP emissions, corresponding to nominal river-entry totals of 463.9, 6.84, 69.15, and 5.32 t, respectively (Table S8). At the source-emission scale, large-scale livestock accounts for 56.6% of COD and 65.1% of TP emissions, whereas the centralized facility accounts for 37.1% of TN emissions. These are inventory-composition shares, not estimates of source contributions at the outlet. Under the revised default S2 missing-data treatment, outlet loads are 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. These differences establish an aggregate inventory–outlet discrepancy. They do not by themselves separate coefficient error from channel transformation or omitted sources.")
    add_body(doc, "The prior-regularized reconciliation remains close to each target by construction (Table 2). The residual deviations of +2.0% to +12.0% describe the chosen observation-error and prior trade-off, not predictive validation. The unidentified term remains 4.5–8.3% of the outlet load under S2.")
    add_caption(doc, "Table 2. Prior-regularized aggregate reconciliation under the revised default S2 missing-data scenario. Deviations are constructional and are not independent validation statistics.")
    missing_summary = pd.read_excel(RESULT_XLSX, sheet_name="Missing_summary")
    missing_detail = pd.read_excel(RESULT_XLSX, sheet_name="Missing_detail")
    default_summary = missing_summary[missing_summary.Scenario == "S2 month-specific (default)"]
    default_detail = missing_detail[missing_detail.Scenario == "S2 month-specific (default)"]
    rows = []
    for _, row in default_summary.iterrows():
        source_sum = default_detail[default_detail.Pollutant == row.Pollutant].Regularized_component_t.sum()
        unknown = row.Reconciled_t - source_sum
        rows.append([
            {"COD": "COD", "氨氮": "NH₃-N", "总氮": "TN", "总磷": "TP"}[row.Pollutant],
            f"{row.Monitored_t:.3f}", f"{row.Reconciled_t:.3f}", f"{row.Deviation_pct:+.1f}%",
            f"{unknown:.3f} ({unknown / row.Monitored_t * 100:.1f}%)",
            f"{int(row.Tier_A_count)}/{int(row.Tier_B_count)}/{int(row.Tier_C_count)}",
        ])
    add_table(doc, ["Pollutant", "Outlet load (t)", "Reconciled load (t)", "Deviation", "Unidentified term", "A/B/C tiers"], rows, font_size=8.5)

    add_heading(doc, "3.2 What the Bayesian reconciliation does—and does not—identify", 2)
    add_body(doc, "The annual Jacobian has rank 1 for every pollutant, with nullity 8 for COD and TN and 9 for NH₃-N and TP (Figure 2d; Table S32). Thus, only one aggregate linear combination is observed. Exactly equivalent solutions make the consequence concrete: for TN, increasing the large-livestock factor from 0.766 to 1.753 while decreasing the centralized-facility factor from 0.746 to 0.528 leaves the regularized source sum unchanged to machine precision; analogous paired solutions exist for all pollutants (Table S33). Source factors are therefore allocations selected by priors and bounds, not recoveries of unique true coefficients.")
    add_picture(doc, FIG / "figure2_identifiability_aware_bayesian.png", 6.5)
    add_caption(doc, "Figure 2. Identifiability-aware Bayesian reconciliation. (a) S2 MAP discrepancy factors. (b) A/B/C tiers redefined as strong/moderate/no resolved prior conflict; they do not denote data identifiability. (c) Four boundary conflicts persist under all missing-data treatments. (d) One annual observation yields Jacobian rank 1 and nullity 8–9.")
    add_body(doc, "Under S2, five pairs meet the strong prior-conflict threshold: large-livestock COD; rural-domestic TP; and TP for large livestock, industrial sources, and the centralized facility. Four of these—the large-livestock COD/TP and industrial/centralized TP pairs—remain at the 0.10 truncation bound under S1–S4. Rural-domestic TP changes from Tier A under S1–S2 to Tier B under S3–S4. Across all pairs, tier counts and some regularized component rankings change with the monitored-load scenario; for example, the highest NH₃-N component shifts between rural domestic and large livestock. The persistent four are therefore field-verification priorities, whereas no tier is presented as a validated coefficient.")
    add_body(doc, "The former narrow MCMC interval for large-livestock TP (0.100–0.152 under S3) is piled against the 0.10 truncation bound. Its narrowness is consequently evidence of boundary conflict within the stated model, not strong data information. Likewise, prior-sensitivity calculations share the same likelihood and do not provide independent validation. Their legitimate use is to reveal whether a flag disappears when prior assumptions change.")
    add_body(doc, "Cross-year application reinforces the limitation. The four original boundary conflicts reappear in the two better monitored years, 2022 and 2023, but the fixed 2022 inventory is paired with monitored COD loads that vary 3.6-fold across 2020–2023. The correction factors therefore absorb interannual hydrology and source changes. We report this as a consistency pattern, not confirmation of source coefficients, and exclude the 2020–2021 sparse archives from replication claims.")
    add_body(doc, "Forward Monte Carlo distributions describe nominal river-entry loads before transport (Figure 3). COD and TP distributions remain far above the outlet load, whereas TN overlaps it more closely. Inventory-side elasticities identify large-scale livestock as the most influential input for nominal COD (0.391), NH₃-N (0.364), and TP (0.528) river-entry totals, and the centralized facility for TN (0.704) (Table S25). These elasticities quantify sensitivity of the pre-transport inventory total, not source contributions at the outlet. Because in-stream loss is absent from this calculation, the comparison cannot distinguish inventory overestimation from channel retention. The analysis is retained only to communicate input uncertainty and to identify influential inventory terms.")
    add_picture(doc, FIG / "figure3_forward_uncertainty.png", 6.5)
    add_caption(doc, "Figure 3. Forward Monte Carlo uncertainty in nominal river-entry loads before in-stream transport for (a) COD, (b) NH₃-N, (c) TN, and (d) TP. Histograms show Monte Carlo probability densities; solid blue lines mark the Monte Carlo means, gray shading denotes the 5th–95th percentile intervals, and red dashed lines mark the default S2 outlet loads. Separation between the inventory-side distributions and outlet loads is an inventory–outlet scale discrepancy and does not independently establish source-coefficient bias or in-stream loss.")

    add_heading(doc, "3.3 Spatial-decay audit and Loess Plateau process interpretation", 2)
    add_body(doc, "Once γ is profiled without the former lower/upper bounds, the distance-decay profiles are essentially flat relative to k = 0 (Figure 4; Table 3). The incremental ΔR² is 0.00003 for COD, 0.00019 for NH₃-N, 0.00333 for TN, and 0.00069 for TP. For COD and TP, the unconstrained k = 0 γ values (0.225 and 0.077) fall below the previous γ ≥ 0.3 bound; the bound made k > 0 appear necessary. NH₃-N, TN, and TP profile minima occur at the edge of the tested k range or require γ well above the former upper bound. These are constraint and scaling effects, not independently estimated attenuation rates.")
    add_picture(doc, FIG / "figure4_spatial_identifiability_profile.png", 6.5)
    add_caption(doc, "Figure 4. Spatial identifiability profiles for (a) COD, (b) NH₃-N, (c) TN, and (d) TP, with γ solved analytically at each k. Blue curves show relative-error SSE divided by its profile minimum; red dashed lines mark the k = 0 flow-only null, and gray dotted lines mark the profile-minimum k rather than an independently identified attenuation rate. Annotations report ΔR² relative to k = 0. The negligible ΔR² shows that monthly outlet data add virtually no distance information under shared or assumed source calendars.")
    add_caption(doc, "Table 3. k = 0 null-model comparison after profiling γ without artificial bounds.")
    spatial = pd.read_excel(SPATIAL_XLSX, sheet_name="Summary")
    rows = []
    for _, row in spatial.iterrows():
        rows.append([
            {"COD": "COD", "氨氮": "NH₃-N", "总氮": "TN", "总磷": "TP"}[row.Pollutant],
            f"{row.R2_k0_free:.4f}", f"{row.R2_best_free:.4f}", f"{row.Delta_R2_free:.5f}",
            f"{row.gamma_k0_free:.3f}", "Not identifiable",
        ])
    add_table(doc, ["Pollutant", "R², k=0", "Best profile R²", "ΔR²", "γ at k=0", "Inference for k"], rows, font_size=8.5)
    add_body(doc, "The absence of an identifiable k does not imply that in-stream processes are unimportant. In Loess Plateau rivers, TP transport reflects particulate settling during low flow and remobilization during storm-driven sediment pulses; the direction and magnitude of apparent retention can therefore change with discharge and particle size (Withers and Jarvie, 2008). Ammonium is transformed by nitrification, while nitrate and total nitrogen are affected by denitrification, oxygen status, residence time, temperature, and benthic exchange (Birgand et al., 2007; Alexander et al., 2009; Beaulieu et al., 2011). Flashy hydrographs shorten residence time during events but increase particulate transport, whereas low-flow periods favor contact with sediments and microbial processing (Alexander et al., 2000; Behrendt and Opitz, 2000; Wollheim et al., 2006; Wollheim et al., 2008). These mechanisms explain why a single constant k is physically incomplete, but the present outlet record cannot estimate more complex process parameters either. Matching process detail to available observations is a recurring challenge in integrated water-quality modelling (Rode et al., 2010); multi-section concentration–discharge profiles and particulate/dissolved measurements are required here.")
    add_body(doc, "Because k is not identified, the previously reported half-life distances, effective-contribution shares, point-source percentages, and emission-to-effective ranking inversions are withdrawn. Distance and nominal river-entry coefficients remain useful descriptors for designing monitoring locations, but they do not support quantitative intervention rankings in this dataset.")

    add_heading(doc, "3.4 Decision-support implications", 2)
    add_body(doc, "MC-SIRC now supports two deliberately limited decisions. First, persistent prior conflicts identify where field audits have the greatest diagnostic value: large-livestock COD/TP and industrial/centralized TP records should be checked against production statistics, discharge permits, treatment performance, and contemporaneous sampling. This is a hypothesis-generation step; low regularized factors are equally compatible with inventory bias, unrepresented channel loss, and compensating model terms. Second, the rank and profile audits identify the measurements needed before source control can be optimized: at least several tributary or control-unit stations, source-specific seasonal discharge information, and process-resolved measurements for nitrogen and particulate phosphorus.")
    add_body(doc, "Operational measures should therefore not be re-ranked from the present factors. Any near-source/far-source strategy is treated as a testable field hypothesis, to be evaluated with intervention or multi-section data before policy use. This distinction protects readers from inheriting the review context: the published workflow produces an evidence map—what is discrepant, what is underdetermined, and what to measure next—rather than a hidden set of calibrated policy coefficients.")

    add_heading(doc, "3.5 Limitations and transfer scope", 2)
    add_body(doc, "The primary limitation is structural: one annual outlet observation cannot identify 9–10 source terms, and one monthly outlet series cannot separate distance decay from a global scale when source inputs share temporal weights. Priors and bounds yield stable computations but do not create information. A second limitation is the MNAR monitoring gap. S2 uses month-specific coverage to avoid global seasonal extrapolation, but months with 9–28% coverage remain uncertain. The four persistent boundary conflicts survive S1–S4; other tier assignments and some rankings do not, and the latter are not used for decisions.")
    add_body(doc, "The cross-year check is not independent because the 2022 inventory is fixed, the monitored load changes substantially, and two years are extremely sparse. Annual inventories and contemporaneous multi-site observations would be needed to separate source changes from hydrology. Finally, the study concerns one Loess Plateau basin. Its source definitions, hydrological seasonality, sediment regime, and monitoring design are not representative of humid plains, coastal urban basins, or semi-arid grasslands without re-parameterization and new identifiability analysis. MC-SIRC is therefore presented as a methodological prototype and audit template, not a demonstrated transferable calibration.")

    add_heading(doc, "4 Conclusions", 1)
    add_body(doc, "MC-SIRC reconciles a detailed watershed source inventory with outlet monitoring while making the limits of that reconciliation explicit. In the Nanchuan case, the revised month-specific missing-data scenario gives outlet loads of 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. Bayesian regularization can allocate the aggregate discrepancy and flag persistent prior conflicts, but the annual Jacobian rank of 1 and nullity of 8–9 show that individual factors are not identifiable. Four boundary conflicts persist across all missing-data scenarios and warrant field verification; they do not establish causal coefficient overestimation.")
    add_body(doc, "The spatial audit reaches the same boundary: after profiling γ, a k = 0 flow-only model is effectively indistinguishable from the best decay profile (ΔR² ≤ 0.0033). Accordingly, half-life distances, effective-contribution shares, ranking inversions, and direct policy prescriptions have been removed. The principal contribution is an identifiability-aware workflow that converts disagreement between inventory and monitoring into transparent aggregate diagnostics and a concrete monitoring roadmap. Multi-section, seasonally resolved, and process-specific data are prerequisites for source-level calibration and transfer beyond this case basin.")

    add_heading(doc, "Data availability statement", 1)
    add_body(doc, "Analysis code is available at https://github.com/Shifa-Zhong/MC-SIRC. The exact public snapshot used as the revision baseline is permanently addressable at https://github.com/Shifa-Zhong/MC-SIRC/tree/126fa0a789751ecf0850b4971244f26cd404a880. Input schemas, parameter settings, fixed random seeds, and commands for regenerating the main tables and figures are listed in SI Text S8 and the repository reproduction guide. An anonymized example package reproduces the file structure without disclosing restricted station records. Raw hourly water-quality data are controlled by the local environmental authority and may be requested from the corresponding author for academic use, subject to approval.")
    add_heading(doc, "Author contributions", 1)
    add_body(doc, "Yujie Wang: data analysis and manuscript writing. Shifa Zhong: study design, data processing, model construction, identifiability analysis, and manuscript revision. Cheng Zhang: GIS spatial analysis. Jinhong Luo and Weifeng Zhang: project supervision and manuscript revision.")
    add_heading(doc, "Funding", 1)
    add_body(doc, "This study was supported by the Shanxi Provincial Key Laboratory for Water Pollution Prevention and Utilization (Project No. 202404010931034).")
    add_heading(doc, "References", 1)
    for reference in sorted(REFERENCES, key=lambda value: value.lower()):
        p = doc.add_paragraph(reference)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)

    for table in doc.tables:
        format_table(table)
    out = PAPER / "Manuscript_Final_MC-SIRC.docx"
    doc.save(out)
    shutil.copy2(out, ARCHIVE_DUP / "Manuscript_Revised_20260824.docx")
    return out


def build_abstract():
    doc = Document(PAPER / "Abstract.docx")
    # Rebuild a compact standalone structured abstract while retaining section settings.
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    set_document_format(doc, line_numbers=False)
    from docx.enum.style import WD_STYLE_TYPE
    if 'Table Grid' not in [style.name for style in doc.styles]:
        doc.styles.add_style('Table Grid', WD_STYLE_TYPE.TABLE)
    p = doc.add_paragraph(TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    p = doc.add_paragraph("Abstract")
    p.runs[0].bold = True
    sections = [
        ("Study region", "The Nanchuan River Basin (1,438 km²) is a flashy hilly-gully catchment on the Loess Plateau, China. The case study combines a 2022 bottom-up inventory with 5,928 observed hourly records at one outlet station."),
        ("Study focus", "MC-SIRC is reformulated as an identifiability-aware workflow combining four missing-data scenarios, prior-regularized Bayesian reconciliation, Jacobian-rank and null-space diagnostics, forward inventory uncertainty, and a k = 0 spatial null comparison."),
        ("New hydrological insights for the region", "One annual outlet constraint identifies only one aggregate linear combination of 9–10 unknowns (rank 1; nullity 8–9). Four boundary prior conflicts persist across missing-data scenarios but are retained only as field-audit priorities. After profiling the global scale, the distance-decay model adds virtually no explanatory power over k = 0 (ΔR² ≤ 0.0033); half-lives, effective shares, and ranking inversions are withdrawn. Multi-section and source-seasonal observations are required for source-specific transport inference."),
    ]
    for label, text in sections:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.add_run(label + ": ").bold = True
        p.add_run(text)
    doc.add_paragraph("Keywords: source inventory; outlet load; structural identifiability; Bayesian regularization; missing-not-at-random data; null model; Loess Plateau")
    doc.save(PAPER / "Abstract.docx")


def build_si():
    doc = Document(PAPER / "SI_Final.docx")
    set_document_format(doc, line_numbers=False)
    from docx.enum.style import WD_STYLE_TYPE
    if 'Table Grid' not in [style.name for style in doc.styles]:
        doc.styles.add_style('Table Grid', WD_STYLE_TYPE.TABLE)
    doc.paragraphs[1].text = TITLE

    # First-use abbreviation definitions.
    anchor = doc.paragraphs[7]
    abbr = simple_insert_after(
        anchor,
        "Abbreviations: maximum a posteriori (MAP); Markov chain Monte Carlo (MCMC); "
        "non-point source (NPS); chemical oxygen demand (COD); ammonia nitrogen (NH₃-N); "
        "total nitrogen (TN); total phosphorus (TP).",
    )
    abbr.paragraph_format.line_spacing = 1.0

    replacements = {
        "Text S1.": ("Text S1. Missing-data processing and annual-load scenarios", "The 2022 archive contains 5,928 unique hourly timestamps from a possible 8,760 (67.7% coverage), so 2,832 timestamps are absent. Gaps cluster in July–November and are treated as MNAR. Within observed timestamps, concentration gaps were linearly interpolated; discharge gaps or invalid values were filled by same-calendar historical medians, adjacent-month medians, and then interpolation. Completely absent hours were handled by four explicit scenarios: S1 observed-only; S2 month-specific coverage scaling (default); S3 annual-mean scaling by 8,760/5,928 = 1.4777; and S4 observed load plus 1.5 times the S2-imputed portion. S2 is preferred because it preserves the observed monthly structure instead of exporting the April–June mean into the poorly observed late-flood-season months."),
        "Text S2.": ("Text S2. Urban NPS inventory disaggregation", "Urban NPS functional-zone weights are retained as descriptive inventory inputs. The former 18-parameter secondary optimization is withdrawn because it introduced no new observation and could not identify zone-specific concentrations. No optimized Cj values are used in the revised manuscript."),
        "Text S3.": ("Text S3. Forward Monte Carlo inventory uncertainty", "The 10,000-iteration analysis samples emissions (normal, CV = 20%) and nominal river-entry coefficients (uniform [0.5, 1.5]×). It produces a pre-transport river-entry-load distribution. Comparison with the outlet load is descriptive and cannot distinguish inventory bias from in-stream loss. The calculation is not a prior-free or independent validation of the Bayesian reconciliation."),
        "Text S4.": ("Text S4. Spatial inputs retained for identifiability auditing", "Non-point sources are aggregated by control unit with representative distances of 8, 12, 22, and 30 km; point-source GPS distances range from 2.3 to 99.6 km. Nominal α = R/E values and distances are descriptive inputs. The revised analysis does not infer half-life distance or effective outlet contribution from them."),
        "Text S5.": ("Text S5. Cross-year consistency check", "The annual-mean reconciliation was repeated for 2020–2023 with the 2022 inventory fixed. This design is a consistency check, not independent validation: factors absorb interannual hydrology and source changes. The 2020–2021 paired hourly coverage is only 1.3% and 1.0%, so replication statements are restricted to 2022–2023 and remain conditional on the fixed inventory."),
        "Text S6.": ("Text S6. Wet/dry comparison", "The 2022 (634 mm) and 2023 (462 mm) archives provide a wet/dry contrast, but the same 2022 inventory is used in both inversions. Apparent closure therefore measures the flexibility of prior-regularized reconciliation under different monitored totals; it is not evidence that source-specific factors are invariant across hydrological regimes."),
        "Text S7.": ("Text S7. Formal identifiability diagnostics", "Annual inversion: Eq. (2) has one observation and 9–10 unknowns. The Jacobian rank is 1 and its nullity is 8–9. Table S33 gives constructive pairs of coefficient allocations that differ substantially yet reproduce the same regularized source sum to numerical precision. Spatial model: with shared monthly weights, L̂m = γwmA(k), so only γA(k) is identifiable. Agricultural and aquaculture calendars create small assumed deviations from shared weighting but do not supply independently observed source seasonality. We analytically profiled γ without the former [0.3, 5] bounds and compared every k with k = 0. Across pollutants, ΔR² is 0.00003–0.00333 and the profile SSE ratio is 1.00008–1.0604. Thus k, half-life, and effective-contribution shares are not reported."),
    }
    for prefix, (title, text) in replacements.items():
        p = find_paragraph(doc, prefix)
        p.text = title
        # The explanatory paragraph is immediately after the title in current SI.
        idx = next(i for i, candidate in enumerate(doc.paragraphs) if candidate._p is p._p)
        if idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.startswith("Table S"):
            doc.paragraphs[idx + 1].text = text
        else:
            simple_insert_after(p, text)

    # Captions and explanatory notes.
    caption_updates = {
        "Table S1.": "Table S1. Overview of research data (data type, content, resolution, and volume).",
        "Table S3.": "Table S3. Missing-data scenarios and revised 2022 annual outlet loads.",
        "Table S5.": "Table S5. Descriptive urban NPS functional-zone weights (not calibrated).",
        "Table S6.": "Table S6. Spatial k = 0 null comparison after profiling γ without its former bounds.",
        "Table S11.": "Table S11. Inventory and outlet discrepancy ratios under default S2. Monitor/Entry is not interpreted as a channel-transport coefficient.",
        "Table S12.": "Table S12. MAP regularized discrepancy factors under default S2. z is prior conflict, not identifiability.",
        "Table S13.": "Table S13. MCMC posterior summaries—COD, annual-mean sensitivity scenario S3.",
        "Table S14.": "Table S14. MCMC posterior summaries—NH₃-N, annual-mean sensitivity scenario S3.",
        "Table S15.": "Table S15. MCMC posterior summaries—TN, annual-mean sensitivity scenario S3.",
        "Table S16.": "Table S16. MCMC posterior summaries—TP, annual-mean sensitivity scenario S3. Boundary-piled intervals must not be interpreted as data identifiability.",
        "Table S17.": "Table S17. MCMC parameter correlations under S3. Correlations summarize a prior-regularized posterior and are not formal identifiability diagnostics.",
        "Table S22.": "Table S22. Prior-conflict diagnostic tiers under default S2.",
        "Table S24.": "Table S24. Forward Monte Carlo uncertainty in nominal river-entry loads before in-stream transport.",
        "Table S26.": "Table S26. Missing-data sensitivity of the four persistent boundary conflicts.",
        "Table S27.": "Table S27. Cross-year MAP factors with the 2022 inventory held fixed (consistency check only).",
        "Table S28.": "Table S28. Cross-year prior-regularized reconciliation summary (constructional; not validation).",
        "Table S29.": "Table S29. Wet/dry consistency comparison with a fixed 2022 inventory.",
        "Table S30.": "Table S30. Four persistent boundary conflicts across archive years; 2020–2021 are data-poor.",
        "Table S31.": "Table S31. Spatial k = 0 null-model comparison with γ profiled analytically.",
        "Table S32.": "Table S32. Annual mass-balance Jacobian rank and nullity.",
        "Table S33.": "Table S33. Constructive examples of exactly equivalent source-factor allocations.",
        "Table S34.": "Table S34. Rank stability of regularized source components across missing-data scenarios.",
        "Table S35.": "Table S35. Point-source emission distribution by river-channel distance (descriptive only).",
    }
    for prefix, value in caption_updates.items():
        try:
            p = find_paragraph(doc, prefix)
        except KeyError:
            continue
        p.text = value
        for run in p.runs:
            run.bold = True

    residual_updates = {
        'Table S18.': 'Table S18. Prior sensitivity—COD, annual-mean sensitivity scenario S3.',
        'Table S19.': 'Table S19. Prior sensitivity—NH₃-N, annual-mean sensitivity scenario S3.',
        'Table S20.': 'Table S20. Prior sensitivity—TN, annual-mean sensitivity scenario S3.',
        'Table S21.': 'Table S21. Prior sensitivity—TP, annual-mean sensitivity scenario S3.',
        'Table S23.': 'Table S23. Observation-error sensitivity under annual-mean scenario S3.',
        'Table S25.': 'Table S25. Elasticity of nominal river-entry load to inventory inputs (not outlet contribution).',
        'A: z>2': 'A: z > 2, strong prior conflict (5 pairs under S2). B: 1 < z ≤ 2, moderate prior conflict (4 pairs). C: z ≤ 1, no resolved prior conflict (25 pairs). These tiers do not imply identifiability or reliability.',
        'Anomalous sources remained': 'The four persistent boundary conflicts remain near 0.100 across the tested observation-error scales. This is sensitivity to the assumed likelihood scale, not independent evidence of source identifiability.',
        'Pipe network coefficient': 'All four original boundary conflicts remain at 0.100 under S1–S4, whereas other tier assignments and some regularized component rankings change. The former urban-NPS secondary-optimization results have been withdrawn.',
        'Table S27 reports': 'The annual-mean results in Table S27 are retained for historical comparability. The corrected 2022 factor is 8,760/5,928 = 1.4777; the former 8,705-hour denominator was erroneous. This S3 total is a sensitivity scenario, not the revised S2 default.',
        '★ = Rating A': '★ marks the four boundary conflicts highlighted in the earlier analysis. Their recurrence is a conditional consistency pattern under a fixed inventory and prior system; it is not evidence that the source factors are data-identifiable.',
        'All cross-year deviations remain': 'The small residual deviations across years reflect the flexibility of a prior-regularized model with many factors and one annual constraint. They are reported for bookkeeping and are not validation statistics.',
        'All four pollutants achieve closure': 'Both years can be reconciled within the chosen observation-error scale, but the fixed 2022 inventory forces the factors to absorb hydrological and unobserved source changes. No causal inventory attribution is made from this comparison.',
        'The cross-year reproducibility': 'The 0.100 boundary recurs in 2022 and 2023, the two less-sparse archives. This supports field rechecking of those inventory entries but does not identify their true coefficients; 2020–2021 are not treated as independent replicates.',
        'Model definitions': 'Under shared monthly weights, Eq. (3) reduces to L̂m = γwmA(k), where A(k) = Σi Eiαi exp(−kdi). The likelihood therefore identifies γA(k), not k and γ separately. Agricultural and aquaculture calendars introduce only assumed, not observed, source-specific temporal contrasts.',
        'To facilitate computation': 'For each k from 0 to 0.30 km⁻¹, γ was solved analytically without bounds and the profile relative-error SSE and R² were calculated. The k = 0 flow-only null was then compared with the profile optimum. The original γ bounds were evaluated separately to expose boundary-created optima.',
        'Note: All paper-setup': 'For COD and TP, γ at k = 0 is 0.225 and 0.077, below the former lower bound 0.3. Imposing that bound makes positive k appear necessary. With γ unbounded above zero, ΔR² relative to k = 0 is ≤ 0.00333 for every pollutant.',
        'None of the three alternative': 'Rank 1 means that only one aggregate linear combination is observed. MAP/MCMC regularization chooses a location in the 8–9 dimensional null space but does not eliminate it. Table S33 provides exact constructive examples.',
        '† At boundary': 'The paired allocations in Table S33 preserve the regularized source sum to numerical precision while changing individual factors substantially. The construction is deliberately favorable because the unidentified-source term is held fixed; non-uniqueness is therefore structural.',
    }
    for prefix, value in residual_updates.items():
        try:
            find_paragraph(doc, prefix).text = value
        except KeyError:
            pass

    # Table S3.
    replace_table(
        doc, 2,
        ["Scenario", "Absent-hour assumption", "COD (t)", "NH₃-N (t)", "TN (t)", "TP (t)"],
        [
            ["S1 observed-only", "Zero additional load (lower bound)", "111.862", "3.903", "49.368", "0.571"],
            ["S2 monthly (default)", "Each month scaled by its own coverage", "140.864", "4.785", "57.774", "0.758"],
            ["S3 annual mean", "Observed mean × 8,760/5,928", "165.301", "5.767", "72.953", "0.843"],
            ["S4 event upper", "Observed + 1.5× S2-imputed portion", "155.365", "5.226", "61.977", "0.852"],
        ], font_size=8.0,
    )

    spatial = pd.read_excel(SPATIAL_XLSX, sheet_name="Summary")
    replace_table(
        doc, 5,
        ["Pollutant", "R² at k=0", "Best profile R²", "ΔR²", "γ at k=0", "Inference"],
        [[{"COD": "COD", "氨氮": "NH₃-N", "总氮": "TN", "总磷": "TP"}[r.Pollutant],
          f"{r.R2_k0_free:.4f}", f"{r.R2_best_free:.4f}", f"{r.Delta_R2_free:.5f}",
          f"{r.gamma_k0_free:.3f}", "k not identifiable"] for _, r in spatial.iterrows()],
        font_size=8.0,
    )

    # Table S11: revised default monitored values and neutral discrepancy labels.
    values = {
        "COD": (1036.52, 463.887, 140.864),
        "NH₃-N": (19.95, 6.836, 4.785),
        "TN": (131.298, 69.148, 57.774),
        "TP": (13.964, 5.320, 0.758),
    }
    replace_table(
        doc, 10,
        ["Pollutant", "Emission (t)", "Nominal entry (t)", "Outlet S2 (t)", "Entry/Emission", "Outlet/Entry", "Outlet/Emission"],
        [[p, f"{e:.3f}", f"{r:.3f}", f"{m:.3f}", f"{r/e:.3f}", f"{m/r:.3f}", f"{m/e:.3f}"]
         for p, (e, r, m) in values.items()], font_size=7.8,
    )

    detail = pd.read_excel(RESULT_XLSX, sheet_name="Missing_detail")
    default = detail[detail.Scenario == "S2 month-specific (default)"].copy()
    source_en = {
        "面-农村生活污染源": "Rural domestic", "面-农业面源": "Agricultural NPS", "畜禽散养": "Household livestock",
        "面-水产养殖": "Aquaculture", "面-城市面源": "Urban NPS", "面-城镇散排": "Dispersed urban",
        "规模畜禽养殖": "Large livestock", "点-工业源": "Industrial", "点-集中式污染治理设施": "Central facility",
    }
    poll_en = {"COD": "COD", "氨氮": "NH₃-N", "总氮": "TN", "总磷": "TP"}
    pivot = {}
    for _, r in default.iterrows():
        pivot.setdefault(r.Source, {})[r.Pollutant] = f"{r.MAP_factor:.3f} ({r.z:.2f}; {r.Prior_conflict_tier})"
    replace_table(
        doc, 11,
        ["Source", "COD f (z; tier)", "NH₃-N f (z; tier)", "TN f (z; tier)", "TP f (z; tier)"],
        [[source_en[s]] + [pivot.get(s, {}).get(p, "—") for p in ["COD", "氨氮", "总氮", "总磷"]]
         for s in source_en], font_size=7.4,
    )

    # Table S22: all tier rows, allowing Word to continue across pages with a repeated header.
    tier_rows = []
    for _, r in default.sort_values(["Prior_conflict_tier", "z"], ascending=[True, False]).iterrows():
        tier_rows.append([poll_en[r.Pollutant], source_en[r.Source], f"{r.Prior_mu:.2f}", f"{r.MAP_factor:.3f}",
                          f"{r.z:.2f}", r.Prior_conflict_tier])
    replace_table(doc, 24, ["Pollutant", "Source", "Prior μ", "MAP f", "z", "Prior-conflict tier"], tier_rows, font_size=7.2)

    # Table S24 MC probabilities relative to S2 outlet loads.
    mc_base = ROOT / "output" / "figures" / "figure3_monte_carlo"
    mc_files = {"COD": "panel_a_mc_samples_COD.csv", "NH₃-N": "panel_b_mc_samples_NH3N.csv",
                "TN": "panel_c_mc_samples_TN.csv", "TP": "panel_d_mc_samples_TP.csv"}
    monitored = {"COD": 140.864, "NH₃-N": 4.785, "TN": 57.774, "TP": 0.758}
    mc_rows = []
    for p, fn in mc_files.items():
        a = pd.read_csv(mc_base / fn).iloc[:, 0].to_numpy()
        mc_rows.append([p, f"{a.mean():.2f}", f"{a.std(ddof=1):.2f}", f"{np.percentile(a,5):.2f}",
                        f"{np.percentile(a,95):.2f}", f"{monitored[p]:.3f}", f"{(a > monitored[p]).mean()*100:.1f}%"])
    replace_table(doc, 26, ["Pollutant", "Mean (t)", "SD (t)", "5th (t)", "95th (t)", "Outlet S2 (t)", "P(entry > outlet)"], mc_rows, font_size=7.8)

    flags = pd.read_excel(RESULT_XLSX, sheet_name="Flag_persistence")
    scenarios = ["S1 observed-only", "S2 month-specific (default)", "S3 annual-mean", "S4 event-weighted upper"]
    flag_rows = []
    for _, r in flags.iterrows():
        flag_rows.append([poll_en[r.Pollutant], source_en[r.Source]] +
                         [f"{r[s + ' factor']:.3f} ({r[s]})" for s in scenarios] + [f"{int(r.Tier_A_scenarios)}/4"])
    replace_table(doc, 28, ["Pollutant", "Source", "S1", "S2", "S3", "S4", "Tier A persistence"], flag_rows, font_size=7.5)

    # New identifiability tables S31-S34.
    replace_table(
        doc, 33,
        ["Pollutant", "k=0 R²", "Best R²", "ΔR²", "SSE ratio k0/best", "k at best", "γ at k=0", "Conclusion"],
        [[poll_en[r.Pollutant], f"{r.R2_k0_free:.4f}", f"{r.R2_best_free:.4f}", f"{r.Delta_R2_free:.5f}",
          f"{r.SSE_ratio_k0_to_best_free:.4f}", f"{r.k_best_free:.3f}", f"{r.gamma_k0_free:.3f}", "Not identifiable"]
         for _, r in spatial.iterrows()], font_size=7.0,
    )
    rank = pd.read_excel(RESULT_XLSX, sheet_name="Annual_rank_nullity")
    replace_table(
        doc, 34,
        ["Pollutant", "Observations", "Source factors", "Unknown terms", "Parameters", "Jacobian rank", "Nullity"],
        [[poll_en[r.Pollutant], int(r.Annual_observations), int(r.Active_source_factors), int(r.Unknown_source_terms),
          int(r.Active_source_factors + r.Unknown_source_terms), int(r.Jacobian_rank), int(r.Nullity)] for _, r in rank.iterrows()],
        font_size=7.8,
    )
    eq = pd.read_excel(RESULT_XLSX, sheet_name="Equivalent_solutions")
    replace_table(
        doc, 35,
        ["Pollutant", "Increased source", "f₁", "f₂", "Decreased source", "f₁", "f₂", "Sum 1 (t)", "Sum 2 (t)"],
        [[poll_en[r.Pollutant], source_en[r.Source_increased], f"{r.f_original_i:.3f}", f"{r.f_alternative_i:.3f}",
          source_en[r.Source_decreased], f"{r.f_original_j:.3f}", f"{r.f_alternative_j:.3f}",
          f"{r.Regularized_source_sum_1_t:.6f}", f"{r.Regularized_source_sum_2_t:.6f}"] for _, r in eq.iterrows()],
        font_size=6.8,
    )
    ranks = pd.read_excel(RESULT_XLSX, sheet_name="Rank_stability")
    replace_table(
        doc, 36,
        ["Scenario", "Pollutant", "Spearman ρ vs S2", "Top regularized component", "Interpretation"],
        [[r.Scenario, poll_en[r.Pollutant], f"{r.Spearman_vs_default:.3f}", source_en[r.Top_source],
          "Ordering sensitivity only; not outlet apportionment"] for _, r in ranks.iterrows()], font_size=7.0,
    )

    # Standardize dagger attachment and table layout; abbreviate the wide S17 headers.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = re.sub(r"\s*†\s*", "†", run.text)
        format_table(table, font_size=7.5 if len(table.columns) >= 9 else 8.0)
    if len(doc.tables) > 17:
        abbreviations = ["Source", "Rural", "Agri. NPS", "HH livest.", "Aqua.", "Urban NPS", "Disp. urban", "Large livest.", "Industrial", "Central fac.", "Unknown"]
        for table_index in [16, 17, 18, 19]:
            table = doc.tables[table_index]
            for j, value in enumerate(abbreviations[:len(table.columns)]):
                table.cell(0, j).text = value
                set_cell_nowrap(table.cell(0, j), True)
            format_table(table, font_size=6.5)

    # Replace Figure S1 with the profile diagnostic; retain monthly-load Figure S2.
    image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//a:blip")]
    if image_paragraphs:
        p = image_paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIG / "figure4_spatial_identifiability_profile.png"), width=Inches(6.5))
    try:
        cap = find_paragraph(doc, "Figure S1.")
        cap.text = "Figure S1. Profile spatial-identifiability audit after analytically profiling γ without its former bounds. The negligible improvement relative to k = 0 shows that the outlet series does not identify k under shared or assumed source calendars."
    except KeyError:
        pass
    try:
        cap = find_paragraph(doc, "Figure S2.")
        cap.text = "Figure S2. Observed 2022 monthly outlet loads and timestamp counts. Low late-flood-season coverage motivates the explicit S1–S4 annual-load scenarios."
    except KeyError:
        pass
    # Remove obsolete MCMC-rating and effective-contribution figures and captions.
    image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//a:blip")]
    for p in image_paragraphs[2:]:
        remove_paragraph(p)
    for prefix in ["Figure S3.", "Figure S4."]:
        try:
            remove_paragraph(find_paragraph(doc, prefix))
        except KeyError:
            pass

    # Add reproducibility section before the first figure.
    fig_anchor = find_paragraph(doc, "Figure S1.")
    p = simple_insert_after(
        find_paragraph(doc, "Table S35."),
        "Text S8. Reproducibility package",
    )
    p.runs[0].bold = True
    p2 = simple_insert_after(
        p,
        "Repository: https://github.com/Shifa-Zhong/MC-SIRC; stable baseline commit: "
        "126fa0a789751ecf0850b4971244f26cd404a880. Required input structure is data/raw/{monitor.xlsx, rain.xlsx, data(1).xlsx} "
        "with generated files under data/processed. Priors and bounds are listed in Table S4; MAP/MCMC/Monte Carlo and optimization seeds are 42. "
        "Main diagnostics are reproduced with python scripts/analysis/revision3_diagnostics.py and python scripts/analysis/spatial_identifiability_unbounded.py; "
        "figures are reproduced with python scripts/reporting/generate_revision3_figures.py. An anonymized example package mirrors the restricted input schema.",
    )
    p2.paragraph_format.line_spacing = 1.0

    # Ensure Table S1 title is bold and all captions are kept with the following table.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Table S"):
            for run in paragraph.runs:
                run.bold = True
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.line_spacing = 1.0

    out = PAPER / "SI_Final.docx"
    doc.save(out)
    shutil.copy2(out, ARCHIVE_DUP / "SI_Revised_20260824.docx")
    return out


def main():
    print(build_manuscript())
    build_abstract()
    print(build_si())


if __name__ == "__main__":
    main()
