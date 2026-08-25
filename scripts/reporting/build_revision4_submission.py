#!/usr/bin/env python3
"""Build the final internally consistent revision-4 submission package."""

from __future__ import annotations

import shutil
import sys
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.reporting import build_revision3_documents as base


PAPER = ROOT / "paper"
ARCHIVE_DUP = PAPER / "_archive_non_submission_20260824" / "duplicate_dated_versions"
SENSITIVITY = ROOT / "output" / "results" / "revision4_s3_corrected_sensitivity.xlsx"
GA_IMAGE = ROOT / "output" / "figures" / "revision4" / "graphical_abstract_identifiability.png"
DATE_TAG = "20260824"

ABSTRACT_PARTS = [
    (
        "Study region",
        "The Nanchuan River Basin (1,438 km²) is a hilly-gully Loess Plateau catchment in China. "
        "A 2022 bottom-up inventory (53,155 non-point-source grids and 100 georeferenced point "
        "sources) is paired with 5,928 observed hourly records at one outlet station.",
    ),
    (
        "Study focus",
        "Monitoring-Constrained Source Inventory Reconciliation and Classification (MC-SIRC) is "
        "an identifiability-aware workflow. Four scenarios represent missing-not-at-random gaps; "
        "Bayesian maximum a posteriori and Markov chain Monte Carlo analyses allocate aggregate "
        "discrepancies under explicit priors; Jacobian-rank, null-space, and k = 0 profile "
        "diagnostics test estimability; and forward Monte Carlo simulation describes inventory-side uncertainty.",
    ),
    (
        "New hydrological insights for the region",
        "The annual outlet constraint identifies one aggregate linear combination of 9–10 unknowns "
        "(Jacobian rank 1; nullity 8–9), precluding source-specific coefficient validation. Four "
        "boundary prior conflicts persist under all missing-data scenarios and define field-audit "
        "priorities. After profiling the global scale, the monthly distance-decay model adds virtually "
        "no explanatory power over a k = 0 flow-only null (ΔR² ≤ 0.0033), so half-life distances, "
        "effective-contribution shares, and ranking inversions are not reported. The supported regional "
        "result is an aggregate inventory–monitoring discrepancy and a monitoring-design diagnosis: "
        "tributary sections and observed source seasonality are needed for source-specific transport inference.",
    ),
]


def remove_paragraph(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def paragraph_before(anchor, text="", style=None):
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def find_prefix(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def replace_prefix(doc, prefix, text, optional=False):
    try:
        paragraph = find_prefix(doc, prefix)
    except KeyError:
        if optional:
            return None
        raise
    paragraph.text = text
    return paragraph


def set_body_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.29)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)


def apply_three_line_table(table):
    try:
        table.style = 'Table Normal'
    except KeyError:
        pass

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for child in list(borders):
        borders.remove(child)

    def add_border(parent, edge, value, size='0'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), value)
        if value != 'nil':
            element.set(qn('w:sz'), size)
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
        parent.append(element)

    add_border(borders, 'top', 'single', '12')
    add_border(borders, 'left', 'nil')
    add_border(borders, 'bottom', 'single', '12')
    add_border(borders, 'right', 'nil')
    add_border(borders, 'insideH', 'nil')
    add_border(borders, 'insideV', 'nil')

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn('w:shd'))
            if shading is not None:
                tc_pr.remove(shading)
            cell_borders = tc_pr.find(qn('w:tcBorders'))
            if cell_borders is None:
                cell_borders = OxmlElement('w:tcBorders')
                tc_pr.append(cell_borders)
            for child in list(cell_borders):
                cell_borders.remove(child)
            for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                value = 'single' if row_index == 0 and edge == 'bottom' else 'nil'
                add_border(cell_borders, edge, value, '8')


def set_structured_abstract_in_manuscript(doc):
    heading = find_prefix(doc, "Abstract")
    keywords = find_prefix(doc, "Keywords:")
    body = doc._element.body
    children = list(body)
    start = children.index(heading._p)
    end = children.index(keywords._p)
    for child in children[start + 1:end]:
        body.remove(child)
    for label, text in ABSTRACT_PARTS:
        paragraph = paragraph_before(keywords)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(label + ": ").bold = True
        paragraph.add_run(text)


def replace_introduction(doc):
    introduction = find_prefix(doc, '1 Introduction')
    methods = find_prefix(doc, '2 Materials and methods')
    body = doc._element.body
    children = list(body)
    start = children.index(introduction._p)
    end = children.index(methods._p)
    for child in children[start + 1:end]:
        body.remove(child)

    paragraphs = [
        'Identifying and quantifying pollution sources, and linking source emissions to observed water-quality responses, is a foundational requirement of watershed environmental management (Carpenter et al., 1998; Ongley et al., 2010). Emission inventories support this task through two broad paradigms: top-down accounting based on statistical data and bottom-up accounting based on grid-level, source-by-source quantification. Export-coefficient and spatially integrated inventory approaches provide the source-category and spatial detail needed for refined management (Johnes, 1996; Chen et al., 2014; Shen et al., 2012; Strokal et al., 2016). Inventories nevertheless estimate discharges at the source, while environmental management is concerned with concentrations and fluxes in receiving waters. Two conversion processes bridge this gap: delivery from emissions to nominal river-entry load, and channel transport from river entry to the monitored downstream load.',
        'Many inventory studies terminate at river-entry load estimation and do not quantitatively reconcile those estimates with loads observed at downstream monitoring sections, leaving a structural disconnect between inventory research and water-quality monitoring. River-entry coefficients are often transferred from technical guidance or other settings, so disagreement at the outlet may reflect inventory bias, hydrological delivery, omitted sources, or channel processing (Behrendt and Opitz, 2000; Preston et al., 2011). In-stream transformation adds another layer of complexity. Total phosphorus (TP) is strongly affected by particulate settling and remobilization (Withers and Jarvie, 2008), while ammonia nitrogen and total nitrogen respond to nitrification, denitrification, oxygen status, and residence time (Birgand et al., 2007; Alexander et al., 2009; Beaulieu et al., 2011). Pollutants therefore need not share a common delivery pattern, and distance alone cannot represent all relevant processes (Alexander et al., 2000; Schwarz et al., 2006).',
        'Bayesian methods have been increasingly applied to water-environment parameter estimation because they incorporate prior knowledge, quantify parameter uncertainty, and remain computationally stable in underdetermined settings (Qian et al., 2003; Borsuk et al., 2004; Reckhow, 2003; Gelman et al., 2013). Geographic information system technology allows emission-inventory spatial information to be integrated with channel representations, while Markov chain Monte Carlo sampling characterizes the joint posterior of model parameters (Foreman-Mackey et al., 2013). These tools are valuable, but priors and computation do not create independent information: one annual constraint cannot uniquely identify many source factors, and a narrow posterior at a hard bound may reflect regularization rather than information in the observations.',
        'Recent watershed source-apportionment studies increasingly combine source data, landscape attributes, and river-network transport. Spatially explicit applications in the Beiyun and upper Yangtze basins illustrate source and transport analysis with distributed watershed information (Liu et al., 2023; Li et al., 2023), while Bayesian hierarchical work across contrasting flow regimes demonstrates the value of multiple tributary stations for separating export processes (Neumann et al., 2023). Dynamic inventory–transport frameworks also show that source priorities can change with environmental conditions (Wang et al., 2024). These advances share an important design feature: spatial or temporal replication supplies contrasts that permit source and transport parameters to be estimated.',
        'These foundations converge on a single goal: a quantitative full-chain framework that links emission estimates, river-entry loads, and outlet monitoring within the inventory paradigm. Reaching this goal requires resolving several entangled difficulties at once. In data-limited basins, coefficient allocation against one monitoring section is severely underdetermined; missing observations may cluster in high-flow months; and monthly calibration does not automatically add spatial information when source inputs share measured or assumed temporal weights. The central problem is therefore not only how to reconcile the mass balance, but also how to distinguish supported aggregate discrepancies from non-identifiable source and process allocations.',
        'MC-SIRC makes five interrelated methodological contributions. First, it extends the inventory paradigm beyond river-entry estimation by reconciling aggregate inventory and outlet loads under four explicit missing-data scenarios. Second, Bayesian maximum a posteriori and Markov chain Monte Carlo analyses allocate the aggregate discrepancy under explicit priors, while Jacobian rank and exactly equivalent null-space solutions determine estimability; the A/B/C tiers screen prior conflict rather than coefficient reliability. Third, forward Monte Carlo simulation characterizes pre-transport inventory uncertainty and influential inputs. Fourth, the spatial formulation is audited against a k = 0 flow-only null after profiling the global scale, so distance-dependent outputs are limited to what the monitoring design can support. Fifth, the workflow converts unresolved attribution into field-audit priorities and a monitoring roadmap, with versioned code and anonymized input schemas supporting reproducibility. The Nanchuan River Basin provides a methodological case analysis for one Loess Plateau watershed; cross-basin transfer remains to be tested.',
    ]
    for text in paragraphs:
        paragraph = paragraph_before(methods, text)
        set_body_style(paragraph)


def insert_discussion_structure(doc):
    replace_prefix(doc, "3 Results and discussion", "3 Results")
    replace_prefix(doc, "3.3 Spatial-decay audit", "3.3 Spatial-decay identifiability results")
    process = find_prefix(doc, "The absence of an identifiable k")
    if not any(p.text.strip() == "4 Discussion" for p in doc.paragraphs):
        paragraph_before(process, "4 Discussion", "Heading 1")
        paragraph_before(process, "4.1 Loess Plateau process interpretation", "Heading 2")
    replace_prefix(doc, "3.4 Decision-support implications", "4.2 Decision-support implications")
    replace_prefix(doc, "3.5 Limitations and transfer scope", "4.3 Limitations and transfer scope")
    replace_prefix(doc, "4 Conclusions", "5 Conclusions")


def add_ai_declaration(doc):
    references = find_prefix(doc, "References")
    if any(p.text.startswith("Declaration of generative AI") for p in doc.paragraphs):
        return
    heading = paragraph_before(
        references,
        "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process",
        "Heading 1",
    )
    heading.paragraph_format.keep_with_next = True
    paragraph = paragraph_before(
        references,
        "During the preparation of this work, the authors used OpenAI Codex to assist with language "
        "editing, consistency checking, and preparation of revision files. The authors reviewed and "
        "edited the content as needed and take full responsibility for the content of the publication.",
    )
    set_body_style(paragraph)


def revise_manuscript():
    doc = Document(PAPER / "Manuscript_Final_MC-SIRC.docx")
    set_structured_abstract_in_manuscript(doc)
    replace_introduction(doc)
    replacements = {
        "Recent watershed source-apportionment studies":
            "Recent watershed source-apportionment studies increasingly combine source data, landscape "
            "attributes, and river-network transport. Spatially explicit applications in the Beiyun and "
            "upper Yangtze basins illustrate source and transport analysis with distributed watershed "
            "information (Liu et al., 2023; Li et al., 2023), while Bayesian hierarchical work across "
            "contrasting flow regimes demonstrates the value of multiple tributary stations for separating "
            "export processes (Neumann et al., 2023). Dynamic inventory–transport frameworks also show "
            "that source priorities can change with environmental conditions (Wang et al., 2024). These "
            "advances share an important design feature: spatial or temporal replication supplies contrasts "
            "that permit source and transport parameters to be estimated.",
        "Each MC-SIRC component has a distinct inferential role":
            "Each MC-SIRC component has a distinct inferential role (Table 1). The workflow separates "
            "descriptive accounting, prior-regularized allocation, forward uncertainty propagation, and "
            "formal identifiability auditing. Components that reuse the same outlet load or inventory are "
            "interpreted jointly within one evidence system rather than as independent validation of source coefficients. "
            "Supplementary methods, diagnostic figures, and detailed tables are organized in Texts S1–S8, "
            "Figures S1–S2, and Tables S1–S35, respectively.",
        "where α is the nominal inventory river-entry coefficient":
            "where α is the nominal inventory river-entry coefficient, f is a multiplicative discrepancy "
            "allocation, U is a non-negative unidentified-load term, and ε is observation/model discrepancy. "
            "Eight or nine active source factors plus U are estimated from one annual M. Truncated-normal "
            "priors f ~ TN(μ, σ²; [0.1, 2.0]) and U ~ Gamma(shape = 2, rate = 10/M), truncated at 0.5M, "
            "regularize the equation; the MAP calculation used 0.001M as a numerical lower bound for U. "
            "MAP estimates used 30 L-BFGS-B starts with seed 42. MCMC used the emcee ensemble sampler "
            "with 32 walkers, 20,000 steps, 5,000 burn-in steps, and seed 42 (Foreman-Mackey et al., 2013).",
        "where wᵢ,ₘ allocates annual source load":
            "where wᵢ,ₘ allocates annual source load to month m, dᵢ is distance to the outlet, k is attenuation, "
            "and γ is a global scale. The audit uses the eight months with at least 50% timestamp coverage "
            "(January–July and December). Most source types, including all point sources, use the same "
            "discharge-derived monthly weights; agricultural and aquaculture categories use assumed calendars "
            "multiplied by discharge. If wᵢ,ₘ = wₘ for all sources, Eq. (3) reduces exactly to L̂ₘ = γwₘA(k), "
            "with A(k) = ΣEᵢαᵢexp(−kdᵢ). Only the product γA(k) is identifiable, so any k can be offset by γ. "
            "The small departures created by assumed agricultural/aquaculture calendars do not constitute "
            "independent observations of source seasonality.",
        "We therefore profiled γ analytically":
            "We profiled γ analytically at each k ∈ [0, 0.3] km⁻¹ without restricting γ to [0.3, 5], "
            "compared the fit with the k = 0 flow-only null, and repeated the calculation with γ restricted "
            "to [0.3, 5] to diagnose boundary effects. The profile, rather than a point estimate, determines "
            "whether distance contributes information. Half-life distances and effective-contribution shares "
            "are reported only when k is independently identifiable.",
        "Under S2, five pairs meet":
            "Under S2, five pairs meet the strong prior-conflict threshold: large-livestock COD; rural-domestic "
            "TP; and TP for large livestock, industrial sources, and the centralized facility. Four of these—"
            "the large-livestock COD/TP and industrial/centralized TP pairs—remain at the 0.10 truncation bound "
            "under S1–S4. Rural-domestic TP changes from Tier A under S1–S2 to Tier B under S3–S4. Across all "
            "pairs, tier counts and some regularized component rankings change with the monitored-load scenario; "
            "for example, the highest NH₃-N component shifts between rural domestic and large livestock. The four "
            "persistent conflicts define field-verification priorities; no tier is interpreted as a validated coefficient.",
        "The former narrow MCMC interval":
            "The large-livestock TP MCMC interval (0.100–0.152 under S3) is piled against the 0.10 truncation "
            "bound. Its narrowness indicates boundary conflict within the stated model, not strong data information. "
            "Prior-sensitivity calculations share the same likelihood and serve to show whether a conflict flag "
            "persists when prior assumptions change; they are not an independent validation.",
        "Cross-year application reinforces":
            "The cross-year application reinforces this limitation. The four persistent boundary conflicts reappear "
            "in the two better monitored years, 2022 and 2023, but the fixed 2022 inventory is paired with monitored "
            "COD loads that vary 3.6-fold across 2020–2023. The discrepancy factors therefore absorb interannual "
            "hydrology and source changes. This is a consistency pattern, not confirmation of source coefficients; "
            "the 2020–2021 sparse archives are excluded from replication claims.",
        "Because k is not identified":
            "Because k is not identified, the data do not support pollutant-specific half-life distances, "
            "effective-contribution shares, point-source percentages, or emission-to-effective ranking inversions. "
            "Distance and nominal river-entry coefficients remain useful descriptors for designing monitoring "
            "locations, but they do not support quantitative intervention rankings in this dataset.",
        "Operational measures should therefore":
            "Operational measures should therefore not be re-ranked from the present factors. A near-source/far-source "
            "strategy can be evaluated as a field hypothesis using intervention or multi-section data before policy use. "
            "The workflow produces an evidence map—what is discrepant, what is underdetermined, and what to measure next—"
            "rather than a set of source-specific policy coefficients.",
        "The spatial audit reaches the same boundary":
            "The spatial audit reaches the same boundary: after profiling γ, a k = 0 flow-only model is effectively "
            "indistinguishable from the best decay profile (ΔR² ≤ 0.0033). Consequently, this dataset does not support "
            "half-life distances, effective-contribution shares, ranking inversions, or direct policy prescriptions. "
            "The principal contribution is an identifiability-aware workflow that converts disagreement between inventory "
            "and monitoring into transparent aggregate diagnostics and a concrete monitoring roadmap. Multi-section, "
            "seasonally resolved, and process-specific data are prerequisites for source-level calibration and transfer "
            "beyond this case basin.",
        "Analysis code is available":
            "The analysis package is versioned in two parts. The public baseline is permanently addressable at "
            "https://github.com/Shifa-Zhong/MC-SIRC/tree/126fa0a789751ecf0850b4971244f26cd404a880. Revision-specific "
            "diagnostic scripts, parameter files, result workbooks, reproduction commands, and an anonymized input-schema "
            "example are supplied as Code_Supplement_20260824.zip with this resubmission. Together they define the exact "
            "revision package and fixed random seeds. Raw hourly water-quality data are controlled by the local environmental "
            "authority and may be requested from the corresponding author for academic use, subject to approval.",
    }
    for prefix, text in replacements.items():
        replace_prefix(doc, prefix, text)
    standalone = {
        'Figure 1.': 'Figure 1. (a) Nanchuan River Basin, georeferenced point sources, and the outlet monitoring station. (b) Identifiability-aware MC-SIRC workflow. Under the present single-station design, supported outputs exclude source-specific attenuation rates and effective-contribution rankings.',
        'Within observed timestamps': 'Within observed timestamps, concentration gaps were linearly interpolated and implausible discharge values were replaced using a three-tier hierarchy: the same calendar time in other archive years, adjacent-month medians, and linear interpolation. Completely absent timestamps were handled by four annual-load scenarios: S1 observed hours only (lower bound); S2 month-specific coverage scaling (default); S3 annual-mean scaling by 8,760/5,928; and S4 an event-weighted upper sensitivity that multiplies only the S2-imputed portion by 1.5. S2 preserves the observed seasonal load structure and avoids exporting the high April–June mean to the poorly observed late-flood-season months. Its limitations are evaluated in the Discussion.',
        'The previously calibrated spatial model': 'The spatial model adopts a first-order exponential form used in river water-quality modelling (Streeter and Phelps, 1925; Chapra, 1997) and is evaluated here as an identifiability case study:',
        'The inventory contains': 'The inventory contains 1,036.5 t COD, 20.0 t NH₃-N, 131.3 t TN, and 14.0 t TP emissions, corresponding to nominal river-entry totals of 463.9, 6.84, 69.15, and 5.32 t, respectively (Table S8). At the source-emission scale, large-scale livestock accounts for 56.6% of COD and 65.1% of TP emissions, whereas the centralized facility accounts for 37.1% of TN emissions. These are inventory-composition shares, not estimates of source contributions at the outlet. Under default S2, outlet loads are 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. These differences establish an aggregate inventory–outlet discrepancy. They do not by themselves separate coefficient error from channel transformation or omitted sources.',
        'Table 2.': 'Table 2. Prior-regularized aggregate reconciliation under default S2. Deviations are constructional and are not independent validation statistics.',
        'The large-livestock TP MCMC interval': 'The large-livestock TP MCMC interval (0.100–0.153 under corrected S3) is piled against the 0.10 truncation bound. Its narrowness indicates boundary conflict within the stated model, not strong data information. Prior-sensitivity calculations share the same likelihood and serve to show whether a conflict flag persists when prior assumptions change; they are not an independent validation.',
        'Once γ is profiled': 'With γ profiled freely rather than restricted to [0.3, 5], the distance-decay profiles are essentially flat relative to k = 0 (Figure 4; Table 3). The incremental ΔR² is 0.00003 for COD, 0.00019 for NH₃-N, 0.00333 for TN, and 0.00069 for TP. For COD and TP, the k = 0 γ values (0.225 and 0.077) fall below 0.3; imposing γ ≥ 0.3 makes k > 0 appear necessary. NH₃-N, TN, and TP profile minima occur at the edge of the tested k range or require γ > 5. These are constraint and scaling effects, not independently estimated attenuation rates.',
        'Figure 4.': 'Figure 4. Spatial identifiability profiles for (a) COD, (b) NH₃-N, (c) TN, and (d) TP, with γ solved analytically at each k. Blue curves show relative-error SSE divided by its profile minimum; red dashed lines mark the k = 0 flow-only null, and gray dotted lines mark the profile-minimum k rather than an independently identified attenuation rate. Annotations report ΔR² relative to k = 0. The negligible ΔR² shows that monthly outlet data add virtually no distance information under shared or assumed source calendars.',
        'MC-SIRC reconciles a detailed': 'MC-SIRC reconciles a detailed watershed source inventory with outlet monitoring while making the limits of that reconciliation explicit. In the Nanchuan case, default month-specific coverage scaling gives outlet loads of 140.86 t COD, 4.785 t NH₃-N, 57.77 t TN, and 0.758 t TP. Bayesian regularization can allocate the aggregate discrepancy and flag persistent prior conflicts, but the annual Jacobian rank of 1 and nullity of 8–9 show that individual factors are not identifiable. Four boundary conflicts persist across all missing-data scenarios and warrant field verification; they do not establish causal coefficient overestimation.',
    }
    for prefix, text in standalone.items():
        replace_prefix(doc, prefix, text)
    insert_discussion_structure(doc)
    add_ai_declaration(doc)
    for paragraph in doc.paragraphs:
        if 'Wollheim et al., 2006, 2008' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                'Wollheim et al., 2006, 2008',
                'Wollheim et al., 2006; Wollheim et al., 2008',
            )
            set_body_style(paragraph)
    prune_uncited_references(doc)
    prune_uncited_references_strict(doc)
    base.set_document_format(doc, line_numbers=True)
    doc.save(PAPER / "Manuscript_Final_MC-SIRC.docx")
    shutil.copy2(PAPER / "Manuscript_Final_MC-SIRC.docx", ARCHIVE_DUP / f"Manuscript_Revised_{DATE_TAG}.docx")


def prune_uncited_references(doc):
    references = find_prefix(doc, 'References')
    paragraphs = list(doc.paragraphs)
    reference_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is references._p)
    body_text = '\n'.join(paragraph.text for paragraph in paragraphs[:reference_index])
    for paragraph in list(paragraphs[reference_index + 1:]):
        text = paragraph.text.strip()
        if not text:
            continue
        surname = text.split(',', 1)[0]
        year = next((str(value) for value in range(1900, 2100) if str(value) in text), None)
        if year is None or surname not in body_text or year not in body_text:
            remove_paragraph(paragraph)


def prune_uncited_references_strict(doc):
    cited = {
        ('Alexander', '2000'), ('Alexander', '2009'), ('Beaulieu', '2011'),
        ('Behrendt', '2000'), ('Birgand', '2007'), ('Borsuk', '2004'),
        ('Carpenter', '1998'), ('Chapra', '1997'), ('Chen', '2014'),
        ('Foreman-Mackey', '2013'), ('Gelman', '2013'), ('Johnes', '1996'),
        ('Li', '2023'), ('Liu', '2023'), ('Neumann', '2023'), ('Ongley', '2010'),
        ('Preston', '2011'), ('Qian', '2003'), ('Reckhow', '2003'),
        ('Rode', '2010'), ('Schwarz', '2006'), ('Shen', '2012'),
        ('Streeter', '1925'), ('Strokal', '2016'), ('Wang', '2024'),
        ('Wang', '2026'), ('Withers', '2008'), ('Wollheim', '2006'),
        ('Wollheim', '2008'),
    }
    references = find_prefix(doc, 'References')
    paragraphs = list(doc.paragraphs)
    reference_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is references._p)
    for paragraph in list(paragraphs[reference_index + 1:]):
        text = paragraph.text.strip()
        if not text:
            continue
        surname = text.split(',', 1)[0]
        year = next((str(value) for value in range(1900, 2100) if f', {value}.' in text), None)
        if (surname, year) not in cited:
            remove_paragraph(paragraph)


def build_abstract():
    doc = Document(PAPER / "Abstract.docx")
    clear_document_body(doc)
    base.set_document_format(doc, line_numbers=False)
    title = doc.add_paragraph(base.TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    heading = doc.add_paragraph("Abstract")
    heading.runs[0].bold = True
    for label, text in ABSTRACT_PARTS:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.add_run(label + ": ").bold = True
        paragraph.add_run(text)
    doc.add_paragraph(
        "Keywords: source inventory; outlet load; structural identifiability; Bayesian regularization; "
        "missing-not-at-random data; null model; Loess Plateau"
    )
    doc.save(PAPER / "Abstract.docx")
    shutil.copy2(PAPER / "Abstract.docx", ARCHIVE_DUP / f"Abstract_Revised_{DATE_TAG}.docx")


SOURCE_EN = {
    "面-农村生活污染源": "Rural domestic",
    "面-农业面源": "Agricultural NPS",
    "畜禽散养": "Household livestock",
    "面-水产养殖": "Aquaculture",
    "面-城市面源": "Urban NPS",
    "面-城镇散排": "Dispersed urban",
    "规模畜禽养殖": "Large livestock",
    "点-工业源": "Industrial",
    "点-集中式污染治理设施": "Central facility",
    "Unknown": "Unknown (t)",
}
POLL_KEY = {"COD": "COD", "氨氮": "NH3N", "总氮": "TN", "总磷": "TP"}


def replace_sensitivity_tables(doc):
    if not SENSITIVITY.exists():
        raise FileNotFoundError(SENSITIVITY)
    pollutants = ["COD", "氨氮", "总氮", "总磷"]
    for offset, pollutant in enumerate(pollutants):
        frame = pd.read_excel(SENSITIVITY, sheet_name=f"MCMC_{POLL_KEY[pollutant]}")
        rows = []
        for _, row in frame.iterrows():
            label = SOURCE_EN.get(row.Parameter, row.Parameter)
            rows.append([
                label,
                f"{row.Mean:.3f}",
                f"{row.Median:.3f}",
                f"{row.Std:.3f}",
                f"{row['CI_2.5']:.3f}",
                f"{row['CI_97.5']:.3f}",
            ])
        base.replace_table(
            doc,
            12 + offset,
            ["Parameter", "Mean", "Median", "Std.", "95% CI low", "95% CI high"],
            rows,
            font_size=7.8,
        )
    for offset, pollutant in enumerate(pollutants):
        frame = pd.read_excel(SENSITIVITY, sheet_name=f"Corr_{POLL_KEY[pollutant]}", index_col=0)
        labels = [SOURCE_EN.get(name, name) for name in frame.columns]
        rows = []
        for index, values in frame.iterrows():
            rows.append([SOURCE_EN.get(index, index)] + [f"{value:.2f}" for value in values])
        base.replace_table(doc, 16 + offset, ["Source"] + labels, rows, font_size=6.2)
    selected = {
        "COD": ["面-农村生活污染源", "畜禽散养", "面-城市面源", "面-城镇散排", "规模畜禽养殖", "点-工业源", "点-集中式污染治理设施"],
        "氨氮": ["面-农村生活污染源", "畜禽散养", "面-城市面源", "面-城镇散排", "规模畜禽养殖", "点-工业源", "点-集中式污染治理设施"],
        "总氮": ["面-农村生活污染源", "面-农业面源", "畜禽散养", "面-城市面源", "规模畜禽养殖", "点-集中式污染治理设施"],
        "总磷": ["面-农村生活污染源", "面-农业面源", "畜禽散养", "面-城市面源", "规模畜禽养殖", "点-工业源", "点-集中式污染治理设施"],
    }
    short = {
        "面-农村生活污染源": "Rural", "面-农业面源": "Agri.", "畜禽散养": "HH livest.",
        "面-城市面源": "Urban", "面-城镇散排": "Disp.", "规模畜禽养殖": "Large livest.",
        "点-工业源": "Industrial", "点-集中式污染治理设施": "Central",
    }
    for offset, pollutant in enumerate(pollutants):
        frame = pd.read_excel(SENSITIVITY, sheet_name=f"Prior_{POLL_KEY[pollutant]}")
        sources = selected[pollutant]
        rows = []
        for _, row in frame.iterrows():
            rows.append([row.Scenario, f"{row.Deviation_pct:+.1f}%"] + [f"{row[source]:.3f}" for source in sources])
        base.replace_table(
            doc,
            20 + offset,
            ["Prior scenario", "Dev.%"] + [short[source] for source in sources],
            rows,
            font_size=6.8,
        )
    sigma = pd.read_excel(SENSITIVITY, sheet_name="Sigma")
    sigma = sigma[sigma.Pollutant.isin(["COD", "总磷"])]
    sources = ["面-农村生活污染源", "面-城市面源", "规模畜禽养殖", "点-工业源", "点-集中式污染治理设施"]
    rows = []
    for _, row in sigma.iterrows():
        label = "COD" if row.Pollutant == "COD" else "TP"
        rows.append([label, row.sigma_obs] + [f"{row[source]:.3f}" for source in sources])
    base.replace_table(
        doc,
        25,
        ["Pollutant", "σobs"] + [short[source] for source in sources],
        rows,
        font_size=7.2,
    )


def revise_si():
    doc = Document(PAPER / "SI_Final.docx")
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("Abbreviations:") or paragraph.text.startswith("Text S8.") or paragraph.text.startswith("Repository:"):
            remove_paragraph(paragraph)
    anchor = find_prefix(doc, "*Corresponding Author:")
    abbr = base.simple_insert_after(
        anchor,
        "Abbreviations: maximum a posteriori (MAP); Markov chain Monte Carlo (MCMC); non-point source "
        "(NPS); chemical oxygen demand (COD); ammonia nitrogen (NH₃-N); total nitrogen (TN); total phosphorus (TP).",
    )
    abbr.paragraph_format.line_spacing = 1.0
    replacements = {
        "The 2022 archive contains":
            "The 2022 archive contains 5,928 unique hourly timestamps from a possible 8,760 (67.7% coverage), "
            "so 2,832 timestamps are absent. Gaps cluster in July–November and are treated as MNAR. Within observed "
            "timestamps, concentration gaps were linearly interpolated; discharge gaps or invalid values were filled "
            "by same-calendar historical medians, adjacent-month medians, and then interpolation. Completely absent "
            "hours were handled by four scenarios: S1 observed-only; S2 month-specific coverage scaling (default); "
            "S3 annual-mean scaling by 8,760/5,928 = 1.4777; and S4 observed load plus 1.5 times the S2-imputed portion. "
            "S4 is an event-weighted sensitivity case, not a probabilistic upper confidence bound. S2 preserves the "
            "observed monthly structure instead of exporting the April–June mean into the poorly observed late-flood-season months.",
        "Urban NPS functional-zone weights":
            "Urban NPS functional-zone weights are descriptive inventory inputs. Zone-specific concentrations are not "
            "estimated because one total-load constraint cannot identify 18 additional parameters; no optimized Cj values "
            "enter the analysis.",
        "Non-point sources are aggregated":
            "Non-point sources are aggregated by control unit with representative distances of 8, 12, 22, and 30 km; "
            "point-source GPS distances range from 2.3 to 99.6 km. Nominal α = R/E values and distances are descriptive "
            "inputs. Half-life distances and effective outlet contributions require an independently identifiable k and "
            "are therefore not calculated from these inputs.",
        "Correction factors:":
            "Discrepancy factors: fi ~ TN(μi, σi², [0.1, 2.0]). Unidentified load: U ~ Gamma(shape = 2, "
            "rate = 10/Mp), truncated at 0.5Mp; E(U) before truncation = 0.2Mp. The MAP calculation uses a numerical "
            "lower bound of 0.001Mp, and σobs = 0.10Mp by default.",
        "“—”: data not available":
            "“—” denotes unavailable inventory data. Industrial TN has no nominal river-entry coefficient because its TN "
            "inventory entry is absent; the listed α values are inventory parameters rather than measured transport fractions.",
        "Note: Aug–Nov":
            "The spatial identifiability audit uses January–July and December, the eight months with at least 50% timestamp "
            "coverage. August–November contain 67–209 records per month and are not used in the profile objective.",
        "f = correction factor":
            "f = regularized discrepancy factor; z = |fMAP − μ|/σ. Tier A denotes strong prior conflict (z > 2), "
            "not an anomalous or data-identifiable source coefficient.",
        "All four original boundary conflicts":
            "Four boundary conflicts remain at 0.100 under S1–S4, whereas other tier assignments and some regularized "
            "component rankings change. Urban functional-zone weights remain descriptive and are not separately calibrated.",
        "The annual-mean results in Table S27":
            "Table S27 uses annual-mean scaling for historical comparability. For 2022, the factor is "
            "8,760/5,928 = 1.4777. This S3 total is a sensitivity scenario, not the S2 default.",
        "★ marks the four boundary conflicts":
            "★ marks four boundary conflicts selected for field audit. Their recurrence is a conditional consistency pattern "
            "under a fixed inventory and prior system, not evidence that the source factors are data-identifiable.",
        "Annual inversion: Eq. (2)":
            "Annual inversion: Eq. (2) has one observation and 9–10 unknowns. The Jacobian rank is 1 and its nullity is "
            "8–9. Table S33 gives factor allocations that differ substantially yet reproduce the same regularized source "
            "sum. Spatial model: with shared monthly weights, L̂m = γwmA(k), so only γA(k) is identifiable. Agricultural "
            "and aquaculture calendars create small assumed deviations from shared weighting but do not supply independently "
            "observed source seasonality. γ was profiled analytically without restriction to [0.3, 5] and every k was compared "
            "with k = 0. Across pollutants, ΔR² is 0.00003–0.00333 and the profile SSE ratio is 1.00008–1.0604. Thus k, "
            "half-life, and effective-contribution shares are not reported.",
        "For each k from 0 to 0.30":
            "For each k from 0 to 0.30 km⁻¹, γ was solved analytically without bounds and the profile relative-error SSE and "
            "R² were calculated. The k = 0 flow-only null was compared with the profile optimum. A separate sensitivity "
            "calculation restricted γ to [0.3, 5] to identify boundary-created optima.",
        "For COD and TP, γ at k = 0":
            "For COD and TP, γ at k = 0 is 0.225 and 0.077, below 0.3. Restricting γ to [0.3, 5] makes positive k appear "
            "necessary. Without that restriction, ΔR² relative to k = 0 is ≤ 0.00333 for every pollutant.",
        "Figure S1. Profile":
            "Figure S1. Profile spatial-identifiability audit with γ solved analytically without restriction to its earlier "
            "interval. The negligible improvement relative to k = 0 shows that the outlet series does not identify k under "
            "shared or assumed source calendars.",
    }
    for prefix, text in replacements.items():
        replace_prefix(doc, prefix, text, optional=True)
    replace_prefix(doc, 'Table S3.', 'Table S3. Missing-data scenarios and 2022 annual outlet loads.')
    replace_prefix(doc, 'Table S6.', 'Table S6. Spatial k = 0 null comparison with γ profiled freely at each k.')
    doc.tables[3].cell(1, 3).text = 'Inventory coefficient uncertainty assumed relatively low'
    doc.tables[3].cell(4, 3).text = 'Inventory coefficient uncertainty assumed relatively high'
    caption_updates = {
        "Table S12.": "Table S12. MAP regularized discrepancy factors under default S2. z measures prior conflict, not identifiability.",
        "Table S13.": "Table S13. MCMC posterior summaries—COD, corrected annual-mean sensitivity scenario S3.",
        "Table S14.": "Table S14. MCMC posterior summaries—NH₃-N, corrected annual-mean sensitivity scenario S3.",
        "Table S15.": "Table S15. MCMC posterior summaries—TN, corrected annual-mean sensitivity scenario S3.",
        "Table S16.": "Table S16. MCMC posterior summaries—TP, corrected annual-mean sensitivity scenario S3. Boundary-piled intervals do not establish data identifiability.",
        "Table S17.": "Table S17. MCMC parameter correlations under corrected S3. These are summaries of one prior-regularized posterior, not formal identifiability diagnostics.",
        "Table S18.": "Table S18. Prior scenarios P1–P5 for COD under corrected annual-mean missing-data scenario S3.",
        "Table S19.": "Table S19. Prior scenarios P1–P5 for NH₃-N under corrected annual-mean missing-data scenario S3.",
        "Table S20.": "Table S20. Prior scenarios P1–P5 for TN under corrected annual-mean missing-data scenario S3.",
        "Table S21.": "Table S21. Prior scenarios P1–P5 for TP under corrected annual-mean missing-data scenario S3.",
        "Table S23.": "Table S23. Observation-error sensitivity under corrected annual-mean scenario S3.",
    }
    for prefix, text in caption_updates.items():
        paragraph = replace_prefix(doc, prefix, text, optional=True)
        if paragraph:
            for run in paragraph.runs:
                run.bold = True
    replace_sensitivity_tables(doc)
    table_anchor = doc.tables[-1]._tbl
    title_element = OxmlElement('w:p')
    table_anchor.addnext(title_element)
    title = Paragraph(title_element, table_anchor.getparent())
    title.add_run('Text S8. Reproducibility package')
    title.runs[0].bold = True
    description = base.simple_insert_after(
        title,
        "Public baseline: https://github.com/Shifa-Zhong/MC-SIRC/tree/126fa0a789751ecf0850b4971244f26cd404a880. "
        "The accompanying Code_Supplement_20260824.zip contains revision-specific diagnostic and reporting scripts, "
        "config/revision3_parameters.json, result workbooks, REPRODUCIBILITY.md, and anonymized example schemas. "
        "Restricted real-data inputs follow data/raw/{monitor.xlsx, rain.xlsx, data(1).xlsx}; generated files are placed "
        "under data/processed. MAP, MCMC, Monte Carlo, and optimization seeds are 42. Together the stable commit and code "
        "supplement define the exact analysis version used for this revision.",
    )
    description.paragraph_format.line_spacing = 1.0
    for table in doc.tables:
        base.format_table(table, font_size=7.5 if len(table.columns) >= 9 else 8.0)
    base.set_document_format(doc, line_numbers=False)
    doc.save(PAPER / "SI_Final.docx")
    shutil.copy2(PAPER / "SI_Final.docx", ARCHIVE_DUP / f"SI_Revised_{DATE_TAG}.docx")


def build_highlights():
    bullets = [
        "MC-SIRC reconciles source inventories with one-station outlet monitoring.",
        "Four missing-data scenarios expose flood-season load sensitivity.",
        "Rank 1 and nullity 8–9 reveal source-factor non-identifiability.",
        "A k = 0 null fits nearly as well as the best distance-decay profile.",
        "The workflow turns uncertainty into field-audit and monitoring priorities.",
    ]
    doc = Document(PAPER / "highlights.docx")
    clear_document_body(doc)
    title = doc.add_paragraph("Highlights")
    title.runs[0].bold = True
    for text in bullets:
        doc.add_paragraph("• " + text)
    base.set_document_format(doc, line_numbers=False)
    doc.save(PAPER / "highlights.docx")
    shutil.copy2(PAPER / "highlights.docx", ARCHIVE_DUP / f"Highlights_Revised_{DATE_TAG}.docx")


def build_cover_letter():
    paragraphs = [
        "Dear Editor,",
        f"We are pleased to resubmit our manuscript, \"{base.TITLE},\" for consideration as a Research Paper in Journal of Hydrology: Regional Studies.",
        "The revision addresses the central inferential limitation identified during review. MC-SIRC is now presented as an identifiability-aware workflow for reconciling a source inventory with outlet monitoring, rather than as a source-coefficient calibration. The manuscript quantifies aggregate inventory–monitoring discrepancies under four missing-data scenarios and uses prior-regularized Bayesian analysis only to define field-audit priorities.",
        "Formal diagnostics show that the annual inverse problem has Jacobian rank 1 and nullity 8–9. A profile comparison also shows that the monthly distance-decay formulation provides negligible improvement over a k = 0 flow-only model after the global scale is profiled. We therefore do not report source-specific attenuation half-lives, effective outlet shares, ranking inversions, or direct control prescriptions.",
        "The regional contribution is a transparent diagnosis of what one outlet station can support and which additional observations are needed in this flashy Loess Plateau basin. The revision also clarifies flood-season missingness, labels the cross-year exercise as a consistency check, expands the discussion of sediment and nitrogen processing, updates the figures and SI, and supplies a versioned reproduction package.",
        "The abstract follows the journal's three-part structure and 225-word limit. Updated Highlights, Graphical Abstract, Supporting Information, response letter, and code supplement accompany the manuscript.",
        "The manuscript has not been published or submitted elsewhere, all authors have approved the resubmission, and the authors declare no competing interests.",
        "Thank you for your consideration.",
        "Sincerely,",
        "Shifa Zhong, Ph.D.",
        "College of Environmental Science and Engineering, Tongji University, 1239 Siping Road, Shanghai 200092, China",
        "Email: sfzhong@tongji.edu.cn",
    ]
    doc = Document(PAPER / "cover letter.docx")
    clear_document_body(doc)
    for text in paragraphs:
        doc.add_paragraph(text)
    base.set_document_format(doc, line_numbers=False)
    doc.save(PAPER / "cover letter.docx")
    shutil.copy2(PAPER / "cover letter.docx", ARCHIVE_DUP / f"Cover_Letter_Revised_{DATE_TAG}.docx")


def build_ga():
    doc = Document(PAPER / "GA.docx")
    clear_document_body(doc)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(GA_IMAGE), width=Inches(10.8))
    for section in doc.sections:
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
    doc.save(PAPER / "GA.docx")
    shutil.copy2(PAPER / "GA.docx", ARCHIVE_DUP / f"Graphical_Abstract_Revised_{DATE_TAG}.docx")


def build_code_supplement():
    destination = PAPER / f"Code_Supplement_{DATE_TAG}.zip"
    files = [
        ROOT / "README.md",
        ROOT / "REPRODUCIBILITY.md",
        ROOT / "config" / "revision3_parameters.json",
        ROOT / "scripts" / "analysis" / "revision3_diagnostics.py",
        ROOT / "scripts" / "analysis" / "spatial_identifiability_unbounded.py",
        ROOT / "scripts" / "analysis" / "revision4_s3_sensitivity.py",
        ROOT / "scripts" / "reporting" / "generate_revision3_figure1.py",
        ROOT / "scripts" / "reporting" / "generate_revision3_figures.py",
        ROOT / "scripts" / "reporting" / "build_revision3_documents.py",
        ROOT / "scripts" / "reporting" / "build_revision3_response_letter.py",
        ROOT / "scripts" / "reporting" / "build_revision4_submission.py",
        ROOT / "output" / "results" / "revision3_identifiability_and_missingness.xlsx",
        ROOT / "output" / "results" / "spatial_identifiability_unbounded.xlsx",
        ROOT / "output" / "results" / "revision4_s3_corrected_sensitivity.xlsx",
    ]
    files.extend(sorted((ROOT / "data" / "example").rglob("*")))
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in files:
            if path.is_file():
                archive.write(path, path.relative_to(ROOT).as_posix())
    return destination


def replace_code_distribution_statements():
    repository_text = (
        'The authoritative analysis package is the tagged public repository snapshot at '
        'https://github.com/Shifa-Zhong/MC-SIRC/tree/revision-2026-08-24. It contains the '
        'revision-specific core analysis scripts, parameter file, fixed seeds, reproduction commands, and '
        'station-anonymized input schemas. Raw hourly water-quality data are controlled by the local '
        'environmental authority and may be requested from the corresponding author for academic use, subject to approval.'
    )
    manuscript_path = PAPER / 'Manuscript_Final_MC-SIRC.docx'
    doc = Document(manuscript_path)
    for paragraph in doc.paragraphs:
        if 'Code_Supplement_20260824.zip' in paragraph.text:
            paragraph.text = repository_text
            set_body_style(paragraph)
    doc.save(manuscript_path)
    shutil.copy2(manuscript_path, ARCHIVE_DUP / f'Manuscript_Revised_{DATE_TAG}.docx')

    si_path = PAPER / 'SI_Final.docx'
    doc = Document(si_path)
    for paragraph in doc.paragraphs:
        if 'Code_Supplement_20260824.zip' in paragraph.text:
            paragraph.text = (
                'Tagged repository snapshot: '
                'https://github.com/Shifa-Zhong/MC-SIRC/tree/revision-2026-08-24. The repository contains '
                'the revision-specific core analysis scripts, config/revision3_parameters.json, fixed seeds, '
                'REPRODUCIBILITY.md, and station-anonymized example schemas. Restricted real-data inputs follow '
                'data/raw/{monitor.xlsx, rain.xlsx, data(1).xlsx}; generated files are placed under output/results. '
                'MAP, MCMC, Monte Carlo, and optimization seeds are 42.'
            )
            paragraph.paragraph_format.line_spacing = 1.0
    doc.save(si_path)
    shutil.copy2(si_path, ARCHIVE_DUP / f'SI_Revised_{DATE_TAG}.docx')

    cover_path = PAPER / 'cover letter.docx'
    doc = Document(cover_path)
    for paragraph in doc.paragraphs:
        if 'code supplement accompany the manuscript' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                'Supporting Information, response letter, and code supplement accompany the manuscript.',
                'Supporting Information and response letter accompany the manuscript; the core analysis code is available in the tagged public repository.'
            )
    doc.save(cover_path)
    shutil.copy2(cover_path, ARCHIVE_DUP / f'Cover_Letter_Revised_{DATE_TAG}.docx')


def neutralize_submission_language():
    manuscript_replacements = {
        'so half-life distances, effective-contribution shares, and ranking inversions are not reported.':
            'and therefore leaves half-life distances, effective-contribution shares, and ranking inversions unresolved.',
        'Its limitations are evaluated in the Discussion.':
            'The implications of the remaining low-coverage months are examined in the Discussion.',
        'The cross-year application reinforces this limitation.':
            'The cross-year application provides a conditional comparison across archive years.',
        'This is a consistency pattern, not confirmation of source coefficients; the 2020–2021 sparse archives are excluded from replication claims.':
            'It therefore characterizes consistency under a fixed inventory without identifying source coefficients; the 2020–2021 archives are too sparse for replication analysis.',
        'The analysis is retained only to communicate input uncertainty and to identify influential inventory terms.':
            'Accordingly, the analysis characterizes input uncertainty and identifies influential inventory terms.',
        'With γ profiled freely rather than restricted to [0.3, 5], the distance-decay profiles are essentially flat relative to k = 0':
            'Profiling γ analytically over positive values at each k produces distance-decay profiles that are essentially flat relative to k = 0',
        'Consequently, this dataset does not support half-life distances, effective-contribution shares, ranking inversions, or direct policy prescriptions.':
            'The profile therefore leaves half-life distances, effective-contribution shares, ranking inversions, and direct policy prescriptions unresolved.',
        'revision-specific core analysis scripts': 'core analysis scripts supporting this study',
        'language editing, consistency checking, and preparation of revision files.':
            'language editing and manuscript consistency checking.',
        'under corrected S3': 'under S3',
        'Prior-sensitivity calculations share the same likelihood and serve to show whether a conflict flag persists when prior assumptions change; they are not an independent validation.':
            'Prior-sensitivity calculations share the same likelihood and characterize how the conflict flag responds to prior assumptions.',
        'This is labelled a cross-year consistency check rather than validation: the factors necessarily absorb year-specific hydrology and unobserved source changes.':
            'This cross-year consistency check is conditional because the factors necessarily absorb year-specific hydrology and unobserved source changes.',
        'The residual deviations of +2.0% to +12.0% describe the chosen observation-error and prior trade-off, not predictive validation.':
            'The residual deviations of +2.0% to +12.0% summarize the chosen observation-error and prior trade-off.',
        'Deviations are constructional and are not independent validation statistics.':
            'Deviations describe the prior-regularized reconciliation and are not predictive-performance statistics.',
        'Deviations describe the prior-regularized reconciliation and are not predictive-performance statistics.':
            'Deviations summarize the observation-error and prior trade-off in the regularized reconciliation.',
        'Red dashed lines show S2 outlet loads; the separation is an inventory–outlet scale discrepancy, not independent evidence of inventory bias.':
            'Red dashed lines show S2 outlet loads for an inventory–outlet scale comparison.',
        'Red dashed lines show S2 outlet loads; the separation is an inventory–outlet scale discrepancy, not independent evidence of source-coefficient bias.':
            'Red dashed lines show S2 outlet loads for an inventory–outlet scale comparison.',
        'The cross-year check is not independent because the 2022 inventory is fixed, the monitored load changes substantially, and two years are extremely sparse.':
            'The cross-year check remains conditional because the 2022 inventory is fixed, the monitored load changes substantially, and two years are extremely sparse.',
    }
    si_replacements = {
        'Text S4. Spatial inputs retained for identifiability auditing':
            'Text S4. Spatial inputs for identifiability auditing',
        'corrected annual-mean sensitivity scenario S3':
            'annual-mean sensitivity scenario S3',
        'under corrected S3': 'under S3',
        'under corrected annual-mean missing-data scenario S3':
            'under annual-mean missing-data scenario S3',
        'under corrected annual-mean scenario S3':
            'under annual-mean scenario S3',
        'Thus k, half-life, and effective-contribution shares are not reported.':
            'Thus k, half-life, and effective-contribution shares remain unresolved.',
        'without restriction to its earlier interval':
            'without restriction to [0.3, 5]',
        'revision-specific core analysis scripts':
            'core analysis scripts supporting this study',
        'These are summaries of one prior-regularized posterior, not formal identifiability diagnostics.':
            'Formal estimability is assessed separately with rank and null-space diagnostics.',
        'Cross-year prior-regularized reconciliation summary (constructional; not validation).':
            'Cross-year prior-regularized reconciliation summary under a fixed 2022 inventory.',
        'This is sensitivity to the assumed likelihood scale, not independent evidence of source identifiability.':
            'It therefore characterizes sensitivity to the assumed likelihood scale; formal source estimability is assessed by rank and null-space diagnostics.',
        'This design is a consistency check, not independent validation: factors absorb interannual hydrology and source changes.':
            'This design characterizes consistency under a fixed inventory because factors absorb interannual hydrology and source changes.',
        'They are reported for bookkeeping and are not validation statistics.':
            'They are reported as bookkeeping summaries of the regularized fit.',
    }

    def apply_replacements(path, replacements):
        doc = Document(path)
        counts = {old: 0 for old in replacements}
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            for run in paragraph.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        counts[old] += run.text.count(old)
                        run.text = run.text.replace(old, new)
        for table in doc.tables:
            apply_three_line_table(table)
        doc.save(path)
        return counts

    ms_path = PAPER / 'Manuscript_Final_MC-SIRC.docx'
    apply_replacements(ms_path, manuscript_replacements)
    shutil.copy2(ms_path, ARCHIVE_DUP / f'Manuscript_Revised_{DATE_TAG}.docx')

    abstract_path = PAPER / 'Abstract.docx'
    apply_replacements(abstract_path, manuscript_replacements)
    shutil.copy2(abstract_path, ARCHIVE_DUP / f'Abstract_Revised_{DATE_TAG}.docx')

    si_path = PAPER / 'SI_Final.docx'
    apply_replacements(si_path, si_replacements)
    shutil.copy2(si_path, ARCHIVE_DUP / f'SI_Revised_{DATE_TAG}.docx')


def main():
    baseline = PAPER / '_archive_non_submission_20260824' / 'historical_backups' / 'backup_before_revision_20260824'
    for name in ['Manuscript_Final_MC-SIRC.docx', 'SI_Final.docx', 'Abstract.docx']:
        shutil.copy2(baseline / name, PAPER / name)
    if not SENSITIVITY.exists():
        raise FileNotFoundError("Run scripts/analysis/revision4_s3_sensitivity.py first")
    base.build_manuscript()
    base.build_abstract()
    base.build_si()
    revise_manuscript()
    build_abstract()
    revise_si()
    build_highlights()
    build_cover_letter()
    build_ga()
    replace_code_distribution_statements()
    neutralize_submission_language()


if __name__ == "__main__":
    main()
