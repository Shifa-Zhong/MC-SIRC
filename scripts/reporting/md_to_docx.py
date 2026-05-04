#!/usr/bin/env python3
"""
将论文 Markdown 转换为格式化的 Word docx 文件
支持：标题层级、正文段落、表格、粗体/斜体、公式占位、有序/无序列表
"""

import sys
import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ── 样式配置 ──────────────────────────────────────────────────────────

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_SIZE_BODY = Pt(10.5)      # 五号
FONT_SIZE_TITLE = Pt(16)       # 小二
FONT_SIZE_H2 = Pt(14)          # 四号
FONT_SIZE_H3 = Pt(12)          # 小四
FONT_SIZE_H4 = Pt(10.5)        # 五号加粗
FONT_SIZE_TABLE = Pt(9)        # 小五
LINE_SPACING = Pt(20)          # 行距


def set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = size
    run.font.name = FONT_EN
    run.font.bold = bold
    run.font.italic = italic
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=Pt(0), after=Pt(0), line=LINE_SPACING):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line


def add_formatted_text(paragraph, text, size=FONT_SIZE_BODY):
    """解析 Markdown 行内格式（**粗体**、*斜体*、$公式$）并添加 runs"""
    # 分割标记
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|\$\$.*?\$\$|\$[^$]+?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True)
        elif part.startswith('$$') and part.endswith('$$'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, italic=True, color=RGBColor(0x66, 0x66, 0x66))
        elif part.startswith('$') and part.endswith('$'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=RGBColor(0x66, 0x66, 0x66))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def parse_table(lines, start_idx):
    """从 Markdown 行列表中解析表格，返回 (rows, end_idx)"""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # 跳过分隔行（只含 |、-、:、空格）
        cleaned = line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
        if not cleaned:
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """添加格式化表格"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, cell_text.strip(), size=FONT_SIZE_TABLE)

    # 首行加粗
    if rows:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

    # 添加表后空行
    doc.add_paragraph()


def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    i = 0
    prev_i = -1
    while i < len(lines):
        if i == prev_i:
            print(f'  WARNING: stuck at line {i}, skipping')
            i += 1
            continue
        prev_i = i
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线
        if stripped == '---':
            i += 1
            continue

        # 标题
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            title_text = stripped.lstrip('#').strip()

            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, before=Pt(12), after=Pt(12))
                run = p.add_run(title_text)
                set_run_font(run, size=FONT_SIZE_TITLE, bold=True)
            elif level == 2:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=Pt(12), after=Pt(6))
                run = p.add_run(title_text)
                set_run_font(run, size=FONT_SIZE_H2, bold=True)
            elif level == 3:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=Pt(8), after=Pt(4))
                run = p.add_run(title_text)
                set_run_font(run, size=FONT_SIZE_H3, bold=True)
            elif level >= 4:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=Pt(6), after=Pt(3))
                run = p.add_run(title_text)
                set_run_font(run, size=FONT_SIZE_H4, bold=True)
            i += 1
            continue

        # 表格标题（**表N ...**）
        if stripped.startswith('**表') and stripped.endswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=Pt(6), after=Pt(3))
            text = stripped[2:-2]
            run = p.add_run(text)
            set_run_font(run, size=FONT_SIZE_BODY, bold=True)
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # 无序列表
        if stripped.startswith('- '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=Pt(2), after=Pt(2))
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            text = '  ' + stripped[2:]
            add_formatted_text(p, text)
            i += 1
            continue

        # 公式块（单行 $$...$$ 或多行）
        if stripped.startswith('$$'):
            if stripped.endswith('$$') and len(stripped) > 4:
                # 单行公式
                formula_text = stripped[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, before=Pt(6), after=Pt(6))
                run = p.add_run(formula_text)
                set_run_font(run, size=FONT_SIZE_BODY, italic=True)
                i += 1
            else:
                # 多行公式
                formula_parts = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    formula_parts.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    i += 1
                formula_text = ' '.join(l for l in formula_parts if l)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, before=Pt(6), after=Pt(6))
                run = p.add_run(formula_text)
                set_run_font(run, size=FONT_SIZE_BODY, italic=True)
            continue

        # 普通段落
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=Pt(0), after=Pt(3))
        p.paragraph_format.first_line_indent = Cm(0.75)
        add_formatted_text(p, stripped)
        i += 1

    doc.save(str(docx_path))
    n_para = len(doc.paragraphs)
    n_table = len(doc.tables)
    try:
        print(f'Saved: {docx_path}')
        print(f'Paragraphs: {n_para}')
        print(f'Tables: {n_table}')
    except (ValueError, OSError):
        pass  # stdout may be closed in some contexts


def main():
    base = Path('D:/shanxi/paper')
    md_file = base / '论文_全链条闭合分析.md'
    docx_file = base / '论文2.docx'

    print(f'Converting: {md_file.name} -> {docx_file.name}')
    convert_md_to_docx(md_file, docx_file)
    print('Done.')


if __name__ == '__main__':
    main()
