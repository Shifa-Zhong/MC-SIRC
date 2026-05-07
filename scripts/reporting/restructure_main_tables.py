#!/usr/bin/env python3
"""
将正文从 8 个表格简化为 2 个 (Tables 1, 2 = 原 Table 4, Table 7),
余表移到 SI 或删除 (信息已在图里), 同步更新所有正文/SI 引用。
另外重新嵌入 5-panel Figure 3。
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from pathlib import Path
import shutil

ROOT = Path('D:/shanxi')
MS = ROOT/'paper/Manuscript_Final_MC-SIRC.docx'
SI = ROOT/'paper/SI_Final.docx'

# 备份当前
shutil.copy2(MS, MS.with_suffix('.docx.tab_bak'))
shutil.copy2(SI, SI.with_suffix('.docx.tab_bak'))
print("备份: *.docx.tab_bak")

# ============================================================
# 第一步: 准备移到 SI 的 3 张原表数据 (Tables 1, 2, 3)
# 从当前 manuscript 提取
# ============================================================
print("\n提取要移到 SI 的表格数据...")
ms = docx.Document(str(MS))

# Table indices (0-based):
# 0: Table 1 (data overview) → move to SI
# 1: Table 2 (emission/river-entry kg) → move to SI
# 2: Table 3 (closure coefficients) → move to SI
# 3: Table 4 (Bayesian summary) → KEEP (rename to Table 1)
# 4: Table 5 (MCMC) → DELETE (redundant with Fig 2b, Fig S2)
# 5: Table 6 (Sensitivity) → DELETE (redundant with Fig 2d, SI S20-S23)
# 6: Table 7 (Spatial decay) → KEEP (rename to Table 2)
# 7: Table 8 (Emis vs Eff ranking) → DELETE (redundant with Fig 4 b/c/d)

# 先把 Tables 1, 2, 3 的内容提取出来 (我们要 deep copy 整个 table element)
tables_to_move = []
for idx, label in [(0, 'Table S2 (Data overview)'),
                    (1, 'Table S3 (Emission and river-entry load by source, kg)'),
                    (2, 'Table S4 (Full-chain coefficients: emission → river-entry → monitored)')]:
    t = ms.tables[idx]
    # Deep copy the underlying XML
    tbl_xml = deepcopy(t._element)
    tables_to_move.append((label, tbl_xml))
    print(f"  提取 Table {idx+1}: {label}")

# ============================================================
# 第二步: 把 Tables 1, 2, 3, 5, 6, 8 从主文删掉 (按从后往前删)
# ============================================================
print("\n从主文删除 6 张表 (Tables 1, 2, 3, 5, 6, 8)...")

# Build index list of tables to delete in REVERSE order (so indices stay valid)
to_delete_indices = [7, 5, 4, 2, 1, 0]  # Table 8, 6, 5, 3, 2, 1
# We need to also remove the caption paragraphs ABOVE each table.
# Find caption paragraph that immediately precedes each table.

def find_caption_before_table(doc, table_idx):
    """Find the paragraph index immediately preceding the given table."""
    # docx body children include <w:p> and <w:tbl> elements in order
    body = doc.element.body
    children = list(body)
    # find the child element that corresponds to table[table_idx]
    target_tbl = doc.tables[table_idx]._element
    target_pos = None
    for i, child in enumerate(children):
        if child is target_tbl:
            target_pos = i
            break
    if target_pos is None or target_pos == 0:
        return None
    # walk backwards to find non-empty paragraph
    for j in range(target_pos - 1, -1, -1):
        if children[j].tag.endswith('}p'):
            # check text content
            text = ''.join(t.text or '' for t in children[j].iter() if t.tag.endswith('}t'))
            if text.strip():
                if text.strip().startswith('Table'):
                    return j
                else:
                    return None
            else:
                continue
        else:
            return None
    return None

for tbl_idx in to_delete_indices:
    # find caption paragraph
    tbl = ms.tables[tbl_idx]
    body = ms.element.body
    children = list(body)
    target_tbl = tbl._element
    target_pos = None
    for i, child in enumerate(children):
        if child is target_tbl:
            target_pos = i
            break
    # Find caption ABOVE
    caption_idx = None
    if target_pos is not None and target_pos > 0:
        for j in range(target_pos - 1, max(0, target_pos - 5), -1):
            if children[j].tag.endswith('}p'):
                text = ''.join(t.text or '' for t in children[j].iter() if t.tag.endswith('}t'))
                if text.strip().startswith('Table '):
                    caption_idx = j
                    break
                elif text.strip() == '':
                    continue
                else:
                    break

    # Delete table
    body.remove(target_tbl)
    print(f"  删除 Table {tbl_idx+1}")
    # Delete caption (now caption_idx still valid since we removed only the table after it)
    if caption_idx is not None:
        # re-fetch children since we removed one
        new_children = list(body)
        body.remove(new_children[caption_idx])
        print(f"    删除对应 caption")

# ============================================================
# 第三步: 把保留的 Tables 4, 7 重新编号为 Tables 1, 2
# 同时改 Table 4 caption 文字
# ============================================================
print("\n重新编号: Table 4 → Table 1, Table 7 → Table 2")

# After deletion, the remaining tables are old Table 4 (now index 0) and old Table 7 (now index 1)
ms_after = docx.Document(str(MS))  # reload to find current state? No, use the same ms object
# But we modified body directly; tables are rebuilt
# Actually let's reload after saving to avoid race conditions
ms.save(str(MS))
ms = docx.Document(str(MS))
print(f"  剩余表数: {len(ms.tables)}")

# Find caption paragraphs for the 2 tables
caption_replacements = [
    ('Table 4. Summary of Bayesian optimization results',
     'Table 1. Summary of Bayesian optimization results (Setup A)'),
    ('Table 7. Spatial attenuation model parameters (monthly calibration)',
     'Table 2. Spatial attenuation model parameters (monthly calibration)'),
]
for old, new in caption_replacements:
    for para in ms.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        if full_text.strip().startswith(old.split('.')[0] + '.'):  # Match "Table 4." or "Table 7."
            # Replace if we find the exact caption start
            if old in full_text:
                new_text = full_text.replace(old, new)
                runs = list(para.runs)
                if runs:
                    runs[0].text = new_text
                    for r in runs[1:]: r.text = ''
                print(f"  重命名 caption: '{old[:40]}...' → '{new[:40]}...'")
                break

# ============================================================
# 第四步: 替换正文中所有 "Table X" 引用
# ============================================================
print("\n替换正文中的 Table 引用...")

# 老 Table → 新位置
# 注意: 必须用 SI 编号 S27, S28, S29 (在 SI 末尾追加), 不能与既有 Table S2/S3 冲突
table_ref_mapping = [
    ('Table 1', 'Table S27'),  # data overview
    ('Table 2', 'Table S28'),  # emission/river-entry
    ('Table 3', 'Figure 3e'),  # closure coefficients (now in fig 3 panel e)
    ('Table 4', 'Table 1'),    # Bayesian summary
    ('Table 5', 'Figure 2b'),  # MCMC posteriors → Fig 2b
    ('Table 6', 'Figure 2d'),  # Sensitivity → Fig 2d
    ('Table 7', 'Table 2'),    # Spatial decay
    ('Table 8', 'Figure 4'),   # Emis vs Eff
]

# 注意: 必须用合理的边界匹配, 否则 "Table 1" 会匹配 "Table 10" 等. 用 \bTable N\b
import re

# 顺序: 先处理大编号 (避免 "Table 1" 替换吃掉 "Table 10")
table_ref_mapping_sorted = sorted(table_ref_mapping, key=lambda kv: -int(kv[0].split()[1]))

for para in ms.paragraphs:
    full_text = ''.join(run.text for run in para.runs)
    # Skip caption paragraphs themselves
    if full_text.strip().startswith('Table 1.') or full_text.strip().startswith('Table 2.'):
        # Don't process the new-numbered captions (their text is correct)
        continue
    if full_text.strip().startswith('Table S'):
        # SI table captions - skip
        continue

    new_text = full_text
    changed = False
    for old, new in table_ref_mapping_sorted:
        # Match \bTable N\b but not "Table NN" (e.g. Table 10)
        # Also not "Table SX" (SI tables)
        pattern = re.compile(rf'\b{re.escape(old)}\b(?!\d)')
        new_text2 = pattern.sub(new, new_text)
        if new_text2 != new_text:
            new_text = new_text2
            changed = True

    if changed:
        runs = list(para.runs)
        if runs:
            runs[0].text = new_text
            for r in runs[1:]: r.text = ''

ms.save(str(MS))
print("  ✓ 正文 Table 引用已替换")

# ============================================================
# 第五步: 重新嵌入 5-panel Figure 3
# ============================================================
print("\n嵌入新版 Figure 3 (5 panels)...")
ms = docx.Document(str(MS))
fig3_path = ROOT/'output/figures/figure3_monte_carlo/figure3_monte_carlo.png'

for i, p in enumerate(ms.paragraphs):
    has_img = any(run._element.xpath('.//w:drawing') for run in p.runs)
    if has_img and i > 0 and 'Figure 3.' in ms.paragraphs[i-1].text:
        for run in p.runs:
            run.text = ''
            for drawing in run._element.xpath('.//w:drawing'):
                drawing.getparent().remove(drawing)
        run = p.add_run()
        run.add_picture(str(fig3_path), width=Inches(6.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"  ✓ Figure 3 重新嵌入 (5-panel)")
        break

# 同时更新 Figure 3 caption
new_fig3_caption = ("Figure 3. Monte Carlo uncertainty propagation (10,000 iterations) and full-chain closure across pollutants. "
                    "(a–d) River-entry load probability density distributions for COD, NH₃-N, TN, and TP, "
                    "generated by sampling source emissions (CV = 20%, normal) and river-entry coefficients "
                    "(uniform [0.5, 1.5]× nominal). Red dashed line: monitored (S2-imputed) load; "
                    "blue dashed line: Bayesian MAP prediction; gray shading: 5–95% MC interval. "
                    "(e) Three-tier closure coefficients across pollutants — Entry/Emission (river-entry coefficient), "
                    "Monitor/Entry (channel transport), and Monitor/Emission (overall transport efficiency). "
                    "Overall transport efficiency varies from 6% (TP, particulate settling dominated) to 55% (TN, "
                    "most conservative behavior).")
for para in ms.paragraphs:
    full_text = ''.join(run.text for run in para.runs)
    if full_text.strip().startswith('Figure 3.'):
        runs = list(para.runs)
        if runs:
            runs[0].text = new_fig3_caption
            for r in runs[1:]: r.text = ''
        print(f"  ✓ Figure 3 caption 已更新")
        break

ms.save(str(MS))

# ============================================================
# 第六步: 把 Tables 1, 2, 3 加到 SI 末尾
# ============================================================
print("\n把原 Tables 1, 2, 3 加到 SI 末尾...")
si = docx.Document(str(SI))

# 在 SI 末尾追加新章节 + 表
si.add_paragraph('')  # spacer
section_p = si.add_paragraph('Additional reference tables (moved from main text)')
section_p.runs[0].font.bold = True
section_p.runs[0].font.size = Pt(11)

new_si_caption_texts = [
    'Table S27. Overview of research data (data type, content, resolution, volume).',
    'Table S28. Source-by-source emission (E) and river-entry (R) load estimates (kg) under Setup A. '
    'See Methods §2.4 and SI Table S7 for the underlying river-entry coefficients.',
    'Table S29. Full-chain coefficients linking emission, river-entry, and monitored loads. '
    'Entry/Emission represents the basin-scale river-entry coefficient; Monitor/Entry represents '
    'channel transport efficiency; Monitor/Emission represents overall transport. '
    'Overall efficiency ranges from 6% (TP) to 55% (TN).',
]

for (label, tbl_xml), caption in zip(tables_to_move, new_si_caption_texts):
    # add caption
    cap_p = si.add_paragraph(caption)
    cap_p.runs[0].font.size = Pt(10)
    # append table XML
    si.element.body.append(deepcopy(tbl_xml))
    si.add_paragraph('')  # spacer
    print(f"  ✓ 追加 {caption[:60]}...")

si.save(str(SI))
print(f"  保存: {SI}")

# ============================================================
# 验证
# ============================================================
print("\n" + "="*60)
print("验证")
print("="*60)
ms2 = docx.Document(str(MS))
print(f"Manuscript: {len(ms2.tables)} tables (target: 2)")
for i, t in enumerate(ms2.tables):
    first_row = [c.text.strip() for c in t.rows[0].cells]
    print(f"  Table {i+1}: {first_row[:4]}...")

si2 = docx.Document(str(SI))
print(f"\nSI: {len(si2.tables)} tables (was 25, now should be 28)")

print("\n完成。备份在 *.docx.tab_bak")
