#!/usr/bin/env python3
"""Run the final audit using the revision-5 figure command in the reproduction guide."""

from __future__ import annotations

from collections import Counter
from pathlib import Path
import json
import subprocess
import sys

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
RESULT = ROOT / "output" / "results" / "final_submission_audit_20260825.json"
REPORT = ROOT / "output" / "results" / "final_submission_audit_20260825.md"


def main():
    subprocess.run([sys.executable, str(ROOT / "scripts" / "reporting" / "audit_final_submission_20260825_v2.py")], check=False)
    data = json.loads(RESULT.read_text(encoding="utf-8"))
    guide = (ROOT / "REPRODUCIBILITY.md").read_text(encoding="utf-8")
    expected = ["revision3_diagnostics.py", "spatial_identifiability_unbounded.py", "revision5_mcmc_diagnostics.py", "generate_revision5_figures.py"]
    for item in data["checks"]:
        if item["check"] == "repository instructions cite existing scripts":
            item["status"] = "PASS" if all(name in guide and (ROOT / ("scripts/analysis/" + name if name.startswith(("revision", "spatial")) else "scripts/reporting/" + name)).exists() for name in expected) else "FAIL"
            item["detail"] = str(expected)
    ms = Document(ROOT / "paper" / "Manuscript_Final_MC-SIRC.docx")
    si = Document(ROOT / "paper" / "SI_Final.docx")
    response = Document(ROOT / "paper" / "Response_Letter.docx")

    def caption(document, prefix):
        return next(paragraph.text for paragraph in document.paragraphs if paragraph.text.startswith(prefix))

    multipanel_captions = {
        "Figure 3": caption(ms, "Figure 3."),
        "Figure 4": caption(ms, "Figure 4."),
        "Figure S1": caption(si, "Figure S1."),
    }
    panel_labels = ["(a) COD", "(b) NH₃-N", "(c) TN", "(d) TP"]
    captions_complete = all(
        all(label in text for label in panel_labels)
        for text in multipanel_captions.values()
    )
    data["checks"].append({
        "status": "PASS" if captions_complete else "FAIL",
        "check": "Figure 3/Figure 4/Figure S1 captions enumerate panels (a)–(d)",
        "detail": str({name: all(label in text for label in panel_labels) for name, text in multipanel_captions.items()}),
    })
    response_text = "\n".join(paragraph.text for paragraph in response.paragraphs)
    response_has_captions = all(text in response_text for text in multipanel_captions.values())
    data["checks"].append({
        "status": "PASS" if response_has_captions else "FAIL",
        "check": "response letter reproduces complete multipanel captions",
        "detail": str({name: text in response_text for name, text in multipanel_captions.items()}),
    })
    counts = Counter(item["status"] for item in data["checks"])
    data["summary"] = dict(counts)
    RESULT.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# Final submission audit — 2026-08-25", "", f"Summary: {dict(counts)}", "", "| Status | Check | Detail |", "|---|---|---|"]
    for item in data["checks"]:
        detail = str(item["detail"]).replace("|", "\\|").replace("\n", " ")
        lines.append(f"| {item['status']} | {item['check']} | {detail} |")
    REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(dict(counts))
    for item in data["checks"]:
        if item["status"] == "FAIL":
            print("FAIL", item["check"], item["detail"])
    print(REPORT)
    raise SystemExit(0 if counts["FAIL"] == 0 else 1)


if __name__ == "__main__":
    main()
